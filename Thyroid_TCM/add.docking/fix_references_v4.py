#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V4修正：
1. 删除幽灵引用[8][9][27][28]的正文标注及参考文献条目
2. 处理[15][16][17]不匹配引用（删除对应正文描述）
3. 合并[12][13]重复引用
4. 自动重新编号正文中的所有引用标注
"""

import shutil
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

PINK = RGBColor(0xFF, 0x69, 0xB4)

def clear_para_text(para):
    """Clear all runs in a paragraph."""
    for run in para.runs:
        run._element.getparent().remove(run._element)

def set_pink_para(para, text, font_size=11, bold=False):
    """Replace paragraph text with pink text."""
    clear_para_text(para)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = PINK
    para.paragraph_format.line_spacing = 1.5
    return run

def delete_paragraph(para):
    """Delete a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)

def build_mapping():
    """Build old citation number -> new citation number mapping."""
    deleted = {8, 9, 13, 15, 16, 17, 27, 28}
    mapping = {}
    new_num = 1
    for old_num in range(1, 45):
        if old_num in deleted:
            mapping[old_num] = None
        else:
            mapping[old_num] = new_num
            new_num += 1
    return mapping

def replace_citations_in_text(text, mapping):
    """Replace all citation numbers in text according to mapping."""
    def replacer(match):
        inner = match.group(1)
        parts = []
        for seg in inner.split(','):
            seg = seg.strip()
            if '-' in seg:
                start, end = seg.split('-')
                start, end = int(start), int(end)
                new_nums = []
                for n in range(start, end + 1):
                    if n in mapping:
                        new_n = mapping[n]
                        if new_n is not None:
                            new_nums.append(new_n)
                if not new_nums:
                    continue  # Skip this segment entirely
                if len(new_nums) == 1:
                    parts.append(str(new_nums[0]))
                else:
                    consecutive = all(new_nums[i] + 1 == new_nums[i+1] for i in range(len(new_nums)-1))
                    if consecutive:
                        parts.append(f"{new_nums[0]}-{new_nums[-1]}")
                    else:
                        parts.append(','.join(str(x) for x in new_nums))
            else:
                n = int(seg)
                if n in mapping:
                    new_n = mapping[n]
                    if new_n is not None:
                        parts.append(str(new_n))
                else:
                    parts.append(str(n))
        
        if not parts:
            return ""
        return "[" + ", ".join(parts) + "]"
    
    # Pattern matches [digits, spaces, commas, dashes]
    return re.sub(r'\[([\d\s,\-]+)\]', replacer, text)

def main():
    src = 'output/2024.12.21v2_docking_updated.docx'
    dst = 'output/2024.12.21v2_docking_updated_v4.docx'
    
    if os.path.exists(dst):
        os.remove(dst)
    shutil.copy2(src, dst)
    doc = Document(dst)
    
    mapping = build_mapping()
    print("Citation mapping (old -> new):")
    for old, new in sorted(mapping.items()):
        if new is None:
            print(f"  [{old}] -> DELETED")
        elif old != new:
            print(f"  [{old}] -> [{new}]")
    
    # ===================== Step 1: Modify in-text citations =====================
    print("\nModifying in-text citations...")
    
    citation_paragraphs = [13, 15, 17, 21, 62, 63, 64, 65, 68]
    
    for idx in citation_paragraphs:
        para = doc.paragraphs[idx]
        original_text = para.text.strip()
        new_text = replace_citations_in_text(original_text, mapping)
        
        # Special handling for paragraph 17: remove sentences with [15][16][17]
        if idx == 17:
            # Keep only the first sentence (formula composition)
            sentences = new_text.split('. ')
            # The first sentence ends with "Xia Ku Cao."
            # We need to be careful with split; let's reconstruct
            # Original: "...Xia Ku Cao. Clinical applications... [15]. Furthermore... [16]. Experimental... [17]."
            # After citation replacement, [15][16][17] are gone, so we need to identify and drop those sentences
            # Better: use the original first sentence
            first_sentence = "The Xiao Ying Tang formula, developed by renowned TCM expert Professor Xu Zhuyin, consists of Chai Hu, Yu Jin, Hou Po, Ban Xia, Dang Gui, Huang Qi, Bai Zhu, Cang Zhu, Ban Zhi Lian, Mao Zhua Cao, and Xia Ku Cao."
            new_text = first_sentence
        
        if new_text != original_text:
            set_pink_para(doc.paragraphs[idx], new_text)
            print(f"  Para {idx}: modified")
    
    # ===================== Step 2: Delete ghost reference entries =====================
    print("\nDeleting ghost reference entries from the list...")
    
    # Find reference list paragraph indices
    ref_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "References":
            ref_start = i
            break
    
    if ref_start is None:
        print("ERROR: Could not find References section!")
        return
    
    # Identify paragraphs to delete by content matching
    # We build a list of (index, paragraph) for the reference section
    ref_paras = []
    for i in range(ref_start + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text or text.startswith("Notes:") or text.startswith("Abbreviations:") or text.startswith("Figure"):
            break
        ref_paras.append((i, p, text))
    
    # Mapping from original position (1-based) to paragraph index
    # ref_paras[0] = [1], ref_paras[1] = [2], etc.
    # We need to delete positions: 8, 9, 13, 15, 16, 17, 27, 28
    delete_positions = [8, 9, 13, 15, 16, 17, 27, 28]
    to_delete = []
    for pos in delete_positions:
        idx = pos - 1  # 0-based index in ref_paras
        if 0 <= idx < len(ref_paras):
            to_delete.append(ref_paras[idx][1])  # paragraph object
            print(f"  Will delete ref [{pos}]: {ref_paras[idx][2][:80]}...")
    
    # Delete paragraphs (order doesn't matter for element removal)
    for para in to_delete:
        delete_paragraph(para)
    
    # ===================== Save =====================
    doc.save(dst)
    print(f"\nV4 document saved to: {dst}")
    print("Summary of changes:")
    print("  - Deleted in-text citations [8], [9], [15], [16], [17], [27], [28]")
    print("  - Merged duplicate [12][13] into [12]")
    print("  - Deleted 8 ghost/duplicate reference entries")
    print("  - Re-numbered all in-text citations to match new reference list")
    print("  - Removed 3 unsupported clinical claims from Introduction")

if __name__ == '__main__':
    main()
