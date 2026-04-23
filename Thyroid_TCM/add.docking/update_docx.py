#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Modify the docx file to add molecular docking result images and descriptions.
New content is marked in red.
"""

import shutil
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, font_name='Times New Roman', font_size=11, bold=False, color=None):
    """Set font properties for a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_red_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=11):
    """Add a paragraph with red text."""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold, color=RGBColor(0xFF, 0x00, 0x00))
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_normal_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=11):
    """Add a normal paragraph."""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_image_with_caption(doc, image_path, caption_text, width_inches=5.5):
    """Add an image with caption."""
    # Add image centered
    para_img = doc.add_paragraph()
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para_img.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    # Add caption in red
    para_cap = doc.add_paragraph()
    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = para_cap.add_run(caption_text)
    set_run_font(run_cap, font_size=10, bold=False, color=RGBColor(0xFF, 0x00, 0x00))
    para_cap.paragraph_format.space_after = Pt(12)
    return para_img, para_cap

def main():
    src = '2024.12.21v2.docx'
    dst = 'output/2024.12.21v2_docking_updated.docx'
    
    # Copy original file to output
    shutil.copy2(src, dst)
    
    # Open the copied document
    doc = Document(dst)
    
    # Find the index of paragraph containing "Figure5:" or last paragraph of Molecular Docking section
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('Figure5:') or text.startswith('Figure 5:'):
            insert_index = i + 1
            break
    
    if insert_index is None:
        # Fallback: insert after "Molecular Docking" section, before "Discussion"
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Discussion':
                insert_index = i
                break
    
    if insert_index is None:
        insert_index = len(doc.paragraphs)
    
    # We will build new elements and insert them at the right place.
    # python-docx doesn't support arbitrary insertion index well, so we use a workaround:
    # add all new paragraphs at the end, then move them to the correct position.
    
    # Collect all new elements
    new_elements = []
    
    # ---- New Content Start ----
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Detailed Molecular Docking Validation of β-Sitosterol')
    set_run_font(r, font_size=12, bold=True, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    new_elements.append(p)
    
    # Description paragraph 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('To further verify the binding stability of β-Sitosterol (CID 222284) with the core targets identified in this study, we performed molecular docking against three representative proteins: 4P5A, 6UNI, and 9AYG. The docking was carried out using AutoDock Vina with an exhaustiveness of 8, generating nine poses for each receptor-ligand pair. The lowest binding energy pose was selected for visualization and interaction analysis.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Description paragraph 2 - 4P5A
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('For the 4P5A protein target, β-Sitosterol achieved a binding energy of −9.23 kcal/mol for the top-ranked pose, which meets the criterion for strong binding affinity (< −7 kcal/mol). The global view (Figure 6A) shows that the ligand sits deeply within the hydrophobic pocket of the receptor. In the local interaction view (Figure 6B), the compound forms close contacts with surrounding residues including HIS-188, VAL-202, ARG-60, SER-59, HIS-61, GLY-62, HIS-87, GLU-66, ARG-88, and ASN-93. These residues are mainly concentrated in the catalytic cleft region, suggesting that β-Sitosterol may exert its inhibitory effect through occupation of the active site.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Figure 6A
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_4P5A_CID222284_20260423_133247/output_global.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 6A. Global view of β-Sitosterol docked into the 4P5A protein binding pocket (binding energy = −9.23 kcal/mol). The receptor is shown as gray cartoon and the ligand as green sticks.')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Figure 6B
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_4P5A_CID222284_20260423_133247/output_local.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 6B. Local interaction view of β-Sitosterol with key residues in the 4P5A binding pocket. Hydrophobic and polar contacts with residues HIS-188, VAL-202, ARG-60, SER-59, HIS-61, GLY-62, HIS-87, GLU-66, ARG-88, and ASN-93 are highlighted.')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Description paragraph 3 - 6UNI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('For the 6UNI protein target, the best docking pose yielded a binding energy of −9.12 kcal/mol. The global docking pose (Figure 7A) indicates that β-Sitosterol is anchored in a narrow groove formed by several α-helices. The local interaction map (Figure 7B) reveals that the ligand is surrounded by residues GLU-374, ARG-105, ILE-110, PHE-108, PHE-220, PHE-241, VAL-240, PRO-242, ILE-300, and PHE-304. A number of these residues are aromatic amino acids, which likely contribute to hydrophobic stacking with the sterol ring system of β-Sitosterol. This binding pattern implies that the compound can stabilize the receptor conformation through multiple van der Waals contacts.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Figure 7A
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_6UNI_CID222284_20260423_132950/output_global.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 7A. Global view of β-Sitosterol docked into the 6UNI protein binding pocket (binding energy = −9.12 kcal/mol).')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Figure 7B
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_6UNI_CID222284_20260423_132950/output_local.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 7B. Local interaction view of β-Sitosterol with key residues in the 6UNI binding pocket. Residues GLU-374, ARG-105, ILE-110, PHE-108, PHE-220, PHE-241, VAL-240, PRO-242, ILE-300, and PHE-304 are shown as green sticks.')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Description paragraph 4 - 9AYG
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('For the 9AYG protein target, β-Sitosterol displayed the strongest binding among the three receptors, with a docking score of −9.68 kcal/mol. The global pose (Figure 8A) demonstrates that the ligand is buried in a large cavity at the interface of two structural domains. The detailed interaction view (Figure 8B) shows that β-Sitosterol is flanked by residues LEU-924, ILE-210, PHE-928, GLU-927, LEU-203, LEU-204, PHE-1001, ALA-997, VAL-1005, GLU-223, and LEU-227. The presence of multiple leucine and phenylalanine residues around the ligand points to a predominantly hydrophobic binding environment, which is well suited to accommodate the lipophilic sterol scaffold of β-Sitosterol.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Figure 8A
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_9AYG_CID222284_20260423_132658/output_global.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 8A. Global view of β-Sitosterol docked into the 9AYG protein binding pocket (binding energy = −9.68 kcal/mol).')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Figure 8B
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('input/docking_9AYG_CID222284_20260423_132658/output_local.png', width=Inches(5.0))
    new_elements.append(p_img)
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run('Figure 8B. Local interaction view of β-Sitosterol with key residues in the 9AYG binding pocket. Residues LEU-924, ILE-210, PHE-928, GLU-927, LEU-203, LEU-204, PHE-1001, ALA-997, VAL-1005, GLU-223, and LEU-227 are displayed as green sticks.')
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Summary paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Taken together, the docking scores for all three protein targets fell well below the −7 kcal/mol threshold, and the top pose for 9AYG even reached −9.68 kcal/mol. These values support the notion that β-Sitosterol binds tightly to the selected targets. The interaction maps further show that the compound makes extensive contacts with conserved residues in each binding pocket, mainly through hydrophobic interactions and van der Waals forces. These computational findings lend additional support to the network pharmacology results presented above and provide a structural basis for the observed anti-thyroid cancer activity of XYT.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # ---- New Content End ----
    
    # Move new elements to the insert position
    # In python-docx, paragraphs are stored as <w:p> elements under the body.
    # We need to move each new paragraph element to just after the insert_index paragraph.
    body = doc.element.body
    # Find the reference element after which we want to insert
    ref_para = doc.paragraphs[insert_index - 1]._element
    
    for elem in new_elements:
        # elem is a Paragraph object; move its _element after ref_para
        body.append(elem._element)  # first append to end to keep it in the document
    
    # Now reorder: we want new elements right after ref_para
    # We can do this by moving elements one by one
    for elem in new_elements:
        ref_para.addnext(elem._element)
        ref_para = elem._element
    
    # Save
    doc.save(dst)
    print(f'Document saved to: {dst}')

if __name__ == '__main__':
    main()
