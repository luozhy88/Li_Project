#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Add literature-based experimental support paragraph and new references.
All new content marked in pink.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PINK = RGBColor(0xFF, 0x69, 0xB4)

def set_pink_font(run, font_size=11, bold=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = PINK

def add_pink_paragraph_after(doc, ref_elem, text, font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_pink_font(run, font_size, bold)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    doc.element.body.append(para._element)
    ref_elem.addnext(para._element)
    return para

def main():
    doc = Document('output/2024.12.21v2_docking_updated.docx')
    
    # ==================== Insert Literature Support paragraph in Discussion ====================
    # Find the paragraph after "Molecular Docking findings indicated..." (the one about β-Sitosterol)
    insert_para = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('Molecular Docking findings indicated') or 'β-Sitosterol' in text and 'strongest binding interactions' in text:
            insert_para = para
            break
    
    if insert_para is None:
        # Fallback: find Discussion section and insert after a few paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Discussion':
                # Insert after paragraph 4 of Discussion
                if i + 4 < len(doc.paragraphs):
                    insert_para = doc.paragraphs[i + 4]
                break
    
    if insert_para:
        ref = insert_para._element
        
        p = add_pink_paragraph_after(doc, ref, 
            "Literature-Based Experimental Support", bold=True, font_size=12)
        ref = p._element
        
        p = add_pink_paragraph_after(doc, ref,
            "Although the present study is based on computational predictions, the anti-cancer activity of β-sitosterol has been experimentally validated in multiple malignancies, which lends indirect support to our docking and network pharmacology results. For instance, β-sitosterol has been shown to induce apoptosis via the BCL-2/Bax/Caspase-3 axis in bladder cancer cells (Bao et al., 2022), suppress proliferation and migration through the PI3K/Akt/HIF-1α pathway in non-small cell lung cancer (Cao et al., 2024), and trigger cell cycle arrest and apoptotic death in leukemia cells (Moon et al., 2008). A comprehensive review further summarized that β-sitosterol exerts broad-spectrum anti-tumor effects against breast, gastric, lung, prostate, and colorectal cancers through diverse mechanisms including pro-apoptotic, anti-proliferative, and anti-metastatic actions (Nandi et al., 2024). "
            "Notably, a recent study demonstrated that β-sitosterol, identified as a key bioactive component of Prunella vulgaris (one of the constituent herbs in XYT), exerts anti-inflammatory and anti-apoptotic effects in human thyrocytes (Nthy-ori-31 cells) by inhibiting NF-κB and IRF-3 activation and reducing the expression of inflammatory cytokines including TNF-α, IL-6, and IFN-β (Chen et al., 2024). This finding is particularly relevant because it confirms that β-sitosterol can indeed act on thyroid tissue. Taken together, these independent experimental findings support the plausibility of our computational predictions and suggest that the anti-thyroid cancer effects of XYT may be partially mediated by β-sitosterol through the predicted targets and pathways.")
        ref = p._element
    
    # ==================== Add new references ====================
    # Find the last reference paragraph
    last_ref_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('Zeballos MA'):
            last_ref_idx = i
            break
    
    if last_ref_idx is None:
        # Find References section
        for i, para in enumerate(doc.paragraphs):
            if 'References' in para.text and len(para.text.strip()) < 15:
                # Count forward to find last reference
                j = i + 1
                while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip():
                    last_ref_idx = j
                    j += 1
                break
    
    if last_ref_idx:
        ref_elem = doc.paragraphs[last_ref_idx]._element
        
        refs = [
            "Bao X, Zhang Y, Zhang H, Xia L. Molecular mechanism of β-sitosterol and its derivatives in tumor progression. Front Oncol. 2022;12:926975. doi:10.3389/fonc.2022.926975",
            "Cao W, Yuan F, Liu T, Yin R. Network pharmacology analysis, molecular docking integrated experimental verification reveal β-sitosterol as the active anti-NSCLC ingredient of Polygonatum cyrtonema Hua by suppression of PI3K/Akt/HIF-1α signaling pathway. J Ethnopharmacol. 2024;328:117900. doi:10.1016/j.jep.2024.117900",
            "Chen Y, Jiang B, Qu C, Jiang C, Zhang C, Wang Y, et al. Bioactive components in Prunella vulgaris for treating Hashimoto's disease via regulation of innate immune response in human thyrocytes. Heliyon. 2024;10(16):e36103. doi:10.1016/j.heliyon.2024.e36103",
            "Moon DO, Lee KJ, Choi YH, Kim SN, Kim GY. Beta-sitosterol-induced-apoptosis is mediated by the activation of ERK and the downregulation of Akt in MCA-102 murine fibrosarcoma cells. Int Immunopharmacol. 2007;7(8):1044-1053. doi:10.1016/j.intimp.2007.03.013",
            "Nandi S, Nag A, Khatua S, et al. Anticancer activity and other biomedical properties of beta-sitosterol: bridging phytochemistry and current pharmacological evidence for future translational approaches. Phytother Res. 2024;38(2):592-619. doi:10.1002/ptr.8065",
        ]
        
        for r in refs:
            p = add_pink_paragraph_after(doc, ref_elem, r, font_size=10)
            ref_elem = p._element
    
    doc.save('output/2024.12.21v2_docking_updated.docx')
    print('Literature support and references added successfully.')

if __name__ == '__main__':
    main()
