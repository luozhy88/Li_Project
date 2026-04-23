#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Modify the docx file to add composite molecular docking image.
New content is marked in red.
"""

import shutil
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, font_name='Times New Roman', font_size=11, bold=False, color=None):
    """Set font properties for a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def main():
    src = '2024.12.21v2.docx'
    dst = 'output/2024.12.21v2_docking_updated.docx'
    
    # Remove old output if exists
    if os.path.exists(dst):
        os.remove(dst)
    
    # Copy original file to output
    shutil.copy2(src, dst)
    
    # Open the copied document
    doc = Document(dst)
    
    # Find the index of paragraph containing "Figure5:" or last paragraph of Molecular Docking section
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith('Figure5:') or text.startswith('Figure 5:'):
            insert_index = i + 1
            break
    
    if insert_index is None:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Discussion':
                insert_index = i
                break
    
    if insert_index is None:
        insert_index = len(doc.paragraphs)
    
    # Collect all new elements
    new_elements = []
    
    # ---- New Content Start ----
    
    # Section header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Detailed Molecular Docking Validation of β-Sitosterol')
    set_run_font(r, font_size=12, bold=True, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    new_elements.append(p)
    
    # Description paragraph 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('To verify the binding stability of β-Sitosterol (CID 222284) with the core targets, we performed molecular docking against three representative proteins: 4P5A, 6UNI, and 9AYG. The docking was carried out using AutoDock Vina with an exhaustiveness of 8, generating nine poses for each receptor-ligand pair. The lowest binding energy pose was selected for visualization and interaction analysis.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Description paragraph 2 - 4P5A
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Row A presents the docking results for the 4P5A protein target. β-Sitosterol achieved a binding energy of −9.23 kcal/mol for the top-ranked pose, satisfying the strong binding criterion (< −7 kcal/mol). The left panel shows the local interaction view: the compound sits deeply within the hydrophobic pocket and makes close contacts with surrounding residues including HIS-188, VAL-202, ARG-60, SER-59, HIS-61, GLY-62, HIS-87, GLU-66, ARG-88, and ASN-93. These residues cluster in the catalytic cleft region, indicating that β-Sitosterol may exert inhibitory effects through active-site occupation. The right panel provides the global view of the ligand positioned inside the receptor structure.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Description paragraph 3 - 6UNI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Row B shows the docking results for the 6UNI protein target. The best pose yielded a binding energy of −9.12 kcal/mol. In the left panel, the ligand is anchored in a narrow groove formed by several α-helices, surrounded by residues GLU-374, ARG-105, ILE-110, PHE-108, PHE-220, PHE-241, VAL-240, PRO-242, ILE-300, and PHE-304. Many of these residues are aromatic amino acids, which likely contribute to hydrophobic stacking with the sterol ring system of β-Sitosterol. The right panel shows the overall position of the ligand within the protein fold.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Description paragraph 4 - 9AYG
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Row C displays the docking results for the 9AYG protein target, where β-Sitosterol showed the strongest binding among the three receptors with a score of −9.68 kcal/mol. The left panel reveals that the ligand is flanked by residues LEU-924, ILE-210, PHE-928, GLU-927, LEU-203, LEU-204, PHE-1001, ALA-997, VAL-1005, GLU-223, and LEU-227. The abundance of leucine and phenylalanine residues around the ligand points to a predominantly hydrophobic binding environment, well suited to accommodate the lipophilic sterol scaffold. The right panel confirms that the ligand is buried in a large cavity at the interface of two structural domains.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # Insert composite image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture('output/figure_composite_docking.png', width=Inches(6.0))
    new_elements.append(p_img)
    
    # Figure caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(
        'Figure 6. Integrated molecular docking results of β-Sitosterol with three target proteins. '
        'Rows A–C correspond to proteins 4P5A, 6UNI, and 9AYG, respectively. '
        'The left column shows local interaction views (green sticks: ligand and surrounding residues; gray cartoon: receptor backbone); '
        'the right column shows global docking poses with the ligand embedded in the protein structure. '
        'Binding energies for the top-ranked poses are −9.23, −9.12, and −9.68 kcal/mol, respectively.'
    )
    set_run_font(r, font_size=10, color=RGBColor(0xFF, 0x00, 0x00))
    p_cap.paragraph_format.space_after = Pt(12)
    new_elements.append(p_cap)
    
    # Summary paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Taken together, the docking scores for all three protein targets fell well below the −7 kcal/mol threshold, and the top pose for 9AYG even reached −9.68 kcal/mol. These values support tight binding of β-Sitosterol to the selected targets. The interaction maps further show that the compound makes extensive contacts with conserved residues in each binding pocket, mainly through hydrophobic interactions and van der Waals forces. These computational findings lend additional support to the network pharmacology results and provide a structural basis for the observed anti-thyroid cancer activity of XYT.')
    set_run_font(r, font_size=11, color=RGBColor(0xFF, 0x00, 0x00))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    new_elements.append(p)
    
    # ---- New Content End ----
    
    # Move new elements to the insert position
    body = doc.element.body
    ref_para = doc.paragraphs[insert_index - 1]._element
    
    for elem in new_elements:
        body.append(elem._element)
    
    for elem in new_elements:
        ref_para.addnext(elem._element)
        ref_para = elem._element
    
    # Save
    doc.save(dst)
    print(f'Document saved to: {dst}')

if __name__ == '__main__':
    main()
