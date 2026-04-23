#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V3版本修改：
1. 修正前后不一致（删除3CLpro、gut microbiome、machine learning等错误内容）
2. 修正语言错误和AI痕迹
3. 删除Discussion中未验证的蛋白提及（2AR9, 7KYO）
4. 新增Limitations段落
5. 所有修改/新增内容标黄色
"""

import shutil
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

YELLOW = RGBColor(0xFF, 0xFF, 0x00)

def clear_para_text(para):
    """Clear all runs in a paragraph."""
    for run in para.runs:
        run._element.getparent().remove(run._element)

def add_yellow_text(para, text, font_size=11, bold=False, alignment=None):
    """Add yellow text to a paragraph, clearing existing content."""
    clear_para_text(para)
    if alignment is not None:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = YELLOW
    para.paragraph_format.line_spacing = 1.5
    return run

def set_yellow_para(para, text, font_size=11, bold=False):
    """Replace paragraph text with yellow text."""
    add_yellow_text(para, text, font_size, bold)

def add_yellow_paragraph_after(doc, ref_para, text, font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a new yellow paragraph after ref_para."""
    new_para = doc.add_paragraph()
    new_para.alignment = alignment
    run = new_para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = YELLOW
    new_para.paragraph_format.line_spacing = 1.5
    new_para.paragraph_format.space_after = Pt(6)
    
    # Move to correct position
    doc.element.body.append(new_para._element)
    ref_para.addnext(new_para._element)
    return new_para

def delete_paragraph(para):
    """Delete a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)

def main():
    src = '2024.12.21v2.docx'
    dst = 'output/2024.12.21v2_docking_updated.docx'
    
    if os.path.exists(dst):
        os.remove(dst)
    shutil.copy2(src, dst)
    doc = Document(dst)
    
    # ===================== 修改段落 =====================
    
    # 段落5: Purpose
    set_yellow_para(doc.paragraphs[5], 
        "Purpose: This study aims to use network pharmacology to compare the active components, key targets, mechanisms of action, and related pathways of three classic traditional Chinese medicine (TCM) formulations—Xiao Ying Tang (XYT), Zhe Bei Mu (ZBM), and Xia Ku Cao (XKC)—in the treatment of thyroid cancer.")
    
    # 段落6: Methods - 删除3CLpro
    set_yellow_para(doc.paragraphs[6], 
        "Methods: We employed network pharmacology and molecular docking to identify the most therapeutically effective active components and analyze their mechanisms of action.")
    
    # 段落7: Results - 重写
    set_yellow_para(doc.paragraphs[7], 
        "Results: ZBM had the fewest action targets (78 targets, 1%), followed by XKC (842 targets, 6%), while XYT had the most (1595 targets, 12%). Gene Ontology analysis indicated that XYT and XKC share similarities in biological processes, molecular functions, and cellular components, particularly in voltage-gated calcium channel complexes. These genes play roles in angiogenesis, calcium homeostasis, steroid hormone responses, and oxidative stress responses. Elevated reactive oxygen species (ROS) are linked to thyroid cancer development because oxidative stress causes DNA damage and genomic instability. In addition, transient receptor potential canonical (TRPC) channels were found to contribute significantly to cancer cell growth. The enriched pathways associated with XYT and XKC included immune response, signal transduction, and apoptosis.")
    
    # 段落8: Conclusion
    set_yellow_para(doc.paragraphs[8], 
        "Conclusion: This study highlights the potential of Xiao Ying Tang (XYT) in the treatment of thyroid cancer because of its strong binding affinities and interactions with key target proteins. Among the three formulations examined, XYT showed the highest gene enrichment and binding affinity, suggesting a notable therapeutic effect. The integration of network pharmacology and molecular docking provides useful insights into the mechanisms of traditional Chinese medicine in thyroid cancer treatment and may guide future research.")
    
    # 段落12: 修正表达
    set_yellow_para(doc.paragraphs[12], 
        "According to data from China's National Cancer Center, thyroid cancer ranked seventh among all malignant tumors in terms of incidence in 2015, and fourth among female malignant tumors. Overall, the prognosis for thyroid cancer in China is relatively good, with a five-year survival rate of 84.3%. However, this is still far below the rate in developed nations such as the United States, where the five-year survival rate is reported to be 98% [2].")
    
    # 段落13: 修正
    set_yellow_para(doc.paragraphs[13], 
        "Currently, Western medicine lacks a perfect curative therapy. The primary treatment strategy is surgical excision of the thyroid gland combined with radioactive iodine therapy or lifelong thyroid hormone replacement medication to control symptoms [3]. Nevertheless, long-term medication can cause side effects that produce considerable discomfort for patients, who also face risks of persistent or recurrent cancer along with local and distant metastases [4-5]. Even molecular targeted drugs such as Dabrafenib, Trametinib, Selpercatinib, and Pralsetinib have limited practical use because of their high cost and uncertain long-term effectiveness [6-9]. Therefore, it is essential to actively pursue effective treatments for thyroid cancer and investigate the underlying regulatory mechanisms.")
    
    # 段落15: 修正
    set_yellow_para(doc.paragraphs[15], 
        "Xia Ku Cao has been shown to exert bidirectional regulatory effects on the immune system; its aqueous extract affects NF-κB and MAP kinase activity and stimulates macrophage activity [11]. Modern pharmacological studies have confirmed that Zhe Bei Mu contains various alkaloids, saponins, flavonoids, and polysaccharides with effects including phlegm elimination, cough relief, anti-inflammatory activity, pain relief, and anti-tumor activity, as well as the ability to reverse tumor resistance [12]. Network pharmacology analyses of Zhe Bei Mu have identified 2,357 disease-related targets and 96 core targets linked to thyroid cancer, including key genes such as NPM1, HSPA5, and HDAC5 [13].")
    
    # 段落16: 修正
    set_yellow_para(doc.paragraphs[16], 
        "Researchers have conducted network pharmacology and molecular docking studies on Zhe Bei Mu and Xia Ku Cao, identifying 108 targets related to the treatment of thyroid nodules. Overlapping targets such as AKT1, tumor protein p53 (TP53), and mitogen-activated protein kinase 1 (MAPK1) were identified as core targets. Both herbs may treat thyroid cancer through multiple components, multiple targets, and multiple pathways [14].")
    
    # 段落17: 修正
    set_yellow_para(doc.paragraphs[17], 
        "The Xiao Ying Tang formula, developed by renowned TCM expert Professor Xu Zhuyin, consists of Chai Hu, Yu Jin, Hou Po, Ban Xia, Dang Gui, Huang Qi, Bai Zhu, Cang Zhu, Ban Zhi Lian, Mao Zhua Cao, and Xia Ku Cao. Clinical applications by Pu Chunli have demonstrated that this formula significantly alleviates neck discomfort in patients while reducing the volume and maximum diameter of thyroid nodules [15]. Furthermore, Meng Dali's clinical trials found that this formula can lower thyroid autoantibody levels while regulating CD4+ and CD8+ ratios to improve immune dysregulation [16]. Experimental studies have confirmed that this formula significantly inhibits follicular thyroid carcinoma by downregulating the HIF-1α/VEGF signaling pathway to suppress angiogenesis [17].")
    
    # 段落19: Materials - 修正 Gene Cards 和 thyroid nodule-related
    set_yellow_para(doc.paragraphs[19], 
        "Metabolite data for Fritillaria thunbergii (ZBM), Prunella vulgaris (XKC), and the Xiao Ying Tang formula (XYT) were obtained from the Traditional Chinese Medicine Systems Pharmacology (TCMSP) database [18]. The chemical identifiers (CIDs) for these metabolites were retrieved from PubChem [19]. Using these CIDs, gene information for the metabolites was collected from the MetaBat, BindingDB, and Guide to Pharmacology databases [20]. GeneCards was also used to obtain genes associated with thyroid cancer [21]. The intersection of genes related to TCM metabolites and thyroid cancer-related genes was then identified for further analysis.")
    
    # 段落20: Methods - 修正 enrich EGG, Scrub.py, my SMILES
    set_yellow_para(doc.paragraphs[20], 
        "Using the enrichGO and enrichKEGG functions from the clusterProfiler package in R, Gene Ontology (GO) and Kyoto Encyclopedia of Genes and Genomes (KEGG) enrichment analyses were performed on the intersecting genes [22]. The enriched biological processes, pathways, and functions linked to thyroid cancer and metabolites were visualized based on these results (Figures 3 and 4). Significantly enriched terms were identified using a significance threshold of p < 0.05 [23]. The SMILES data derived from CIDs were converted into Structure Data Format (SDF) files using Open Babel for molecular docking experiments [24]. These SDF files were used as ligand inputs for docking simulations. Protein structures corresponding to the intersecting genes were retrieved from the Protein Data Bank (PDB) through the RCSB website [25]. The binding affinities between metabolites and target proteins were predicted using AutoDock Vina [26].")
    
    # 段落21: 修正
    set_yellow_para(doc.paragraphs[21], 
        "Bar graphs were used to display the results of the KEGG and GO enrichment analyses, as shown in Figures 3 and 4. These figures provide insights into possible mechanisms of action by highlighting key biological pathways and processes involved in metabolite-disease relationships [27-28].")
    
    # 段落26: Results开头 - 完全重写，删除gut microbiome和machine learning
    set_yellow_para(doc.paragraphs[26], 
        "To identify potential therapeutic targets and mechanisms, we systematically analyzed the overlapping genes between the three TCM formulations and thyroid cancer. The following sections present the results of target screening, functional enrichment analysis, and molecular docking validation.")
    
    # 段落28: Venn diagram结果 - 修正
    set_yellow_para(doc.paragraphs[28], 
        "Venn diagrams were constructed for the three traditional Chinese medicine (TCM) formulations—ZBM, XKC, and XYT—to compare disease-related genes with TCM-related genes. For ZBM (Figure 2A), 12,909 disease-related genes and 29 TCM-related genes were identified; 78 of these genes overlapped, accounting for 1% of the total. For XKC (Figure 2B), 12,145 disease-related genes and 252 TCM-related genes were identified, with 842 overlapping genes (6% of the total). For XYT (Figure 2C), 11,392 disease-related genes and 536 TCM-related genes were identified, with 1,595 overlapping genes (12% of the total). These overlapping genes represent potential therapeutic targets associated with each TCM formulation.")
    
    # 段落34: GO结果开头
    set_yellow_para(doc.paragraphs[34], 
        "To comprehensively understand the functional differences among the ZBM, XKC, and XYT groups, we compared the three groups across Biological Processes (BP), Molecular Functions (MF), and Cellular Components (CC) based on GO enrichment analysis results.")
    
    # 段落35: GO-BP
    set_yellow_para(doc.paragraphs[35], 
        "As illustrated in Figure 3, the XYT group exhibited the highest representation across most entries, particularly in key biological processes related to tumors, inflammation, and proliferation. For example, within biological processes (BP), XYT dominated in categories such as response to oxidative stress, response to steroid hormones, calcium ion homeostasis, and vascular processes in the circulatory system. The XKC group followed a similar trend but with lower representation than XYT, while the ZBM group had the lowest representation across most entries.")
    
    # 段落36: GO-MF
    set_yellow_para(doc.paragraphs[36], 
        "In terms of molecular functions (MF), the XYT group again led in several entries, especially in voltage-gated ion channel activity, transcription factor binding, and oxidoreductase activity. The XKC group ranked second to XYT, whereas ZBM had the fewest gene enrichments.")
    
    # 段落37: GO-CC
    set_yellow_para(doc.paragraphs[37], 
        "Regarding cellular components (CC), XYT dominated multiple entries related to cell membranes, synapses, and ion channels, such as voltage-gated calcium channel complexes and synaptic membranes. The performance of XKC was similar to that of XYT; however, ZBM showed a significantly lower number of associated genes.")
    
    # 段落38: GO总结
    set_yellow_para(doc.paragraphs[38], 
        "Overall, the XYT group had the highest representation across most entries, followed by XKC and then ZBM. This indicates that XYT has the highest degree of gene enrichment in functional annotations and may exert a more substantial impact on biological processes related to tumors, inflammation, and proliferation. Additionally, XYT and XKC exhibited similar trends across multiple entries, suggesting a degree of functional similarity between these two formulations.")
    
    # 段落40: KEGG开头
    set_yellow_para(doc.paragraphs[40], 
        "To gain a comprehensive understanding of functional differences among ZBM, XKC, and XYT within signaling pathways, we conducted comparative analyses based on KEGG enrichment results.")
    
    # 段落41: KEGG结果1
    set_yellow_para(doc.paragraphs[41], 
        "According to the results, the XYT group had the highest enrichment in the majority of signaling pathways, especially those related to signal transduction and apoptosis. ZBM showed comparatively lower enrichment across most pathways, whereas XKC had the next-highest overall enrichment level.")
    
    # 段落42: KEGG癌症通路
    set_yellow_para(doc.paragraphs[42], 
        "In cancer-related pathways—such as PD-L1 expression and PD-1 checkpoint pathways, as well as various specific cancer pathways—the XYT group displayed the highest enrichment levels, followed by XKC. This suggests that both groups may play significant roles in regulating tumor-associated pathways.")
    
    # 段落43: KEGG免疫和增殖
    set_yellow_para(doc.paragraphs[43], 
        "Both XYT and XKC showed comparable enrichment patterns in pathways linked to the immune system and inflammatory response, which were noticeably greater than those seen in ZBM. This pattern was especially evident in pathways related to rheumatoid arthritis and autoimmune disorders. In proliferation-related pathways such as cell growth and death, as well as signaling molecules and interactions, XYT again demonstrated the highest enrichment levels while XKC closely followed. This indicates that both groups may share similar functional characteristics in regulating cell proliferation.")
    
    # 段落44: KEGG总结
    set_yellow_para(doc.paragraphs[44], 
        "XYT and XKC showed comparable enrichment patterns across a variety of signaling pathways, including those related to metabolism, signal transduction, and immune responses, which suggests a functional resemblance between these two formulations. ZBM, on the other hand, displayed lower enrichment values across the majority of pathways, indicating that its contribution to related biological activities may be relatively limited.")
    
    # 段落47: Molecular Docking开头
    set_yellow_para(doc.paragraphs[47], 
        "The analysis reveals significant differences in binding energies among compounds from ZBM, XKC, and XYT across 20 compound-protein pairs. Molecular docking and interaction studies between core gene target proteins and TCM components were conducted using AutoDock Vina. It is generally accepted that the lower the binding energy between a ligand and a receptor, the stronger the binding affinity. Binding energies less than 0 kcal/mol indicate spontaneous binding, values below −5 kcal/mol suggest good binding affinity, and values below −7 kcal/mol indicate a highly active binding conformation.")
    
    # 段落48: 对接结果比较
    set_yellow_para(doc.paragraphs[48], 
        "Compared with the other compounds, XYT had the lowest free energy of binding to proteins and the strongest binding capabilities, as it was most frequently represented across all entries. XKC ranked second, also showing strong presence across many entries. ZBM, in contrast, was the least represented, suggesting weaker binding capacity and relatively higher binding energy.")
    
    # 段落49: β-Sitosterol具体对接
    set_yellow_para(doc.paragraphs[49], 
        "Specifically, among the compound-protein interactions examined, β-Sitosterol (CID 222284) showed strong binding affinity with proteins 6UNI, 9AYG, and 4P5A. XYT demonstrated the best performance at these protein targets, while ZBM showed minimal representation, reflecting its weaker binding energy.")
    
    # 段落50: 对接模式总结
    set_yellow_para(doc.paragraphs[50], 
        "Furthermore, XYT and XKC showed consistent patterns in binding to proteins, suggesting that their functional mechanisms may overlap. These two formulations may therefore have comparable roles in certain biological processes or pathways related to thyroid cancer. In summary, XYT appears to have the strongest functional impact among the three compound types, followed by XKC, while ZBM's contribution seems relatively minor. This discrepancy may arise from their varying binding capabilities at different protein targets.")
    
    # 段落53: Discussion - 删除3CLpro
    set_yellow_para(doc.paragraphs[53], 
        "This study employed network pharmacology to compare the active components, key targets, mechanisms of action, and related pathways of three classic traditional Chinese medicine formulations for the treatment of thyroid tumors: XYT, XKC, and ZBM. We also identified the most therapeutically effective active components through molecular docking to analyze the mechanisms by which these traditional formulations treat thyroid cancer.")
    
    # 段落54: Discussion - 修正大量语法
    set_yellow_para(doc.paragraphs[54], 
        "The findings showed that ZBM (78 targets, 1%), XKC (842 targets, 6%), and XYT (1595 targets, 12%) differed substantially in the number of action targets. Gene Ontology analysis indicated that XYT and XKC are comparable across biological processes (BP), molecular functions (MF), and cellular components (CC). These genes are essential components of voltage-gated calcium channel complexes and are involved in regulating transcription factor binding, oxidoreductase activity, and voltage-gated ion channel activity. They play important roles in calcium homeostasis, angiogenesis, steroid hormone responses, and oxidative stress responses. Reactive oxygen species (ROS) are linked to cancer and other chronic inflammatory disorders. Under normal conditions, the production of ROS and the maintenance of antioxidants are balanced to preserve cellular redox state. Thyroid cancer and oxidative responses are related; increased ROS levels may accelerate the growth of malignant thyroid cancer [29]. Oxidative stress induces DNA damage and genomic instability, sustains proliferation pathways and cell survival, and promotes angiogenesis and metastasis, thereby supporting multiple stages of tumor development. Thyroid cancer tissues exhibit higher levels of thyroid hormones compared with normal thyroid tissue [30].")
    
    # 段落55: TRPC - 修正
    set_yellow_para(doc.paragraphs[55], 
        "Transient receptor potential canonical (TRPC) channels are non-selective calcium channels activated by various stimuli that allow calcium ions to enter cells. Several TRPC channels contribute to cancer cell proliferation. Studies have shown that TRPC1 and TRPC2 influence the migration and adhesion of follicular thyroid cancer cells as well as RAC, calcium, and MMP2 [31]. These ion channels are also widely present in vascular endothelial cells; calcium-regulating proteins can modulate calcium ion influx from these channels into endothelial cells, subsequently regulating tumor angiogenesis [32]. Thus, Xiao Ying Tang and Xia Ku Cao may play significant roles in regulating redox reactions and inhibiting calcium channels to suppress angiogenesis in thyroid cancer.")
    
    # 段落56: 凋亡通路
    set_yellow_para(doc.paragraphs[56], 
        "XYT and XKC target genes are associated with enriched pathways such as apoptosis, signal transduction, immune system regulation, and pathological responses. Abnormal apoptotic pathways within tumor cells can impede the normal clearance of abnormal cells. The Bcl-2-associated gene BAG-1 is a unique anti-apoptotic gene that binds to the Bcl-2 protein, inhibiting the release of apoptotic factors such as cytochrome c and thereby exerting anti-apoptotic effects. Research indicates that 74% of thyroid cancer cells express BAG-1, which is significantly higher than in thyroid adenomas or normal thyroid tissues [33]. Recent studies have used nanomaterials to induce apoptosis for anti-tumor benefits [34].")
    
    # 段落57: 分子对接讨论 - 删除2AR9和7KYO（无数据支持）
    set_yellow_para(doc.paragraphs[57], 
        "Molecular docking findings indicated a hierarchical binding affinity pattern, with ZBM having the lowest affinity, followed by XKC, and XYT showing the strongest binding. Among the compounds examined, β-Sitosterol (CID 222284) had the strongest binding interactions with protein structures 6UNI, 9AYG, and 4P5A, demonstrating favorable binding affinities. The top-ranked binding molecules are engaged in several pathways, including those involving CYP3A4, CACNA1H, and POLB. CYP3A4 is an important enzyme for the metabolism of chemotherapeutic drugs in thyroid malignancies; for example, Vandetanib, a tyrosine kinase inhibitor used to treat medullary thyroid cancer, is largely metabolized by CYP3A4 [35]. CYP3A4 expression is regulated by several signaling mechanisms. PTHrP, a regulatory factor in cancer cachexia, reduces CYP3A4 production through the cAMP/PKA/PKC/NF-κB pathway. Inhibiting these pathways can restore this suppression [36].")
    
    # 段落58: 信号通路
    set_yellow_para(doc.paragraphs[58], 
        "Major signaling pathways in thyroid cancer include MAPK, PI3K, Wnt/β-catenin, Shh, Keap1/Nrf2, JAK-STAT, and NF-κB. Wnt/β-catenin signaling activation promotes the growth of thyroid carcinoma by stabilizing β-catenin [37]. Keap1/Nrf2 signaling is involved in both benign and malignant thyroid disorders and can serve as a prognostic indicator or therapeutic target [38]. The activation of MAPK signaling is closely associated with several forms of thyroid malignancy, notably papillary thyroid carcinoma (PTC) and follicular thyroid carcinoma (FTC), making MAPK and PI3K pathways key activation mechanisms in thyroid tumorigenesis [39].")
    
    # 段落60: Conclusion
    set_yellow_para(doc.paragraphs[60], 
        "This study investigated the potential mechanisms of XYT in the treatment of thyroid cancer using network pharmacology. Our findings indicate that active components such as β-Sitosterol may produce therapeutic benefits by interacting with target proteins including 6UNI, 9AYG, and 4P5A. Molecular docking experiments revealed that the major active components had high binding affinities for the target proteins, providing a basis for further research. Our network pharmacology analysis and molecular docking validation suggest that the pharmacodynamic effects of β-Sitosterol are likely mediated through interactions between its active molecules and target proteins. However, its pharmacokinetic properties require additional investigation. While this research provides interesting insights, it is important to recognize its limitations. Further pharmacological and clinical studies are needed to confirm our findings.")
    
    # 段落64: Acknowledgments - 修正He meticulous
    set_yellow_para(doc.paragraphs[64], 
        "We would like to express our sincere gratitude to all those who contributed to this research. First and foremost, we thank Dr. Pengfei Yang, who assisted with data collection and analysis. His meticulous attention to detail was crucial in ensuring the accuracy of our findings. Additionally, we wish to acknowledge the financial support provided by the construction project of the inheritance studio of Wang Wenchun, a nationally renowned expert in traditional Chinese medicine, under the National Administration of Traditional Chinese Medicine (Document No. Guo Zhong Yi Yao Ren Jiao Fa [2010] 59); Shanghai Clinical Research Center of Traditional Chinese Medicine Oncology Science and Technology Commission of Shanghai Municipality (21MC1930500).")
    
    # ===================== 新增 Limitations 段落 =====================
    # 在Conclusion段落(60)之后、Ethics段落(61)之前插入
    ref_para = doc.paragraphs[60]._element
    
    lim_title = add_yellow_paragraph_after(doc, ref_para, "Limitations", bold=True, font_size=12)
    ref_para = lim_title._element
    
    lim_text = add_yellow_paragraph_after(doc, ref_para, 
        "Several limitations of this study should be acknowledged. First, the findings are based entirely on computational predictions from network pharmacology and molecular docking, without experimental validation at the cellular or animal level. Second, the TCMSP database has known limitations in data completeness and accuracy; therefore, the identified active components and targets may not fully represent the true pharmacological profile of the studied formulations. Third, molecular docking provides only a static snapshot of ligand-receptor interactions and does not account for dynamic conformational changes, solvent effects, or protein flexibility. Fourth, this study focused on only three representative protein targets for docking validation, which may not capture the full complexity of the polypharmacology involved. Finally, the clinical relevance of the predicted pathways and binding interactions remains to be confirmed through in vitro and in vivo experiments, as well as clinical trials.")
    ref_para = lim_text._element
    
    # ===================== 插入整合图（保留之前的红色内容，但这里我们重新从头做）=====================
    # 实际上，V3应该是在V2的基础上继续修改。但用户要求V3的修改标黄色。
    # 由于我们已经修改了原文段落为黄色，现在需要把V2中的红色对接内容保留，并新增/修改的部分标黄色。
    # 为了简化，我重新从头构建V3：基于原文，先做黄色修改，然后再在Figure5之后插入对接内容（标黄色）。
    
    # 先找到Figure5的位置
    fig5_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith('Figure5:') or para.text.strip().startswith('Figure 5:'):
            fig5_idx = i
            break
    
    if fig5_idx:
        ref = doc.paragraphs[fig5_idx]._element
        
        # 插入对接内容（标黄色）
        p = add_yellow_paragraph_after(doc, ref, "Detailed Molecular Docking Validation of β-Sitosterol", bold=True, font_size=12)
        ref = p._element
        
        p = add_yellow_paragraph_after(doc, ref, 
            "To verify the binding stability of β-Sitosterol (CID 222284) with the core targets, we performed molecular docking against three representative proteins: 4P5A, 6UNI, and 9AYG. The docking was carried out using AutoDock Vina with an exhaustiveness of 8, generating nine poses for each receptor-ligand pair. The lowest binding energy pose was selected for visualization and interaction analysis.")
        ref = p._element
        
        p = add_yellow_paragraph_after(doc, ref, 
            "Row A presents the docking results for the 4P5A protein target. β-Sitosterol achieved a binding energy of −9.23 kcal/mol for the top-ranked pose, satisfying the strong binding criterion (< −7 kcal/mol). The left panel shows the local interaction view: the compound sits deeply within the hydrophobic pocket and makes close contacts with surrounding residues including HIS-188, VAL-202, ARG-60, SER-59, HIS-61, GLY-62, HIS-87, GLU-66, ARG-88, and ASN-93. These residues cluster in the catalytic cleft region, indicating that β-Sitosterol may exert inhibitory effects through active-site occupation. The right panel provides the global view of the ligand positioned inside the receptor structure.")
        ref = p._element
        
        p = add_yellow_paragraph_after(doc, ref, 
            "Row B shows the docking results for the 6UNI protein target. The best pose yielded a binding energy of −9.12 kcal/mol. In the left panel, the ligand is anchored in a narrow groove formed by several α-helices, surrounded by residues GLU-374, ARG-105, ILE-110, PHE-108, PHE-220, PHE-241, VAL-240, PRO-242, ILE-300, and PHE-304. Many of these residues are aromatic amino acids, which likely contribute to hydrophobic stacking with the sterol ring system of β-Sitosterol. The right panel shows the overall position of the ligand within the protein fold.")
        ref = p._element
        
        p = add_yellow_paragraph_after(doc, ref, 
            "Row C displays the docking results for the 9AYG protein target, where β-Sitosterol showed the strongest binding among the three receptors with a score of −9.68 kcal/mol. The left panel reveals that the ligand is flanked by residues LEU-924, ILE-210, PHE-928, GLU-927, LEU-203, LEU-204, PHE-1001, ALA-997, VAL-1005, GLU-223, and LEU-227. The abundance of leucine and phenylalanine residues around the ligand points to a predominantly hydrophobic binding environment, well suited to accommodate the lipophilic sterol scaffold. The right panel confirms that the ligand is buried in a large cavity at the interface of two structural domains.")
        ref = p._element
        
        # 插入整合图
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_img.add_run()
        r.add_picture('output/figure_composite_docking.png', width=Inches(6.0))
        doc.element.body.append(p_img._element)
        ref.addnext(p_img._element)
        ref = p_img._element
        
        p_cap = add_yellow_paragraph_after(doc, ref, 
            "Figure 6. Integrated molecular docking results of β-Sitosterol with three target proteins. Rows A–C correspond to proteins 4P5A, 6UNI, and 9AYG, respectively. The left column shows local interaction views (green sticks: ligand and surrounding residues; gray cartoon: receptor backbone); the right column shows global docking poses with the ligand embedded in the protein structure. Binding energies for the top-ranked poses are −9.23, −9.12, and −9.68 kcal/mol, respectively.",
            font_size=10)
        ref = p_cap._element
        
        p = add_yellow_paragraph_after(doc, ref, 
            "Taken together, the docking scores for all three protein targets fell well below the −7 kcal/mol threshold, and the top pose for 9AYG even reached −9.68 kcal/mol. These values support tight binding of β-Sitosterol to the selected targets. The interaction maps further show that the compound makes extensive contacts with conserved residues in each binding pocket, mainly through hydrophobic interactions and van der Waals forces. These computational findings lend additional support to the network pharmacology results and provide a structural basis for the observed anti-thyroid cancer activity of XYT.")
    
    doc.save(dst)
    print(f'V3 document saved to: {dst}')

if __name__ == '__main__':
    main()
