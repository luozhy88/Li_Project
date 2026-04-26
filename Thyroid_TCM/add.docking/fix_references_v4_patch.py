#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V4补充补丁：修正遗漏的引用段落
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

PINK = RGBColor(0xFF, 0x69, 0xB4)

def clear_para_text(para):
    for run in para.runs:
        run._element.getparent().remove(run._element)

def set_pink_para(para, text, font_size=11, bold=False):
    clear_para_text(para)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = PINK
    para.paragraph_format.line_spacing = 1.5
    return run

def build_mapping():
    deleted = {8, 9, 13, 15, 16, 17, 27, 28}
    mapping = {}
    new_num = 1
    for old_num in range(1, 45):
        if old_num in deleted:
            mapping[old_num] = None
        else:
            mapping[old_num] = new_num
            new_num += 1
    return mapping

def replace_citations_in_text(text, mapping):
    def replacer(match):
        inner = match.group(1)
        parts = []
        for seg in inner.split(','):
            seg = seg.strip()
            if '-' in seg:
                start, end = seg.split('-')
                start, end = int(start), int(end)
                new_nums = []
                for n in range(start, end + 1):
                    if n in mapping:
                        new_n = mapping[n]
                        if new_n is not None:
                            new_nums.append(new_n)
                if not new_nums:
                    continue
                if len(new_nums) == 1:
                    parts.append(str(new_nums[0]))
                else:
                    consecutive = all(new_nums[i] + 1 == new_nums[i+1] for i in range(len(new_nums)-1))
                    if consecutive:
                        parts.append(f"{new_nums[0]}-{new_nums[-1]}")
                    else:
                        parts.append(','.join(str(x) for x in new_nums))
            else:
                n = int(seg)
                if n in mapping:
                    new_n = mapping[n]
                    if new_n is not None:
                        parts.append(str(new_n))
                else:
                    parts.append(str(n))
        
        if not parts:
            return ""
        return "[" + ", ".join(parts) + "]"
    
    return re.sub(r'\[([\d\s,\-]+)\]', replacer, text)

def main():
    dst = 'output/2024.12.21v2_docking_updated_v4.docx'
    doc = Document(dst)
    mapping = build_mapping()
    
    # Fix missing paragraphs
    missing_paras = [14, 16, 19, 20]
    
    for idx in missing_paras:
        para = doc.paragraphs[idx]
        original_text = para.text.strip()
        new_text = replace_citations_in_text(original_text, mapping)
        if new_text != original_text:
            set_pink_para(doc.paragraphs[idx], new_text)
            print(f"Para {idx}: modified")
            print(f"  Before: {original_text[:120]}")
            print(f"  After:  {new_text[:120]}")
    
    doc.save(dst)
    print(f"\nPatched document saved to: {dst}")

if __name__ == '__main__':
    main()
