#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import RGBColor

doc = Document('output/2024.12.21v2_docking_updated.docx')

YELLOW = RGBColor(0xFF, 0xFF, 0x00)
PINK = RGBColor(0xFF, 0x69, 0xB4)  # Hot pink, more visible

changed = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.color and run.font.color.rgb and run.font.color.rgb == YELLOW:
            run.font.color.rgb = PINK
            changed += 1

# Also check tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.color and run.font.color.rgb and run.font.color.rgb == YELLOW:
                        run.font.color.rgb = PINK
                        changed += 1

doc.save('output/2024.12.21v2_docking_updated.docx')
print(f'Changed {changed} runs from yellow to pink.')
