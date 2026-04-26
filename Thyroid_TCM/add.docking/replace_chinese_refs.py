#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将[3]和[11]中文期刊引用替换为英文综述
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

PINK = RGBColor(0xFF, 0x69, 0xB4)

def clear_para_text(para):
    for run in para.runs:
        run._element.getparent().remove(run._element)

def set_pink_para(para, text, font_size=11, bold=False):
    clear_para_text(para)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = PINK
    para.paragraph_format.line_spacing = 1.5
    return run

def main():
    dst = 'output/2024.12.21v2_docking_updated.docx'
    doc = Document(dst)
    
    # [3] Replacement: ATA 2015 Guidelines
    ref3_new = ("Haugen BR, Alexander EK, Bible KC, Doherty GM, Mandel SJ, Nikiforov YE, et al. "
                "2015 American Thyroid Association Management Guidelines for Adult Patients with Thyroid Nodules and Differentiated Thyroid Cancer: "
                "The American Thyroid Association Guidelines Task Force on Thyroid Nodules and Differentiated Thyroid Cancer. "
                "Thyroid. 2016;26(1):1–133. doi:10.1089/thy.2015.0020")
    
    # [7] (originally [11]) Replacement: Prunella vulgaris review
    ref7_new = ("Wang SJ, Wang XH, Dai YY, Ma MH, Rahman K, Nian H, Zhang H. "
                "Prunella vulgaris: a comprehensive review of chemical constituents, pharmacological effects and clinical applications. "
                "Curr Pharm Des. 2019;25(3):359–369. doi:10.2174/1381612825666190313121608")
    
    # Find and replace in reference list
    in_refs = False
    ref_count = 0
    modified = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "References":
            in_refs = True
            continue
        if in_refs and text:
            if text.startswith("Notes:") or text.startswith("Abbreviations:") or text.startswith("Figure"):
                break
            ref_count += 1
            
            if ref_count == 3:
                set_pink_para(para, ref3_new)
                modified.append((3, text[:80], ref3_new[:80]))
            elif ref_count == 7:
                set_pink_para(para, ref7_new)
                modified.append((7, text[:80], ref7_new[:80]))
    
    doc.save(dst)
    print(f"Document saved to: {dst}")
    print("\nReplacements made:")
    for num, old, new in modified:
        print(f"\n[{num}] Replaced:")
        print(f"  OLD: {old}...")
        print(f"  NEW: {new}...")

if __name__ == '__main__':
    main()
