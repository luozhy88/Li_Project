#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V5最终版：一次性完成所有引用修正
1. 删除幽灵引用[8][9][27][28]
2. 删除不匹配引用[15][16][17]及对应正文
3. 删除重复引用[13]（与[12]合并）
4. 删除伪造引用[4][5]
5. 重新编号所有正文引用和参考文献列表
"""

import shutil
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

PINK = RGBColor(0xFF, 0x69, 0xB4)

def clear_para_text(para):
    for run in para.runs:
        run._element.getparent().remove(run._element)

def set_pink_para(para, text, font_size=11, bold=False):
    clear_para_text(para)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = PINK
    para.paragraph_format.line_spacing = 1.5
    return run

def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)

def build_mapping():
    deleted = {4, 5, 8, 9, 13, 15, 16, 17, 27, 28}
    mapping = {}
    new_num = 1
    for old_num in range(1, 45):
        if old_num in deleted:
            mapping[old_num] = None
        else:
            mapping[old_num] = new_num
            new_num += 1
    return mapping

def replace_citations_in_text(text, mapping):
    def replacer(match):
        inner = match.group(1)
        parts = []
        for seg in inner.split(','):
            seg = seg.strip()
            if '-' in seg:
                start, end = seg.split('-')
                start, end = int(start), int(end)
                new_nums = []
                for n in range(start, end + 1):
                    if n in mapping:
                        new_n = mapping[n]
                        if new_n is not None:
                            new_nums.append(new_n)
                if not new_nums:
                    continue
                if len(new_nums) == 1:
                    parts.append(str(new_nums[0]))
                else:
                    consecutive = all(new_nums[i] + 1 == new_nums[i+1] for i in range(len(new_nums)-1))
                    if consecutive:
                        parts.append(f"{new_nums[0]}-{new_nums[-1]}")
                    else:
                        parts.append(','.join(str(x) for x in new_nums))
            else:
                n = int(seg)
                if n in mapping:
                    new_n = mapping[n]
                    if new_n is not None:
                        parts.append(str(new_n))
                else:
                    parts.append(str(n))
        
        if not parts:
            return ""
        return "[" + ", ".join(parts) + "]"
    
    return re.sub(r'\[([\d\s,\-]+)\]', replacer, text)

def main():
    src = 'output/2024.12.21v2_docking_updated.docx'
    dst = 'output/2024.12.21v2_docking_updated_v5.docx'
    
    if os.path.exists(dst):
        os.remove(dst)
    shutil.copy2(src, dst)
    doc = Document(dst)
    
    mapping = build_mapping()
    print("Citation mapping (old -> new):")
    for old, new in sorted(mapping.items()):
        if new is None:
            print(f"  [{old}] -> DELETED")
        elif old != new:
            print(f"  [{old}] -> [{new}]")
    
    # ===================== Step 1: Modify in-text citations =====================
    print("\nModifying in-text citations...")
    citation_paragraphs = [11, 12, 13, 14, 15, 16, 17, 19, 20, 21, 62, 63, 64, 65, 68]
    
    for idx in citation_paragraphs:
        para = doc.paragraphs[idx]
        original_text = para.text.strip()
        new_text = replace_citations_in_text(original_text, mapping)
        
        # Special handling for paragraph 17: remove sentences with [15][16][17]
        if idx == 17:
            first_sentence = "The Xiao Ying Tang formula, developed by renowned TCM expert Professor Xu Zhuyin, consists of Chai Hu, Yu Jin, Hou Po, Ban Xia, Dang Gui, Huang Qi, Bai Zhu, Cang Zhu, Ban Zhi Lian, Mao Zhua Cao, and Xia Ku Cao."
            new_text = first_sentence
        
        if new_text != original_text:
            set_pink_para(doc.paragraphs[idx], new_text)
            print(f"  Para {idx}: modified")
    
    # ===================== Step 2: Delete ghost/fake reference entries =====================
    print("\nDeleting ghost/fake/duplicate reference entries...")
    
    ref_start = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "References":
            ref_start = i
            break
    
    if ref_start is None:
        print("ERROR: Could not find References section!")
        return
    
    ref_paras = []
    for i in range(ref_start + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text or text.startswith("Notes:") or text.startswith("Abbreviations:") or text.startswith("Figure"):
            break
        ref_paras.append((i, p, text))
    
    delete_positions = [4, 5, 8, 9, 13, 15, 16, 17, 27, 28]
    to_delete = []
    for pos in delete_positions:
        idx = pos - 1
        if 0 <= idx < len(ref_paras):
            to_delete.append(ref_paras[idx][1])
            print(f"  Will delete ref [{pos}]: {ref_paras[idx][2][:80]}...")
    
    for para in to_delete:
        delete_paragraph(para)
    
    doc.save(dst)
    print(f"\nV5 document saved to: {dst}")
    print("Summary of changes:")
    print("  - Deleted in-text citations: [4], [5], [8], [9], [15], [16], [17], [27], [28]")
    print("  - Merged duplicate [12][13] into [12]")
    print("  - Deleted 10 ghost/fake/duplicate reference entries")
    print("  - Re-numbered all in-text citations to match new reference list")
    print("  - Removed 3 unsupported clinical claims from Introduction")

if __name__ == '__main__':
    main()
