#!/usr/bin/env python3
"""在最终版 v3 论文中，将 2.4 节替换为详细的计算方法描述。"""
import json
from docx import Document
from docx.shared import Pt, Cm

INPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_final.docx'
OUTPUT_DOC = INPUT_DOC
JSON_PATH = '/tmp/method_paras.json'


def main():
    doc = Document(INPUT_DOC)
    method_paras = json.load(open(JSON_PATH, encoding='utf-8'))

    # 找到原"数据分析包括："那段
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("数据分析包括："):
            target_idx = i
            break

    if target_idx is None:
        print("未找到目标段落")
        return

    target_p = doc.paragraphs[target_idx]
    insert_after = target_p

    # 依次插入新方法描述段落
    for bold_prefix, body_text in method_paras:
        para = doc.add_paragraph()
        para.add_run(bold_prefix).bold = True
        para.add_run(body_text)
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(10.5)
            run.font.name = "宋体"
        insert_after._element.addnext(para._element)
        insert_after = para

    # 删除原简短段落
    target_p._element.getparent().remove(target_p._element)

    doc.save(OUTPUT_DOC)
    print(f"已更新: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
