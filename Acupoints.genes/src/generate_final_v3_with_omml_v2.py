#!/usr/bin/env python3
"""
生成最终版 v3 论文（修正版），方法学部分的公式使用 Word OMML 数学公式，
并修正文本衔接问题。
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import parse_xml

OUTPUT_DIR = 'output/paper_figures'
BASE_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v2.docx'
OUTPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_final.docx'
PLATFORM_URL = 'https://acupoint-gene-db.luozhy88.workers.dev'
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def om_run(text): return f'<m:r xmlns:m="{MATH_NS}"><m:t>{text}</m:t></m:r>'
def om_sub(base, sub): return f'<m:sSub xmlns:m="{MATH_NS}"><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'
def om_frac(num, den): return f'<m:f xmlns:m="{MATH_NS}"><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'
def om_sum(sub_text, sup_text, expr):
    return (f'<m:nary xmlns:m="{MATH_NS}"><m:naryPr><m:accChr m:val="∑"/><m:limLoc m:val="underOver"/></m:naryPr>'
            f'<m:sub><m:r><m:t>{sub_text}</m:t></m:r></m:sub><m:sup><m:r><m:t>{sup_text}</m:t></m:r></m:sup>'
            f'<m:e>{expr}</m:e></m:nary>')
def om_delim(left, right, expr):
    return (f'<m:d xmlns:m="{MATH_NS}"><m:dPr><m:begChr m:val="{left}"/><m:endChr m:val="{right}"/></m:dPr>'
            f'<m:e>{expr}</m:e></m:d>')
def make_omml(*parts): return f'<m:oMath xmlns:m="{MATH_NS}">{"".join(parts)}</m:oMath>'


def add_omml(para, xml):
    para._element.append(parse_xml(xml))


def rfmt(run, size=10.5, name='宋体'):
    run.font.size = Pt(size)
    run.font.name = name


# ===== OMML 公式定义 =====

# f(x) = Σ_{i=1}^{N} I(x ∈ D_i)
F1 = make_omml(om_run('f(x)='), om_sum('i=1', 'N',
    om_run('I(x∈D') + om_sub(om_run('D'), om_run('i')) + om_run(')')))

# C(a,g) = Σ_{i=1}^{N} I(a ∈ D_i ∧ g ∈ D_i)
F2 = make_omml(om_run('C(a,g)='), om_sum('i=1', 'N',
    om_run('I(a∈D') + om_sub(om_run('D'), om_run('i')) + om_run('∧g∈D') + om_sub(om_run('D'), om_run('i')) + om_run(')')))

# J(a,g) = |D_a ∩ D_g| / |D_a ∪ D_g|
F3 = make_omml(om_run('J(a,g)='), om_frac(
    om_run('|D') + om_sub(om_run('D'), om_run('a')) + om_run('∩D') + om_sub(om_run('D'), om_run('g')) + om_run('|'),
    om_run('|D') + om_sub(om_run('D'), om_run('a')) + om_run('∪D') + om_sub(om_run('D'), om_run('g')) + om_run('|')))

# k(v) = deg(v) / (|V| - 1)
F4 = make_omml(om_run('k(v)='), om_frac(om_run('deg(v)'), om_run('|V|−1')))

# s(v) = Σ_{u∈N(v)} w(u,v)
F5 = make_omml(om_run('s(v)='), om_sum('u∈N(v)', '', om_run('w(u,v)')))

# B(v) = Σ_{s≠v≠t} σ_st(v) / σ_st
F6 = make_omml(om_run('B(v)='), om_sum('s≠v≠t', '',
    om_frac(om_sub(om_run('σ'), om_run('st')) + om_run('(v)'), om_sub(om_run('σ'), om_run('st')))))

# A·x = λ·x
F7 = make_omml(om_run('A·x=λ·x'))

# Q = 1/(2m) Σ_{ij} [A_{ij} - (k_i·k_j)/(2m)] δ(c_i, c_j)
F8 = make_omml(om_run('Q='), om_frac(om_run('1'), om_run('2m')),
    om_sum('ij', '',
        om_delim('[', ']', om_sub(om_run('A'), om_run('ij')) + om_run('−') +
            om_frac(om_sub(om_run('k'), om_run('i')) + om_run('·') + om_sub(om_run('k'), om_run('j')), om_run('2m')))
        + om_run(' δ') + om_delim('(', ')',
            om_sub(om_run('c'), om_run('i')) + om_run(',') + om_sub(om_run('c'), om_run('j')))))

# m = Σ_{ij} A_{ij} / 2
F9 = make_omml(om_run('m='), om_frac(om_sum('ij', '', om_sub(om_run('A'), om_run('ij'))), om_run('2')))

# P(X=k) = C(M,k)·C(N-M,n-k) / C(N,n)
F10 = make_omml(om_run('P(X=k)='), om_frac(om_run('C(M,k)·C(N−M,n−k)'), om_run('C(N,n)')))

# P(X≥k) = Σ_{x=k}^{min(n,M)} C(M,x)·C(N-M,n-x) / C(N,n)
F11 = make_omml(om_run('P(X≥k)='), om_sum('x=k', 'min(n,M)',
    om_frac(om_run('C(M,x)·C(N−M,n−x)'), om_run('C(N,n)'))))


FIGURE_PLAN = [
    ('3.1　文献收录概况',     f'{OUTPUT_DIR}/fig1_dashboard.png',      '图1  数据库统计概览与文献年份分布', 15),
    ('3.2　穴位分布特征',     f'{OUTPUT_DIR}/fig3_top_acupoints.png',  '图2  Top 10 高频研究穴位', 12),
    ('3.3　高频关联基因与基因家族', f'{OUTPUT_DIR}/fig2_top_genes.png',      '图3  Top 10 高频关联基因', 12),
    ('3.5　穴位-基因共现分析', f'{OUTPUT_DIR}/fig4_network.png',        '图4  穴位-基因共现网络图（共现频次≥3）', 15),
    ('3.6.1　网络拓扑分析',    f'{OUTPUT_DIR}/fig7_topology.png',       '图5  网络拓扑分析：基因度中心性与中介中心性 Top 10', 15),
    ('3.6.2　聚类分析（Louvain社区发现）', f'{OUTPUT_DIR}/fig8_clusters.png', '图6  Louvain 聚类分析：各社区节点构成', 12),
    ('3.6.4　超几何分布显著性检验', f'{OUTPUT_DIR}/fig9_hypergeo.png', '图7  超几何检验：Top 10 显著穴位-基因对', 12),
    ('4　讨论',              f'{OUTPUT_DIR}/fig5_acupoint_query.png', '图8  数据库平台"穴位→基因"查询模块示例（以足三里为例）', 14),
]
PLATFORM_FIG2 = (f'{OUTPUT_DIR}/fig6_gene_query.png', '图9  数据库平台"基因→穴位"查询模块示例（以TNF为例）', 12)


def find_para(doc, substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    return None


def insert_pic(doc, target, path, width, caption):
    if not os.path.exists(path):
        return False
    pp = doc.add_paragraph()
    pp.alignment = 1
    pp.add_run().add_picture(path, width=Cm(width))
    cp = doc.add_paragraph(caption)
    cp.alignment = 1
    cp.paragraph_format.space_after = Pt(12)
    for r in cp.runs:
        rfmt(r)
    target._element.addnext(pp._element)
    pp._element.addnext(cp._element)
    return True


def add_method_para(doc, insert_after, bold_text, parts):
    """
    parts: list of (type, content)
    type: 'text' | 'omml'
    """
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    rfmt(r)
    for typ, content in parts:
        if typ == 'text':
            r = p.add_run(content)
            rfmt(r)
        elif typ == 'omml':
            add_omml(p, content)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p._element)
    return p


def main():
    doc = Document(BASE_DOC)

    # --- 摘要添加平台网址 ---
    for p in doc.paragraphs:
        if p.text.strip().startswith('目的') and '为揭示针灸作用的分子机制' in p.text:
            p.add_run(f'同时，基于该数据库构建了在线检索与分析平台（{PLATFORM_URL}），为研究者提供便捷的文献数据查询服务。')
            break

    # --- 方法学添加平台说明 ---
    for p in doc.paragraphs:
        if 'SQLite' in p.text and 'articles（文献表）' in p.text:
            plat = doc.add_paragraph(
                f'此外，基于Cloudflare Worker构建了轻量级Web数据检索平台（{PLATFORM_URL}），'
                '提供数据概览、穴位查基因、基因查穴位、共现网络可视化及文献浏览等功能，'
                '便于研究者在线查询与浏览数据库内容。平台采用服务器端渲染（SSR）策略，'
                '所有数据以JSON格式内联部署，无需后端数据库连接，确保高可用性与快速响应。')
            plat.paragraph_format.first_line_indent = Cm(0.74)
            for r in plat.runs:
                rfmt(r)
            p._element.addnext(plat._element)
            break

    # --- 替换数据分析方法段落 ---
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('数据分析包括：'):
            target_idx = i
            break
    if target_idx is None:
        print("未找到目标段落")
        return
    target_p = doc.paragraphs[target_idx]
    after = target_p

    # （1）频次统计
    after = add_method_para(doc, after, '（1）频次统计：', [
        ('text', '运用 Python 3.13 的 pandas（v3.0.1）与 collections 模块，对穴位名称、基因符号及疾病实体的出现频次进行计数统计。'),
        ('text', '频次计算公式为：'),
        ('omml', F1),
        ('text', '，其中 N 为纳入文献总数（N=344），D_i 为第 i 篇文献提取的实体集合，I(·) 为指示函数，当实体 x 出现在文献 D_i 中时取值为 1，否则为 0。对于基因家族（如 AKT、PI3K 等），在频次统计中单独标注为「基因家族/通路」，不纳入单一基因的频次排序，以避免重复计数。'),
    ])

    # （2）共现分析与网络构建
    after = add_method_para(doc, after, '（2）共现分析与网络构建：', [
        ('text', '采用共现频率（co-occurrence frequency）衡量穴位-基因关联强度。对于穴位 a 与基因 g，其共现频率定义为：'),
        ('omml', F2),
        ('text', '，即同时包含穴位 a 和基因 g 的文献篇数。进一步计算 Jaccard 相似系数以消除高频实体的偏倚：'),
        ('omml', F3),
        ('text', '，其中 D_a、D_g 分别为包含穴位 a 和基因 g 的文献集合。共现网络以 NetworkX（v3.6.1）构建为无向二分图 G=(V,E)，节点集合 V=A∪G 分为穴位层 A 与基因层 G，边权 w(a,g)=C(a,g)（共现频次≥3 的边纳入可视化）。网络可视化采用 Kamada-Kawai 力导向布局，节点大小映射加权度中心性，边透明度映射共现频次。'),
    ])

    # （3）网络拓扑分析
    after = add_method_para(doc, after, '（3）网络拓扑分析：', [
        ('text', '基于 NetworkX（v3.6.1）计算以下拓扑指标：①度中心性（degree centrality）：'),
        ('omml', F4),
        ('text', '，其中 deg(v) 为节点 v 的邻居数，反映节点的局部连接密度；②加权度中心性（weighted degree）：'),
        ('omml', F5),
        ('text', '，即关联边权之和，反映节点在共现频次层面的影响力；③中介中心性（betweenness centrality）：'),
        ('omml', F6),
        ('text', '，其中 σ_st 为节点 s 到 t 的最短路径总数，σ_st(v) 为经过 v 的最短路径数，反映节点作为信息流通「桥梁」的能力；④特征向量中心性（eigenvector centrality）：'),
        ('omml', F7),
        ('text', '，求解邻接矩阵 A 的主特征向量，节点得分与其邻居的中心性加权相关。上述计算均调用 NetworkX 内置函数（degree_centrality、betweenness_centrality、eigenvector_centrality），其中 betweenness 计算采用 Brandes 算法（时间复杂度 O(|V||E|)）。'),
    ])

    # （4）Louvain聚类
    after = add_method_para(doc, after, '（4）聚类分析（Louvain 社区发现）：', [
        ('text', '采用 Louvain 社区发现算法（python-louvain 0.16）对共现网络进行模块划分。该算法以模块度（modularity）最大化为目标：'),
        ('omml', F8),
        ('text', '，其中 '),
        ('omml', F9),
        ('text', ' 为网络总边权，A_ij 为节点 i,j 的邻接矩阵元素，k_i、k_j 为节点度，δ(·) 为克罗内克函数（当 i,j 属于同一社区时取 1，否则为 0）。算法通过两阶段迭代实现：阶段一将每个节点视为独立社区，依次将节点移动到能使模块度增量 ΔQ 最大的邻居社区；阶段二将同一社区内节点收缩为超级节点，重构网络后重复阶段一，直至模块度不再提升。本研究设置分辨率参数 γ=1.0（默认），最终得到 8 个社区。'),
    ])

    # （5）功能富集分析
    after = add_method_para(doc, after, '（5）功能富集分析：', [
        ('text', '选取网络中度中心性排名前 50 的基因作为核心基因集，通过 Enrichr 在线数据库（https://maayanlab.cloud/Enrichr/）进行 GO（Gene Ontology）生物学过程（Biological Process）与 KEGG 通路富集分析。富集显著性采用 Fisher 精确检验（右侧检验），概率质量函数为：'),
        ('omml', F10),
        ('text', '，其中 N 为背景基因总数（本数据库有效基因 1997 个），M 为某一 GO/KEGG 条目注释的基因数，n 为核心基因集大小（n=50），k 为核心基因集中属于该条目的基因数。累积 p 值计算为：'),
        ('omml', F11),
        ('text', '。当 p<0.05 时，认为该通路富集显著。多重检验校正采用 Benjamini-Hochberg 法计算校正后 p 值（FDR），以 FDR<0.05 作为显著性阈值。'),
    ])

    # （6）超几何分布检验
    after = add_method_para(doc, after, '（6）显著性检验（超几何分布）：', [
        ('text', '为评估穴位-基因共现是否显著高于随机期望，采用超几何分布检验（scipy.stats.hypergeom，v1.17.1）。设总体文献数 N=344，其中包含特定基因 g 的文献数为 M（基因 g 的边际频次），从总体中随机抽取 n 篇文献（n 为包含特定穴位 a 的文献数，即穴位 a 的边际频次），观察到的共现文献数为 k=C(a,g)。超几何分布概率质量函数为：'),
        ('omml', F10),
        ('text', '。单侧 p 值计算为：'),
        ('omml', F11),
        ('text', '。当 p<0.05 时，认为该穴位-基因对的共现显著高于随机期望。需要强调的是，显著性检验评估的是「研究关注度」的统计学意义，而非直接的生物学因果关系；共现显著仅表明该穴位-基因对在文献中被共同提及的频率高于随机水平，不能推断穴位对基因表达具有直接调控作用。'),
    ])

    # （7）软件环境
    after = add_method_para(doc, after, '（7）软件环境：', [
        ('text', '本研究数据分析均在 Python 3.13 环境下完成，主要依赖包及版本包括：pandas 3.0.1（数据处理）、NetworkX 3.6.1（网络构建与拓扑分析）、python-louvain 0.16（社区发现）、scipy 1.17.1（统计检验）、matplotlib 3.10.8（可视化）、numpy 2.4.3（数值计算）。富集分析通过 Enrichr API 在线提交基因列表获取结果。在解读所有分析结果时，须严格区分「共现关注度」与「实验验证的调控关系」，本数据库的定位是文献数据整合与假设生成工具，而非因果推断平台。'),
    ])

    # 删除原简短段落
    target_p._element.getparent().remove(target_p._element)

    # --- 插入图片 ---
    inserted = []
    for section_key, pic_path, caption, width in FIGURE_PLAN:
        section_para = find_para(doc, section_key)
        if section_para is None:
            print(f'警告：未找到小节 "{section_key}"')
            continue
        found = False
        target = None
        for p in doc.paragraphs:
            if found and p.text.strip():
                target = p
                break
            if p._element is section_para._element:
                found = True
        if target and insert_pic(doc, target, pic_path, width, caption):
            inserted.append(caption)
            if section_key == '4　讨论':
                p2, c2, w2 = PLATFORM_FIG2
                # 找到刚插入的图8图注
                fig8_cap = None
                for p in doc.paragraphs:
                    if '图8' in p.text:
                        fig8_cap = p
                if fig8_cap and insert_pic(doc, fig8_cap, p2, w2, c2):
                    inserted.append(c2)

    # --- 讨论部分添加平台价值 ---
    for p in doc.paragraphs:
        if p.text.strip().startswith('本研究构建了针灸穴位-基因关联数据库'):
            p.add_run(f'数据库平台（{PLATFORM_URL}）提供了在线检索、共现网络可视化及文献溯源功能，'
                      '可供研究者快速获取特定穴位或基因的关联信息，为后续实验设计提供文献依据。')
            break

    doc.save(OUTPUT_DOC)
    print(f'最终文档已保存: {OUTPUT_DOC}')
    print(f'共插入 {len(inserted)} 张图片')


if __name__ == '__main__':
    main()
