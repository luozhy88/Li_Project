#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成平台分析图表并插入Word文档
"""

import os
import json
import subprocess
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import font_manager as fm
import networkx as nx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = 'output/paper_figures'
os.makedirs(OUTPUT_DIR, exist_ok=True)

BASE_URL = 'http://ag.db.qyjc.top'

# 加载中文字体
CJK_FONT = fm.FontProperties(fname='/usr/share/fonts/google-noto-cjk/NotoSansCJK-Medium.ttc')
CJK_FONT_LIGHT = fm.FontProperties(fname='/usr/share/fonts/google-noto-cjk/NotoSansCJK-Light.ttc')
CJK_FONT_BOLD = fm.FontProperties(fname='/usr/share/fonts/google-noto-cjk/NotoSansCJK-Bold.ttc')


def fetch_api(endpoint):
    url = f'{BASE_URL}{endpoint}'
    try:
        result = subprocess.run(
            ['curl', '-s', '-A', 'Mozilla/5.0', '--max-time', '15', url],
            capture_output=True, text=True, timeout=20
        )
        return json.loads(result.stdout)
    except Exception as e:
        print(f'API fetch failed: {endpoint} - {e}')
        return None


def fig1_dashboard():
    stats = fetch_api('/api/stats')
    year_data = fetch_api('/api/year_data')
    if not stats or not year_data:
        return None

    fig = plt.figure(figsize=(14, 6))
    gs = fig.add_gridspec(1, 2, width_ratios=[1, 2])

    ax1 = fig.add_subplot(gs[0])
    ax1.axis('off')
    labels = ['收录文献', '识别基因', '涉及穴位', '相关疾病', '全文获取']
    values = [stats['articles'], stats['genes'], stats['acupoints'], stats['diseases'], stats['fulltext']]
    colors = ['#1a5fb4', '#26a269', '#c061cb', '#f59e0b', '#ef4444']
    
    for i, (l, v, c) in enumerate(zip(labels, values, colors)):
        rect = plt.Rectangle((0.1, 0.78 - i*0.18), 0.8, 0.14, facecolor=c, alpha=0.15, edgecolor=c, linewidth=2, transform=ax1.transAxes)
        ax1.add_patch(rect)
        ax1.text(0.5, 0.85 - i*0.18, str(v), fontsize=20, fontweight='bold', color=c, ha='center', transform=ax1.transAxes, fontproperties=CJK_FONT)
        ax1.text(0.5, 0.80 - i*0.18, l, fontsize=12, color='#555', ha='center', transform=ax1.transAxes, fontproperties=CJK_FONT_LIGHT)
    ax1.set_title('数据库统计概览', fontsize=14, fontweight='bold', pad=20, fontproperties=CJK_FONT)

    ax2 = fig.add_subplot(gs[1])
    years = [d['year'] for d in year_data]
    counts = [d['count'] for d in year_data]
    ax2.fill_between(years, counts, alpha=0.3, color='#1a5fb4')
    ax2.plot(years, counts, marker='o', color='#1a5fb4', linewidth=2, markersize=6)
    ax2.set_title('文献年份分布趋势', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax2.set_xlabel('年份', fontsize=12, fontproperties=CJK_FONT_LIGHT)
    ax2.set_ylabel('收录文献数', fontsize=12, fontproperties=CJK_FONT_LIGHT)
    ax2.grid(axis='y', alpha=0.3)
    for y, c in zip(years, counts):
        if c > 5:
            ax2.annotate(str(c), (y, c), textcoords="offset points", xytext=(0, 8), ha='center', fontsize=9)

    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig1_dashboard.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def fig2_top_genes():
    data = fetch_api('/api/top_genes?limit=10')
    if not data:
        return None

    genes = [d['gene'] for d in data][::-1]
    counts = [d['count'] for d in data][::-1]

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(genes, counts, color='#3584e4', edgecolor='white', height=0.6)
    for bar, c in zip(bars, counts):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2, str(c), va='center', fontsize=10, color='#333')
    ax.set_xlabel('共现文献数', fontsize=12, fontproperties=CJK_FONT_LIGHT)
    ax.set_title('Top 10 高频关联基因', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax.set_xlim(0, max(counts) * 1.15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig2_top_genes.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def fig3_top_acupoints():
    data = fetch_api('/api/top_acupoints?limit=10')
    if not data:
        return None

    names = [f"{d['name_en']}\n{d['code']}" if d.get('code') else d['name_en'] for d in data][::-1]
    counts = [d['count'] for d in data][::-1]

    fig, ax = plt.subplots(figsize=(10, 6))
    colors = plt.cm.Purples([0.4 + 0.06 * i for i in range(len(names))])[::-1]
    bars = ax.barh(names, counts, color=colors, edgecolor='white', height=0.6)
    for bar, c in zip(bars, counts):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, str(c), va='center', fontsize=10, color='#333')
    ax.set_xlabel('共现文献数', fontsize=12, fontproperties=CJK_FONT_LIGHT)
    ax.set_title('Top 10 高频研究穴位', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax.set_xlim(0, max(counts) * 1.15)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig3_top_acupoints.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def fig4_network():
    data = fetch_api('/api/network?min_cooccur=5&limit=80')
    if not data:
        return None

    G = nx.Graph()
    acupoint_nodes = set()
    gene_nodes = set()
    
    for d in data:
        apt = d['acupoint']
        gene = d['gene']
        count = d['count']
        G.add_edge(apt, gene, weight=count)
        acupoint_nodes.add(apt)
        gene_nodes.add(gene)

    fig, ax = plt.subplots(figsize=(14, 10))
    pos = nx.spring_layout(G, k=2.5, iterations=50, seed=42)
    
    edges = G.edges()
    weights = [G[u][v]['weight'] for u, v in edges]
    nx.draw_networkx_edges(G, pos, width=[w * 0.4 for w in weights], alpha=0.5, edge_color='#999', ax=ax)
    
    apt_pos = {n: pos[n] for n in acupoint_nodes if n in pos}
    nx.draw_networkx_nodes(G, apt_pos, nodelist=list(acupoint_nodes & set(pos.keys())), 
                           node_color='#c061cb', node_size=[G.degree(n) * 80 + 200 for n in acupoint_nodes if n in pos], 
                           alpha=0.9, ax=ax)
    
    gene_pos = {n: pos[n] for n in gene_nodes if n in pos}
    nx.draw_networkx_nodes(G, gene_pos, nodelist=list(gene_nodes & set(pos.keys())), 
                           node_color='#3584e4', node_size=[G.degree(n) * 60 + 150 for n in gene_nodes if n in pos], 
                           alpha=0.85, ax=ax)
    
    nx.draw_networkx_labels(G, pos, font_size=9, font_family='sans-serif', ax=ax)
    
    ax.set_title('穴位-基因共现网络图（共现频次≥5）', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax.axis('off')
    
    from matplotlib.lines import Line2D
    legend_elements = [
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#c061cb', markersize=12, label='穴位'),
        Line2D([0], [0], marker='o', color='w', markerfacecolor='#3584e4', markersize=12, label='基因'),
    ]
    ax.legend(handles=legend_elements, loc='upper right', fontsize=11, prop=CJK_FONT_LIGHT)
    
    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig4_network.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def fig5_acupoint_query():
    data = fetch_api('/api/acupoint_genes?acupoint=Zusanli')
    if not data:
        return None
    data = data[:15]

    genes = [d['gene'] for d in data][::-1]
    cooccurs = [d['cooccur'] for d in data][::-1]
    jaccards = [d['jaccard'] for d in data][::-1]

    fig, ax1 = plt.subplots(figsize=(12, 7))
    bars = ax1.barh(genes, cooccurs, color='#1a5fb4', alpha=0.8, label='共现频次')
    for bar, c in zip(bars, cooccurs):
        ax1.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2, str(c), va='center', fontsize=9)
    ax1.set_xlabel('共现频次', fontsize=12, color='#1a5fb4', fontproperties=CJK_FONT_LIGHT)
    ax1.tick_params(axis='x', labelcolor='#1a5fb4')
    
    ax2 = ax1.twiny()
    ax2.plot(jaccards, genes, 'ro-', markersize=5, label='Jaccard系数', alpha=0.7)
    ax2.set_xlabel('Jaccard系数', fontsize=12, color='#d32f2f', fontproperties=CJK_FONT_LIGHT)
    ax2.tick_params(axis='x', labelcolor='#d32f2f')
    ax2.set_xlim(0, max(jaccards) * 1.3)

    ax1.set_title('足三里（Zusanli）关联基因分析', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    fig.legend(loc='lower right', bbox_to_anchor=(0.95, 0.05), fontsize=10, prop=CJK_FONT_LIGHT)
    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig5_acupoint_query.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def fig6_gene_query():
    data = fetch_api('/api/gene_acupoints?gene=TNF')
    if not data:
        return None
    data = [d for d in data if d['cooccur'] >= 2][:15]

    apts = [f"{d['acupoint']}" for d in data][::-1]
    cooccurs = [d['cooccur'] for d in data][::-1]

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(apts, cooccurs, color='#26a269', edgecolor='white', height=0.6)
    for bar, c in zip(bars, cooccurs):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2, str(c), va='center', fontsize=10)
    ax.set_xlabel('共现频次', fontsize=12, fontproperties=CJK_FONT_LIGHT)
    ax.set_title('TNF 关联穴位分析', fontsize=14, fontweight='bold', fontproperties=CJK_FONT)
    ax.set_xlim(0, max(cooccurs) * 1.2)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = f'{OUTPUT_DIR}/fig6_gene_query.png'
    plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def find_para_by_text(doc, text_substr):
    """返回第一个包含指定文本的段落对象"""
    for p in doc.paragraphs:
        if text_substr in p.text:
            return p
    return None


def insert_picture_after(doc, target_para, picture_path, width_cm, caption_text):
    """在指定段落后插入图片和图注"""
    if target_para is None:
        return
    
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(picture_path, width=Cm(width_cm))
    
    cap_para = doc.add_paragraph(caption_text)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(12)
    for run in cap_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
    
    target_para._element.addnext(pic_para._element)
    pic_para._element.addnext(cap_para._element)


def insert_figures_to_docx(docx_path, output_path):
    doc = Document(docx_path)
    
    # 找到所有目标段落（从后往前插入，避免前面插入影响后面的定位）
    p_discuss = find_para_by_text(doc, '4　讨论')
    p35 = find_para_by_text(doc, '3.5　穴位-基因关联挖掘')
    p33 = find_para_by_text(doc, '3.3　高频关联基因')
    p32 = find_para_by_text(doc, '3.2　穴位分布特征')
    p31 = find_para_by_text(doc, '3.1　文献收录概况')
    
    print('找到段落:', {
        '3.1': p31.text[:20] if p31 else None,
        '3.2': p32.text[:20] if p32 else None,
        '3.3': p33.text[:20] if p33 else None,
        '3.5': p35.text[:20] if p35 else None,
        '4': p_discuss.text[:20] if p_discuss else None,
    })
    
    # 辅助函数：在某一小节标题后的第一个内容段落后插入
    def insert_after_section(section_para, pic_path, width_cm, caption):
        if section_para is None:
            return
        # 找到该标题后面的第一个有文字的段落
        found = False
        for p in doc.paragraphs:
            if found and p.text.strip():
                insert_picture_after(doc, p, pic_path, width_cm, caption)
                return
            if p._element is section_para._element:
                found = True
    
    # 按从后往前的顺序插入，避免索引错乱
    # 1. 先插入 3.6 平台界面展示（在"4 讨论"前）
    if p_discuss is not None:
        heading = doc.add_paragraph('3.6  数据库平台界面展示')
        heading.style = doc.styles['Heading 2']
        p_discuss._element.addprevious(heading._element)
        
        desc = doc.add_paragraph(
            '为便于科研人员检索与浏览针灸穴位-基因关联数据，本研究基于Cloudflare Worker构建了轻量级Web分析平台（http://ag.db.qyjc.top）。'
            '平台提供数据概览、穴位查基因、基因查穴位、共现网络可视化及文献浏览五大核心模块。'
            '图5、图6分别展示了"穴位→基因"与"基因→穴位"双向查询模块的典型分析结果。'
        )
        desc.paragraph_format.first_line_indent = Cm(0.74)
        heading._element.addnext(desc._element)
        
        pic5 = doc.add_paragraph()
        pic5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic5.add_run().add_picture(f'{OUTPUT_DIR}/fig5_acupoint_query.png', width=Cm(14))
        cap5 = doc.add_paragraph('图5  平台"穴位→基因"查询模块示例（以足三里为例）')
        cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap5.runs:
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
        desc._element.addnext(cap5._element)
        cap5._element.addprevious(pic5._element)
        
        pic6 = doc.add_paragraph()
        pic6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic6.add_run().add_picture(f'{OUTPUT_DIR}/fig6_gene_query.png', width=Cm(12))
        cap6 = doc.add_paragraph('图6  平台"基因→穴位"查询模块示例（以TNF为例）')
        cap6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap6.runs:
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
        cap5._element.addnext(cap6._element)
        cap6._element.addprevious(pic6._element)
    
    # 2. 插入图4到3.5后面
    insert_after_section(p35, f'{OUTPUT_DIR}/fig4_network.png', 15, '图4  穴位-基因共现网络图（共现频次≥5）')
    
    # 3. 插入图2到3.3后面
    insert_after_section(p33, f'{OUTPUT_DIR}/fig2_top_genes.png', 12, '图2  Top 10 高频关联基因分布')
    
    # 4. 插入图3到3.2后面
    insert_after_section(p32, f'{OUTPUT_DIR}/fig3_top_acupoints.png', 12, '图3  Top 10 高频研究穴位分布')
    
    # 5. 插入图1到3.1后面
    insert_after_section(p31, f'{OUTPUT_DIR}/fig1_dashboard.png', 15, '图1  数据库平台数据概览与文献年份分布')
    
    doc.save(output_path)
    print(f'文档已保存: {output_path}')


def main():
    print('正在生成图1：数据库平台数据概览...')
    fig1_dashboard()
    print('正在生成图2：Top 10 高频关联基因...')
    fig2_top_genes()
    print('正在生成图3：Top 10 高频研究穴位...')
    fig3_top_acupoints()
    print('正在生成图4：穴位-基因共现网络...')
    fig4_network()
    print('正在生成图5：足三里关联基因分析...')
    fig5_acupoint_query()
    print('正在生成图6：TNF关联穴位分析...')
    fig6_gene_query()
    
    print('正在插入图片到Word文档...')
    insert_figures_to_docx(
        'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿.docx',
        'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_含图.docx'
    )
    print('全部完成！')


if __name__ == '__main__':
    main()
