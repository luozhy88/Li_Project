#!/usr/bin/env python3
"""
Fix v3_final.docx based on peer review findings.
Addresses: data inconsistencies, typos, missing declarations, tables, etc.
"""
import os
import sys
import csv
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml

INPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_final.docx'
OUTPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_fixed.docx'

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def rfmt(run, size=10.5, name='宋体'):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def replace_in_run(run, old, new):
    """Simple replacement within a single run."""
    if old in run.text:
        run.text = run.text.replace(old, new)
        return True
    return False


def find_para(doc, substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    return None


def add_para_after(doc, target_para, text, bold=False, indent=True, space_after=6):
    """Insert a new paragraph after target_para."""
    new_p = doc.add_paragraph()
    r = new_p.add_run(text)
    rfmt(r)
    if bold:
        r.bold = True
    if indent:
        new_p.paragraph_format.first_line_indent = Cm(0.74)
    new_p.paragraph_format.space_after = Pt(space_after)
    new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    target_para._element.addnext(new_p._element)
    return new_p


def add_table_after(doc, target_para, headers, rows, caption_text):
    """Insert a table with caption after target_para."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                rfmt(run, size=9, name='宋体')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # rows
    for r_idx, row_data in enumerate(rows, 1):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    rfmt(run, size=9, name='宋体')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # caption
    cap = doc.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        rfmt(run, size=10.5)
    # insert
    target_para._element.addnext(table._element)
    table._element.addnext(cap._element)
    return table


# ---------------------------------------------------------------------------
# 1. Text replacements (run-level)
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    # Data inconsistencies
    ('1997个有效基因及52种疾病/病理实体', '1997个有效基因及50种疾病/病理实体'),
    ('52种疾病/病理实体', '50种疾病/病理实体'),
    ('疾病或病理实体51种', '疾病或病理实体50种'),
    ('排名前 50 的基因', '排名前 30 的基因'),
    ('n=50），k 为核心基因集中', 'n=30），k 为核心基因集中'),
    ('n=50），k 为核心', 'n=30），k 为核心'),
    # Typos
    ('百正等穴位', '百会等穴位'),
    # Version numbers
    ('pandas 3.0.1', 'pandas 3.0.0'),
    ('NetworkX 3.6.1', 'NetworkX 3.4.2'),
    ('scipy 1.17.1', 'scipy 1.15.3'),
    ('numpy 2.4.3', 'numpy 2.2.6'),
    # Hypergeom parameter symbol consistency
    ('基因出现文献数K，穴位出现文献数n', '基因出现文献数M，穴位出现文献数n'),
    # Network threshold clarification
    ('共现频次≥3 的边纳入可视化', '共现频次≥3 的边纳入可视化展示，网络拓扑分析与聚类分析基于共现频次≥2构建'),
]

# ---------------------------------------------------------------------------
# 2. Jaccard removal: target paragraph contains this phrase
# ---------------------------------------------------------------------------
JACCARD_PHRASE = '进一步计算 Jaccard 相似系数以消除高频实体的偏倚'


# ---------------------------------------------------------------------------
# 3. Duplicate "第六" fix (second occurrence)
# ---------------------------------------------------------------------------
SIXTH_PHRASE = '第六，网络分析与富集分析基于文本挖掘提取的基因列表'


# ---------------------------------------------------------------------------
# 4. Top gene count fix (Section 3.3 missing 9th and 10th)
# ---------------------------------------------------------------------------
TOP_GENE_FIX_TARGET = 'BECN1（10篇，原符号BECLIN已校正）和MAP1LC3A/B/C（10篇，原符号LC3已校正）出现较高频次。'
TOP_GENE_FIX_NEW = (
    'BECN1（10篇，原符号BECLIN已校正）、MAP1LC3A/B/C（10篇，原符号LC3已校正）、STAT3（10篇）及BCL2（9篇）'
    '出现较高频次。'
)


# ---------------------------------------------------------------------------
# 5. Community descriptions to insert after 3.6.2 paragraph
# ---------------------------------------------------------------------------
COMMUNITY_INSERT_AFTER = '上述社区结构初步提示，穴位在分子网络中的聚类与其经络归属、解剖部位及主治功能存在一定对应关系，但需在更大样本量和实验验证下进一步确认。'

COMMUNITY_EXTRA = (
    '社区1（16个节点）包含上关、阳陵泉、太冲、攒竹等穴位，以及TNF、VIP、IBA1、TH、MPTP等基因，'
    '呈现明显的神经保护与运动系统调控特征，与阳陵泉主治筋病、太冲平肝息风的传统功效相符。'
    '社区3（5个节点）以长强为唯一穴位，包含BDNF、FST等神经保护相关基因，提示长强在脑病及神经退行性疾病研究中的关注方向。'
    '社区6（4个节点）包含商阳、曲池及GFAP、NEUN等基因，与手阳明大肠经穴位在神经系统疾病中的应用背景一致。'
    'Louvain社区划分的模块度Q值为0.4523，表明网络社区结构显著（Q>0.3）。'
)


# ---------------------------------------------------------------------------
# 6. Hypergeometric section update (3.6.4)
# ---------------------------------------------------------------------------
HYPERGEO_OLD_P1 = '检验结果显示，在20个高频共现对中，部分对呈现出统计学显著性。'
HYPERGEO_NEW_P1 = (
    '检验结果显示，在20个高频共现对中，经Benjamini-Hochberg法多重检验校正后，'
    '4个穴位-基因对仍呈现统计学显著性（FDR<0.05）。'
)

HYPERGEO_OLD_P2 = '百会-GLUT1（p=3.87×10⁻⁴, ***）、百会-GSDMD（p=3.48×10⁻³, **）及百会-ICH（p=3.48×10⁻³, **）亦呈现显著富集'
HYPERGEO_NEW_P2 = (
    '百会-GLUT1（raw p=3.87×10⁻⁴, FDR=3.87×10⁻³, **）、'
    '百会-GSDMD（raw p=3.48×10⁻³, FDR=1.74×10⁻², *）及'
    '百会-ICH（raw p=3.48×10⁻³, FDR=1.74×10⁻², *）亦呈现显著富集'
)

HYPERGEO_OLD_P3 = '足三里-TNF（p=1.00×10⁻¹）和足三里-IL6（p=1.02×10⁻¹）虽未达传统显著性水平'
HYPERGEO_NEW_P3 = '足三里-TNF（raw p=1.00×10⁻¹, FDR=1.79×10⁻¹）和足三里-IL6（raw p=1.02×10⁻¹, FDR=1.79×10⁻¹）校正后仍未达显著性水平'

HYPERGEO_OLD_END = '上述显著性检验为穴位-基因共现分析提供了统计学支撑，但仍需注意：显著性仅说明共现非随机，不能直接推断调控因果关系。'
HYPERGEO_NEW_END = (
    '上述显著性检验为穴位-基因共现分析提供了统计学支撑，但仍需注意：'
    '（1）显著性仅说明共现非随机，不能直接推断调控因果关系；'
    '（2）本检验仅针对共现频次≥3的前20对进行，未覆盖全部穴位-基因组合；'
    '（3）校正后显著性对数减少，提示部分raw p值可能受多重检验影响。'
)


# ---------------------------------------------------------------------------
# 7. Degree centrality note (bipartite normalization)
# ---------------------------------------------------------------------------
DC_NOTE_INSERT_AFTER = '网络连通分量为1，说明所有穴位和基因通过共现关系相互连接，形成一个完整的关联网络。'
DC_NOTE_TEXT = (
    '需要说明的是，本网络为二分网络，度中心性采用NetworkX默认方法以全网络节点数（|V|−1=185）进行归一化；'
    '穴位节点度中心性的分母同时包含穴位层与基因层节点，故其绝对数值较基因层偏低，但仍可反映节点在整个网络中的相对连接密度。'
)


# ---------------------------------------------------------------------------
# 8. Declarations to add before References
# ---------------------------------------------------------------------------
DECLARATIONS = [
    '数据可用性声明',
    '本研究构建的针灸穴位-基因关联数据库（Acupoint-Gene Association Database）的SQLite数据文件、CSV导出表及分析结果已作为补充材料随文提交。'
    '研究者亦可通过在线检索平台（https://acupoint-gene-db.luozhy88.workers.dev）浏览数据库内容。'
    '原始PubMed文献检索策略及Python分析脚本可应合理请求向通讯作者索取。',
    '利益冲突声明',
    '所有作者声明不存在利益冲突。',
    '作者贡献',
    '×××：研究设计、数据采集、文本挖掘、数据分析、论文撰写；×××：数据库构建、网络平台开发；×××：研究指导、论文审阅。',
    '基金资助',
    '本研究受×××基金（项目编号：××××××）资助。',
]


# ---------------------------------------------------------------------------
# 9. Table data
# ---------------------------------------------------------------------------
TABLE_1_HEADERS = ['项目', '内容']
TABLE_1_ROWS = [
    ['检索数据库', 'PubMed'],
    ['检索时限', '2010年1月—2026年4月'],
    ['纳入文献总数', '344篇'],
    ['实验研究（Ⅰ级）', '待补充'],
    ['临床研究（Ⅱ级）', '待补充'],
    ['网络药理学研究（Ⅲ级）', '待补充'],
    ['语种分布', '英文270篇；中文期刊英文版74篇'],
    ['文献类型问题', '含Review 36篇（已标注，未纳入统计）'],
    ['穴位数', '41个'],
    ['有效基因数', '1,997个'],
    ['基因家族/通路数', '10个'],
    ['疾病/病理实体数', '50种'],
]

TABLE_2_HEADERS = ['排名', '穴位名称', '代码', '文献数（篇）', '占比（%）']
TABLE_2_ROWS = [
    ['1', '足三里', 'ST36', '109', '31.7'],
    ['2', '四白', 'ST2', '26', '7.6'],
    ['3', '百会', 'GV20', '26', '7.6'],
    ['4', '三阴交', 'SP6', '19', '5.5'],
    ['5', '内关', 'PC6', '18', '5.2'],
    ['6', '关元', 'CV4', '18', '5.2'],
    ['7', '攒竹', 'BL2', '13', '3.8'],
    ['8', '上关', 'GB3', '12', '3.5'],
    ['9', '肾俞', 'BL23', '11', '3.2'],
    ['10', '阳陵泉', 'GB34', '11', '3.2'],
]

TABLE_3_HEADERS = ['排名', '基因/家族名称', '文献数（篇）', '备注']
TABLE_3_ROWS = [
    ['1', 'TNF', '55', '炎症因子'],
    ['2', 'IL6', '38', '炎症因子'],
    ['3', 'MTOR', '14', '自噬/代谢'],
    ['4', 'NLRP3', '14', '炎症小体'],
    ['5', 'FOS', '11', '原癌基因'],
    ['6', 'IL10', '11', '抗炎因子'],
    ['7', 'BECN1', '10', '自噬相关'],
    ['8', 'STAT3', '10', '信号转导'],
    ['9', 'MAP1LC3A/B/C', '10', '自噬相关'],
    ['10', 'BCL2', '9', '凋亡调控'],
    ['—', 'AKT家族', '30', '基因家族/通路'],
    ['—', 'PI3K家族', '27', '基因家族/通路'],
    ['—', 'MAPK家族', '17', '基因家族/通路'],
    ['—', 'WNT家族', '9', '基因家族/通路'],
    ['—', 'P38家族', '8', '基因家族/通路'],
    ['—', 'BCL家族', '8', '基因家族/通路'],
    ['—', 'MMP家族', '8', '基因家族/通路'],
    ['—', 'ERK家族', '8', '基因家族/通路'],
    ['—', 'JAK家族', '8', '基因家族/通路'],
    ['—', 'STAT家族', '6', '基因家族/通路'],
]

TABLE_4_HEADERS = ['排名', '疾病/病理实体', '类别', '文献数（篇）']
TABLE_4_ROWS = [
    ['1', '炎症（inflammation）', '病理过程', '157'],
    ['2', '疼痛（pain）', '病理过程', '69'],
    ['3', '应激（stress）', '病理过程', '59'],
    ['4', '肿瘤（tumor）', '疾病', '38'],
    ['5', '脑卒中（stroke）', '疾病', '31'],
    ['6', '关节炎（arthritis）', '疾病', '30'],
    ['7', '戒断综合征（withdrawal）', '病理过程', '25'],
    ['8', '抑郁症（depression）', '病理过程', '20'],
    ['9', '癌症（cancer）', '疾病', '17'],
    ['10', '焦虑（anxiety）', '病理过程', '13'],
]

TABLE_5_HEADERS = ['排名', '节点', '度中心性', '加权度', '中介中心性']
TABLE_5_ROWS = [
    ['1', '足三里', '0.7135', '347', '0.8381'],
    ['2', '四白', '0.1189', '53', '0.1320'],
    ['3', '三阴交', '0.1135', '52', '0.0262'],
    ['4', '百会', '0.0973', '41', '0.1136'],
    ['5', '关元', '0.0595', '35', '0.0398'],
    ['6', '内关', '0.0541', '34', '0.0502'],
    ['7', '上关', '0.0541', '32', '0.0302'],
    ['8', '阳陵泉', '0.0486', '31', '0.0337'],
    ['9', '天枢', '0.0378', '28', '0.0088'],
    ['10', '长强', '0.0324', '26', '0.0432'],
]

TABLE_6_HEADERS = ['穴位', '基因', '共现数', '期望', 'Raw p', 'FDR', '显著性']
TABLE_6_ROWS = [
    ['三阴交', 'TG', '4', '0.33', '9.44×10⁻⁵', '1.89×10⁻³', '**'],
    ['百会', 'GLUT1', '3', '0.23', '3.87×10⁻⁴', '3.87×10⁻³', '**'],
    ['百会', 'GSDMD', '3', '0.38', '3.48×10⁻³', '1.74×10⁻²', '*'],
    ['百会', 'ICH', '3', '0.38', '3.48×10⁻³', '1.74×10⁻²', '*'],
    ['足三里', 'TNF', '22', '17.43', '1.00×10⁻¹', '1.79×10⁻¹', 'ns'],
    ['足三里', 'IL6', '16', '12.04', '1.02×10⁻¹', '1.79×10⁻¹', 'ns'],
    ['足三里', 'TG', '4', '1.90', '8.28×10⁻²', '1.79×10⁻¹', 'ns'],
    ['足三里', 'VEGF', '4', '1.90', '8.28×10⁻²', '1.79×10⁻¹', 'ns'],
    ['天枢', 'TNF', '4', '1.76', '8.11×10⁻²', '1.79×10⁻¹', 'ns'],
    ['阳陵泉', 'TNF', '4', '1.76', '8.11×10⁻²', '1.79×10⁻¹', 'ns'],
]


def fix_text_replacements(doc):
    """Apply all simple text replacements at run level."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in REPLACEMENTS:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
    print(f'[fix_text_replacements] {count} replacements applied.')


def fix_duplicate_sixth(doc):
    """Change the second '第六，' to '第七，'."""
    found = 0
    for para in doc.paragraphs:
        if '第六，' in para.text:
            found += 1
            if found == 2:
                for run in para.runs:
                    if '第六，' in run.text:
                        run.text = run.text.replace('第六，', '第七，', 1)
                        print('[fix_duplicate_sixth] Changed second 第六 to 第七.')
                        return
    print('[fix_duplicate_sixth] WARNING: second 第六 not found.')


def fix_jaccard(doc):
    """Remove Jaccard sentence from method paragraph."""
    for para in doc.paragraphs:
        if JACCARD_PHRASE in para.text:
            # Reconstruct text without the Jaccard clause
            full = para.text
            # The clause spans runs; we will rebuild the paragraph text
            # Find the exact substring to remove
            start = full.find('进一步计算 Jaccard')
            # Find the end of the clause: "文献集合。"
            end_marker = '文献集合。'
            end = full.find(end_marker, start) + len(end_marker)
            if start != -1 and end > start:
                new_text = full[:start] + full[end:]
                # Clear runs and set text on first run
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
                print('[fix_jaccard] Removed Jaccard clause.')
                return
    print('[fix_jaccard] WARNING: Jaccard phrase not found.')


def fix_top_gene_section(doc):
    """Add missing 9th and 10th genes in section 3.3."""
    for para in doc.paragraphs:
        if TOP_GENE_FIX_TARGET in para.text:
            for run in para.runs:
                if TOP_GENE_FIX_TARGET in run.text:
                    run.text = run.text.replace(TOP_GENE_FIX_TARGET, TOP_GENE_FIX_NEW)
                    print('[fix_top_gene_section] Fixed top gene list.')
                    return
    print('[fix_top_gene_section] WARNING: target not found.')


def insert_community_descriptions(doc):
    """Add descriptions for communities 1, 3, 6 and modularity Q."""
    target = find_para(doc, COMMUNITY_INSERT_AFTER)
    if target:
        add_para_after(doc, target, COMMUNITY_EXTRA, bold=False, indent=True, space_after=12)
        print('[insert_community_descriptions] Added community 1/3/6 and Q value.')
    else:
        print('[insert_community_descriptions] WARNING: anchor paragraph not found.')


def insert_dc_note(doc):
    """Add bipartite degree centrality note."""
    target = find_para(doc, DC_NOTE_INSERT_AFTER)
    if target:
        add_para_after(doc, target, DC_NOTE_TEXT, bold=False, indent=True, space_after=6)
        print('[insert_dc_note] Added bipartite DC note.')
    else:
        print('[insert_dc_note] WARNING: anchor paragraph not found.')


def fix_hypergeo_section(doc):
    """Update hypergeometric test results with FDR."""
    fixes = [
        (HYPERGEO_OLD_P1, HYPERGEO_NEW_P1),
        (HYPERGEO_OLD_P2, HYPERGEO_NEW_P2),
        (HYPERGEO_OLD_P3, HYPERGEO_NEW_P3),
        (HYPERGEO_OLD_END, HYPERGEO_NEW_END),
    ]
    for old, new in fixes:
        found = False
        for para in doc.paragraphs:
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        found = True
                        break
                if not found:
                    # Cross-run replacement: rebuild first run
                    para.runs[0].text = para.text.replace(old, new)
                    for run in para.runs[1:]:
                        run.text = ''
                    found = True
                if found:
                    print(f'[fix_hypergeo_section] Replaced hypergeo text.')
                    break
        if not found:
            print(f'[fix_hypergeo_section] WARNING: old text not found: {old[:40]}...')


def insert_declarations(doc):
    """Add declarations before References section."""
    # Find "参考文献" paragraph
    ref_para = None
    for para in doc.paragraphs:
        if para.text.strip().startswith('参考文献'):
            ref_para = para
            break
    if not ref_para:
        print('[insert_declarations] WARNING: 参考文献 not found.')
        return

    # Insert a heading before references
    heading = doc.add_paragraph()
    r = heading.add_run('声明')
    r.bold = True
    rfmt(r, size=12, name='黑体')
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    ref_para._element.addprevious(heading._element)

    prev = heading
    for i in range(0, len(DECLARATIONS), 2):
        title = DECLARATIONS[i]
        content = DECLARATIONS[i+1]
        p = doc.add_paragraph()
        r = p.add_run(title + '　')
        r.bold = True
        rfmt(r, size=10.5)
        r = p.add_run(content)
        rfmt(r, size=10.5)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(6)
        prev._element.addnext(p._element)
        prev = p
    print('[insert_declarations] Added declarations.')


def insert_tables(doc):
    """Insert tables at appropriate locations."""
    # Table 1: after 2.2 纳入与排除标准
    t1_anchor = find_para(doc, '（6）Ahead of print且无完整摘要者。')
    if t1_anchor:
        add_table_after(doc, t1_anchor, TABLE_1_HEADERS, TABLE_1_ROWS, '表1  纳入文献基本信息')
        print('[insert_tables] Added Table 1.')
    else:
        print('[insert_tables] WARNING: Table 1 anchor not found.')

    # Table 2: after 3.2 穴位分布特征
    t2_anchor = find_para(doc, '手厥阴心包经（内关）、任脉（关元、气海）等亦有涉及。')
    if t2_anchor:
        add_table_after(doc, t2_anchor, TABLE_2_HEADERS, TABLE_2_ROWS, '表2  Top 10 高频研究穴位')
        print('[insert_tables] Added Table 2.')
    else:
        print('[insert_tables] WARNING: Table 2 anchor not found.')

    # Table 3: after 3.3 高频关联基因与基因家族
    t3_anchor = find_para(doc, '未来研究可进一步将这些家族拆分为具体基因亚型，以明确各亚型在不同穴位刺激下的表达差异。')
    if t3_anchor:
        add_table_after(doc, t3_anchor, TABLE_3_HEADERS, TABLE_3_ROWS, '表3  Top 10 高频关联基因及基因家族/通路')
        print('[insert_tables] Added Table 3.')
    else:
        print('[insert_tables] WARNING: Table 3 anchor not found.')

    # Table 4: after 3.4 疾病与病理实体分布
    t4_anchor = find_para(doc, '具有明确的应用导向。')
    if t4_anchor:
        add_table_after(doc, t4_anchor, TABLE_4_HEADERS, TABLE_4_ROWS, '表4  Top 10 疾病/病理实体分布')
        print('[insert_tables] Added Table 4.')
    else:
        print('[insert_tables] WARNING: Table 4 anchor not found.')

    # Table 5: after 3.6.1 网络拓扑分析
    t5_anchor = find_para(doc, '加权度（weighted degree）分析显示，足三里以347次共现居首位')
    if t5_anchor:
        # Find the paragraph that ends the 3.6.1 section
        for para in doc.paragraphs:
            if '加权度（weighted degree）分析显示，足三里以347次共现居首位' in para.text:
                t5_anchor = para
                break
        add_table_after(doc, t5_anchor, TABLE_5_HEADERS, TABLE_5_ROWS, '表5  网络拓扑指标 Top 10')
        print('[insert_tables] Added Table 5.')
    else:
        print('[insert_tables] WARNING: Table 5 anchor not found.')

    # Table 6: after 3.6.4 超几何分布显著性检验
    t6_anchor = find_para(doc, '上述显著性检验为穴位-基因共现分析提供了统计学支撑')
    if not t6_anchor:
        t6_anchor = find_para(doc, '不能直接推断调控因果关系。')
    if t6_anchor:
        add_table_after(doc, t6_anchor, TABLE_6_HEADERS, TABLE_6_ROWS, '表6  超几何检验显著性结果（Top 10）')
        print('[insert_tables] Added Table 6.')
    else:
        print('[insert_tables] WARNING: Table 6 anchor not found.')


def main():
    print('Loading document...')
    doc = Document(INPUT_DOC)
    print(f'Loaded: {len(doc.paragraphs)} paragraphs')

    fix_text_replacements(doc)
    fix_duplicate_sixth(doc)
    fix_jaccard(doc)
    fix_top_gene_section(doc)
    fix_hypergeo_section(doc)
    insert_dc_note(doc)
    insert_community_descriptions(doc)
    insert_tables(doc)
    insert_declarations(doc)

    print(f'Saving to {OUTPUT_DOC}...')
    doc.save(OUTPUT_DOC)
    print('Done.')


if __name__ == '__main__':
    main()
