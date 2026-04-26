#!/usr/bin/env python3
"""
更新 v3 Word 文档：
1. 替换已有的 fig1-fig5 图片
2. 在 3.6 节各小节后面插入 fig7、fig8、fig9
"""
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = 'output/paper_figures'
V3_PATH = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3.docx'
V3_OUTPUT = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3.docx'

# 已有图片替换映射
REPLACE_MAP = {
    '图1': f'{OUTPUT_DIR}/fig1_dashboard.png',
    '图2': f'{OUTPUT_DIR}/fig2_top_genes.png',
    '图3': f'{OUTPUT_DIR}/fig3_top_acupoints.png',
    '图4': f'{OUTPUT_DIR}/fig4_network.png',
    '图5': f'{OUTPUT_DIR}/fig5_acupoint_query.png',
}

# 新增图片插入映射（段落文本关键字 -> 图片路径和标题）
INSERT_MAP = {
    '3.6.1　网络拓扑分析': (
        f'{OUTPUT_DIR}/fig7_topology.png',
        '图7  网络拓扑分析：基因度中心性与中介中心性 Top 10',
        15
    ),
    '3.6.2　聚类分析（Louvain社区发现）': (
        f'{OUTPUT_DIR}/fig8_clusters.png',
        '图8  Louvain 聚类分析：各社区节点构成',
        12
    ),
    '3.6.4　超几何分布显著性检验': (
        f'{OUTPUT_DIR}/fig9_hypergeo.png',
        '图9  超几何检验：Top 10 显著穴位-基因对',
        12
    ),
}

def has_image(para):
    return len(para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0

def insert_picture_after(doc, target_para, picture_path, width_cm, caption_text):
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(picture_path, width=Cm(width_cm))
    
    cap_para = doc.add_paragraph(caption_text)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(12)
    for run in cap_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
    
    target_para._element.addnext(pic_para._element)
    pic_para._element.addnext(cap_para._element)

def main():
    doc = Document(V3_PATH)
    
    # Step 1: 替换已有图片
    replaced = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        matched_key = None
        for key in REPLACE_MAP:
            if text.startswith(key):
                matched_key = key
                break
        
        if not matched_key:
            continue
        
        pic_path = REPLACE_MAP[matched_key]
        # 向前查找最近的包含图片的段落
        target_para = None
        for j in range(i-1, max(-1, i-6), -1):
            if has_image(doc.paragraphs[j]):
                target_para = doc.paragraphs[j]
                break
        
        if target_para is None:
            print(f'段落 {i} ({matched_key}): 未找到前面的图片段落')
            continue
        
        # 删除所有 run 并插入新图片
        for run in target_para.runs:
            run._element.getparent().remove(run._element)
        target_para.add_run().add_picture(pic_path, width=Cm(15))
        target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        replaced.append((i, matched_key))
        print(f'段落 {i} ({matched_key}): 已替换')
    
    # Step 2: 在 3.6 节各小节后面插入新图片
    inserted = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text not in INSERT_MAP:
            continue
        
        pic_path, caption, width = INSERT_MAP[text]
        if not os.path.exists(pic_path):
            print(f'段落 {i} ({text}): 图片不存在 {pic_path}')
            continue
        
        # 找到该标题后面的第一个有文字的段落
        found = False
        insert_after = None
        for p in doc.paragraphs:
            if found and p.text.strip():
                insert_after = p
                break
            if p._element is para._element:
                found = True
        
        if insert_after is None:
            print(f'段落 {i} ({text}): 未找到后续段落')
            continue
        
        insert_picture_after(doc, insert_after, pic_path, width, caption)
        inserted.append((i, text))
        print(f'段落 {i} ({text}): 已插入 {caption}')
    
    doc.save(V3_OUTPUT)
    print(f'\n已保存: {V3_OUTPUT}')
    print(f'替换 {len(replaced)} 张已有图片，插入 {len(inserted)} 张新图片')

if __name__ == '__main__':
    main()
