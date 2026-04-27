#!/usr/bin/env python3
"""
Patch 3: Insert enrichment figures and remove unnecessary tables.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_fixed.docx'
GO_PIC = 'output/analysis_results/fig_go_bp_enrichment.png'
KEGG_PIC = 'output/analysis_results/fig_kegg_enrichment.png'


def rfmt(run, size=10.5, name='宋体'):
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def insert_picture_after(doc, target_para, pic_path, width_cm, caption_text):
    if not os.path.exists(pic_path):
        print(f'WARNING: {pic_path} not found')
        return False
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(pic_path, width=Cm(width_cm))
    cp = doc.add_paragraph(caption_text)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cp.runs:
        rfmt(r)
    cp.paragraph_format.space_after = Pt(12)
    target_para._element.addnext(pp._element)
    pp._element.addnext(cp._element)
    return True


def find_para(doc, substr):
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    return None


def remove_tables_by_caption(doc, caption_keywords):
    """Remove tables whose following paragraph contains any of the keywords."""
    removed = 0
    body = doc.element.body
    # tables and their following paragraphs are interleaved in body
    for tbl in list(doc.tables):
        # Find the table element index in body
        tbl_idx = None
        for i, child in enumerate(body):
            if child is tbl._element:
                tbl_idx = i
                break
        if tbl_idx is None:
            continue
        # Check next sibling paragraphs for caption
        for sibling in body[tbl_idx+1:tbl_idx+4]:
            if sibling.tag.endswith('p'):
                para_text = ''.join(t.text or '' for t in sibling.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
                for kw in caption_keywords:
                    if kw in para_text:
                        # Remove caption paragraph
                        body.remove(sibling)
                        # Remove table
                        body.remove(tbl._element)
                        removed += 1
                        print(f'[remove_tables] Removed: {kw}')
                        break
                break
    print(f'[remove_tables] Total removed: {removed}')


def main():
    doc = Document(DOC)

    # 1. Insert enrichment figures after 3.6.3 section
    anchor = find_para(doc, 'Th17细胞分化（Th17 Cell Differentiation')
    if anchor:
        # Insert after the paragraph that ends 3.6.3 text
        # Find the paragraph containing the last KEGG result
        for p in doc.paragraphs:
            if 'JAK-STAT通路的显著富集提示' in p.text:
                anchor = p
                break
        if insert_picture_after(doc, anchor, GO_PIC, 14, '图10  GO生物学过程富集分析气泡图（Top 10）'):
            print('[main] Inserted GO enrichment figure.')
        # Find the caption paragraph of GO figure
        go_cap = None
        for p in doc.paragraphs:
            if '图10' in p.text:
                go_cap = p
                break
        if go_cap and insert_picture_after(doc, go_cap, KEGG_PIC, 14, '图11  KEGG通路富集分析气泡图（Top 10）'):
            print('[main] Inserted KEGG enrichment figure.')
    else:
        print('[main] WARNING: 3.6.3 anchor not found.')

    # 2. Remove tables 2, 3, 4, 6 by caption keywords
    remove_tables_by_caption(doc, [
        '表2  Top 10 高频研究穴位',
        '表3  Top 10 高频关联基因及基因家族/通路',
        '表4  Top 10 疾病/病理实体分布',
        '表6  超几何检验显著性结果（Top 10）',
    ])

    doc.save(DOC)
    print(f'Saved: {DOC}')


if __name__ == '__main__':
    main()
