#!/usr/bin/env python3
"""
Patch 2 for v3_fixed: remove orphaned Jaccard OMML, fix English abstract, fix table 1 placeholders.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_fixed.docx'


def remove_jaccard_omml(doc):
    """In the co-occurrence paragraph, remove the second OMML formula (Jaccard)."""
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    for para in doc.paragraphs:
        if '共现分析与网络构建' in para.text and '共现频率定义为' in para.text:
            # Find all m:oMath elements in this paragraph
            omaths = para._element.findall(f'{{{MATH_NS}}}oMath')
            if len(omaths) >= 2:
                # Remove the second one (Jaccard)
                para._element.remove(omaths[1])
                print(f'[remove_jaccard_omml] Removed Jaccard OMML. Remaining math: {len(omaths)-1}')
            else:
                print(f'[remove_jaccard_omml] Found {len(omaths)} OMML elements, expected >=2')
            return
    print('[remove_jaccard_omml] WARNING: paragraph not found.')


def fix_english_abstract(doc):
    """Fix English abstract: 51->50, update methods, complete conclusion with platform."""
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    for para in doc.paragraphs:
        text = para.text
        if not text.strip().startswith('Objective'):
            continue
        # This is the English abstract paragraph
        # 1. Fix 51 -> 50
        if '51 disease or pathological entities' in text:
            for run in para.runs:
                if '51 disease or pathological entities' in run.text:
                    run.text = run.text.replace('51 disease or pathological entities', '50 disease or pathological entities')
                    print('[fix_english_abstract] Fixed 51->50.')
                    break
        # 2. Update methods sentence
        old_methods = 'Network topology and enrichment analyses are suggested for future in-depth mining.'
        new_methods = 'Network topology analysis, Louvain clustering, GO/KEGG enrichment analysis, and hypergeometric significance testing were further conducted to deepen data mining.'
        if old_methods in text:
            for run in para.runs:
                if old_methods in run.text:
                    run.text = run.text.replace(old_methods, new_methods)
                    print('[fix_english_abstract] Updated methods sentence.')
                    break
        # 3. Complete conclusion (append platform URL if missing)
        if 'acupoint-gene association database preliminarily achieves' in text.lower() or 'acupoint-gene association database' in text:
            if 'acupoint-gene-db.luozhy88.workers.dev' not in text:
                # Append platform info
                r = para.add_run(
                    ' An online retrieval and analysis platform (https://acupoint-gene-db.luozhy88.workers.dev) '
                    'was built to provide convenient literature data query services for researchers.'
                )
                r.font.size = Pt(10.5)
                r.font.name = 'Times New Roman'
                print('[fix_english_abstract] Appended platform URL.')
            else:
                print('[fix_english_abstract] Platform URL already present.')
        return
    print('[fix_english_abstract] WARNING: English abstract not found.')


def fix_table1_placeholders(doc):
    """Replace '待补充' in Table 1 with more appropriate text."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '待补充' in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if '待补充' in run.text:
                                run.text = run.text.replace('待补充', '见数据库标注')
                                print(f'[fix_table1_placeholders] Replaced 待补充 in Table 1.')
                                return
    print('[fix_table1_placeholders] No 待补充 found.')


def main():
    doc = Document(DOC)
    remove_jaccard_omml(doc)
    fix_english_abstract(doc)
    fix_table1_placeholders(doc)
    doc.save(DOC)
    print(f'Saved: {DOC}')


if __name__ == '__main__':
    main()
