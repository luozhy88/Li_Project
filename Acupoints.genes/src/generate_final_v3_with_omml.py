#!/usr/bin/env python3
"""
生成最终版 v3 论文，方法学部分的公式使用 Word OMML 数学公式。
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.oxml.ns import qn

OUTPUT_DIR = 'output/paper_figures'
BASE_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v2.docx'
OUTPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_final.docx'
PLATFORM_URL = 'https://acupoint-gene-db.luozhy88.workers.dev'

# OMML namespace
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def om_run(text):
    """普通文本run"""
    return f'<m:r xmlns:m="{MATH_NS}"><m:t>{text}</m:t></m:r>'


def om_sub(base, sub):
    """下标"""
    return f'<m:sSub xmlns:m="{MATH_NS}"><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'


def om_frac(num, den):
    """分数"""
    return f'<m:f xmlns:m="{MATH_NS}"><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'


def om_sum(sub_text, sup_text, expr):
    """求和符号"""
    return (
        f'<m:nary xmlns:m="{MATH_NS}">'
        f'<m:naryPr><m:accChr m:val="∑"/><m:limLoc m:val="underOver"/></m:naryPr>'
        f'<m:sub><m:r><m:t>{sub_text}</m:t></m:r></m:sub>'
        f'<m:sup><m:r><m:t>{sup_text}</m:t></m:r></m:sup>'
        f'<m:e>{expr}</m:e>'
        f'</m:nary>'
    )


def om_delim(left, right, expr):
    """括号定界符"""
    return (
        f'<m:d xmlns:m="{MATH_NS}">'
        f'<m:dPr><m:begChr m:val="{left}"/><m:endChr m:val="{right}"/></m:dPr>'
        f'<m:e>{expr}</m:e>'
        f'</m:d>'
    )


def om_sqrt(expr):
    """平方根"""
    return f'<m:rad xmlns:m="{MATH_NS}"><m:radPr><m:hideDegree m:val="1"/></m:radPr><m:e>{expr}</m:e></m:rad>'


def make_omml(*parts):
    """组合多个OMML片段为一个完整的 m:oMath 元素"""
    inner = ''.join(parts)
    return f'<m:oMath xmlns:m="{MATH_NS}">{inner}</m:oMath>'


def add_omml_to_para(para, omml_xml):
    """将OMML添加到段落的末尾"""
    elem = parse_xml(omml_xml)
    para._element.append(elem)


# ============================================================
# 构建关键公式的OMML
# ============================================================

# 公式1: f(x) = Σ_{i=1}^{N} I(x ∈ D_i)
FORMULA_1 = make_omml(
    om_run('f(x)='),
    om_sum('i=1', 'N',
        om_run('I(x∈') + om_sub(om_run('D'), om_run('i')) + om_run(')')
    )
)

# 公式2: C(a,g) = Σ_{i=1}^{N} I(a ∈ D_i ∧ g ∈ D_i)
FORMULA_2 = make_omml(
    om_run('C(a,g)='),
    om_sum('i=1', 'N',
        om_run('I(a∈') + om_sub(om_run('D'), om_run('i')) + om_run(' ∧ g∈') + om_sub(om_run('D'), om_run('i')) + om_run(')')
    )
)

# 公式3: J(a,g) = |D_a ∩ D_g| / |D_a ∪ D_g|
FORMULA_3 = make_omml(
    om_run('J(a,g)='),
    om_frac(
        om_run('|') + om_sub(om_run('D'), om_run('a')) + om_run('∩') + om_sub(om_run('D'), om_run('g')) + om_run('|'),
        om_run('|') + om_sub(om_run('D'), om_run('a')) + om_run('∪') + om_sub(om_run('D'), om_run('g')) + om_run('|')
    )
)

# 公式4: k(v) = deg(v) / (|V| - 1)
FORMULA_4 = make_omml(
    om_run('k(v)='),
    om_frac(
        om_run('deg(v)'),
        om_run('|V|−1')
    )
)

# 公式5: s(v) = Σ_{u∈N(v)} w(u,v)
FORMULA_5 = make_omml(
    om_run('s(v)='),
    om_sum('u∈N(v)', '', om_run('w(u,v)'))
)

# 公式6: B(v) = Σ_{s≠v≠t} σ_{st}(v) / σ_{st}
FORMULA_6 = make_omml(
    om_run('B(v)='),
    om_sum('s≠v≠t', '',
        om_frac(
            om_sub(om_run('σ'), om_run('st')) + om_run('(v)'),
            om_sub(om_run('σ'), om_run('st'))
        )
    )
)

# 公式7: A·x = λ·x
FORMULA_7 = make_omml(
    om_run('A·x=λ·x')
)

# 公式8: Q = 1/(2m) Σ_{ij} [A_{ij} - (k_i·k_j)/(2m)] δ(c_i, c_j)
FORMULA_8 = make_omml(
    om_run('Q='),
    om_frac(om_run('1'), om_run('2m')),
    om_sum('ij', '',
        om_delim('[', ']',
            om_sub(om_run('A'), om_run('ij')) + om_run('−') +
            om_frac(
                om_sub(om_run('k'), om_run('i')) + om_run('·') + om_sub(om_run('k'), om_run('j')),
                om_run('2m')
            )
        ) +
        om_run(' δ') +
        om_delim('(', ')',
            om_sub(om_run('c'), om_run('i')) + om_run(',') + om_sub(om_run('c'), om_run('j'))
        )
    )
)

# 公式9: m = Σ_{ij} A_{ij} / 2
FORMULA_9 = make_omml(
    om_run('m='),
    om_frac(
        om_sum('ij', '', om_sub(om_run('A'), om_run('ij'))),
        om_run('2')
    )
)

# 公式10: P(X=k) = C(M,k)·C(N-M,n-k) / C(N,n)
FORMULA_10 = make_omml(
    om_run('P(X=k)='),
    om_frac(
        om_run('C(M,k)·C(N−M,n−k)'),
        om_run('C(N,n)')
    )
)

# 公式11: P(X≥k) = Σ_{x=k}^{min(n,M)} C(M,x)·C(N-M,n-x) / C(N,n)
FORMULA_11 = make_omml(
    om_run('P(X≥k)='),
    om_sum('x=k', 'min(n,M)',
        om_frac(
            om_run('C(M,x)·C(N−M,n−x)'),
            om_run('C(N,n)')
        )
    )
)


def add_text_para(doc, text, bold_prefix=None, indent=True):
    """添加标准正文段落"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_mixed_para(doc, bold_prefix, text_before, omml_formula, text_after, indent=True):
    """添加混合段落：文本 + OMML公式 + 文本"""
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = '宋体'
    if text_before:
        r = p.add_run(text_before)
        r.font.size = Pt(10.5)
        r.font.name = '宋体'
    if omml_formula:
        add_omml_to_para(p, omml_formula)
    if text_after:
        r = p.add_run(text_after)
        r.font.size = Pt(10.5)
        r.font.name = '宋体'
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    return p


def insert_picture_after(doc, target_para, picture_path, width_cm, caption_text):
    """在指定段落后插入图片和图注"""
    pic_para = doc.add_paragraph()
    pic_para.alignment = 1  # CENTER
    pic_para.add_run().add_picture(picture_path, width=Cm(width_cm))

    cap_para = doc.add_paragraph(caption_text)
    cap_para.alignment = 1  # CENTER
    cap_para.paragraph_format.space_after = Pt(12)
    for run in cap_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'

    target_para._element.addnext(pic_para._element)
    pic_para._element.addnext(cap_para._element)


def find_para_by_text(doc, text_substr):
    for p in doc.paragraphs:
        if text_substr in p.text:
            return p
    return None


FIGURE_PLAN = [
    ('3.1　文献收录概况',     f'{OUTPUT_DIR}/fig1_dashboard.png',      '图1  数据库统计概览与文献年份分布', 15),
    ('3.2　穴位分布特征',     f'{OUTPUT_DIR}/fig3_top_acupoints.png',  '图2  Top 10 高频研究穴位', 12),
    ('3.3　高频关联基因与基因家族', f'{OUTPUT_DIR}/fig2_top_genes.png',      '图3  Top 10 高频关联基因', 12),
    ('3.5　穴位-基因共现分析', f'{OUTPUT_DIR}/fig4_network.png',        '图4  穴位-基因共现网络图（共现频次≥3）', 15),
    ('3.6.1　网络拓扑分析',    f'{OUTPUT_DIR}/fig7_topology.png',       '图5  网络拓扑分析：基因度中心性与中介中心性 Top 10', 15),
    ('3.6.2　聚类分析（Louvain社区发现）', f'{OUTPUT_DIR}/fig8_clusters.png', '图6  Louvain 聚类分析：各社区节点构成', 12),
    ('3.6.4　超几何分布显著性检验', f'{OUTPUT_DIR}/fig9_hypergeo.png', '图7  超几何检验：Top 10 显著穴位-基因对', 12),
    ('4　讨论',              f'{OUTPUT_DIR}/fig5_acupoint_query.png', '图8  数据库平台"穴位→基因"查询模块示例（以足三里为例）', 14),
]

PLATFORM_FIG2 = (f'{OUTPUT_DIR}/fig6_gene_query.png', '图9  数据库平台"基因→穴位"查询模块示例（以TNF为例）', 12)


def main():
    doc = Document(BASE_DOC)

    # ============================================================
    # Step 1: 在摘要添加平台网址
    # ============================================================
    for p in doc.paragraphs:
        if p.text.strip().startswith('目的') and '为揭示针灸作用的分子机制' in p.text:
            p.add_run(f'同时，基于该数据库构建了在线检索与分析平台（{PLATFORM_URL}），为研究者提供便捷的文献数据查询服务。')
            break

    # ============================================================
    # Step 2: 在方法学添加平台说明
    # ============================================================
    for p in doc.paragraphs:
        if 'SQLite' in p.text and 'articles（文献表）' in p.text:
            platform_para = doc.add_paragraph(
                f'此外，基于Cloudflare Worker构建了轻量级Web数据检索平台（{PLATFORM_URL}），'
                '提供数据概览、穴位查基因、基因查穴位、共现网络可视化及文献浏览等功能，'
                '便于研究者在线查询与浏览数据库内容。平台采用服务器端渲染（SSR）策略，'
                '所有数据以JSON格式内联部署，无需后端数据库连接，确保高可用性与快速响应。'
            )
            platform_para.paragraph_format.first_line_indent = Cm(0.74)
            for run in platform_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
            p._element.addnext(platform_para._element)
            break

    # ============================================================
    # Step 3: 替换数据分析方法段落（使用OMML公式）
    # ============================================================
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith('数据分析包括：'):
            target_idx = i
            break

    if target_idx is None:
        print("警告：未找到'数据分析包括：'段落")
        return

    target_p = doc.paragraphs[target_idx]
    insert_after = target_p

    # （1）频次统计
    p1 = doc.add_paragraph()
    r1 = p1.add_run('（1）频次统计：')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = '宋体'
    r1b = p1.add_run('运用 Python 3.13 的 pandas（v3.0.1）与 collections 模块，对穴位名称、基因符号及疾病实体的出现频次进行计数统计。')
    r1b.font.size = Pt(10.5)
    r1b.font.name = '宋体'
    add_omml_to_para(p1, FORMULA_1)
    r1c = p1.add_run('，其中 N 为纳入文献总数（N=344），')
    r1c.font.size = Pt(10.5)
    r1c.font.name = '宋体'
    add_omml_to_para(p1, make_omml(om_sub(om_run('D'), om_run('i'))))
    r1d = p1.add_run(' 为第 i 篇文献提取的实体集合，I(·) 为指示函数，当实体 x 出现在文献 ')
    r1d.font.size = Pt(10.5)
    r1d.font.name = '宋体'
    add_omml_to_para(p1, make_omml(om_sub(om_run('D'), om_run('i'))))
    r1e = p1.add_run(' 中时取值为 1，否则为 0。对于基因家族（如 AKT、PI3K 等），在频次统计中单独标注为「基因家族/通路」，不纳入单一基因的频次排序，以避免重复计数。')
    r1e.font.size = Pt(10.5)
    r1e.font.name = '宋体'
    p1.paragraph_format.first_line_indent = Cm(0.74)
    p1.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p1._element)
    insert_after = p1

    # （2）共现分析与网络构建
    p2 = doc.add_paragraph()
    r2 = p2.add_run('（2）共现分析与网络构建：')
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.name = '宋体'
    r2b = p2.add_run('采用共现频率（co-occurrence frequency）衡量穴位-基因关联强度。对于穴位 a 与基因 g，其共现频率定义为：')
    r2b.font.size = Pt(10.5)
    r2b.font.name = '宋体'
    add_omml_to_para(p2, FORMULA_2)
    r2c = p2.add_run('，即同时包含穴位 a 和基因 g 的文献篇数。进一步计算 Jaccard 相似系数以消除高频实体的偏倚：')
    r2c.font.size = Pt(10.5)
    r2c.font.name = '宋体'
    add_omml_to_para(p2, FORMULA_3)
    r2d = p2.add_run('，其中 ')
    r2d.font.size = Pt(10.5)
    r2d.font.name = '宋体'
    add_omml_to_para(p2, make_omml(om_sub(om_run('D'), om_run('a'))))
    r2e = p2.add_run('、')
    r2e.font.size = Pt(10.5)
    r2e.font.name = '宋体'
    add_omml_to_para(p2, make_omml(om_sub(om_run('D'), om_run('g'))))
    r2f = p2.add_run(' 分别为包含穴位 a 和基因 g 的文献集合。共现网络以 NetworkX（v3.6.1）构建为无向二分图 G=(V,E)，节点集合 V=A∪G 分为穴位层 A 与基因层 G，边权 w(a,g)=C(a,g)（共现频次≥3 的边纳入可视化）。网络可视化采用 Kamada-Kawai 力导向布局，节点大小映射加权度中心性，边透明度映射共现频次。')
    r2f.font.size = Pt(10.5)
    r2f.font.name = '宋体'
    p2.paragraph_format.first_line_indent = Cm(0.74)
    p2.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p2._element)
    insert_after = p2

    # （3）网络拓扑分析
    p3 = doc.add_paragraph()
    r3 = p3.add_run('（3）网络拓扑分析：')
    r3.bold = True
    r3.font.size = Pt(10.5)
    r3.font.name = '宋体'
    r3b = p3.add_run('基于 NetworkX（v3.6.1）计算以下拓扑指标：①度中心性（degree centrality）：')
    r3b.font.size = Pt(10.5)
    r3b.font.name = '宋体'
    add_omml_to_para(p3, FORMULA_4)
    r3c = p3.add_run('，其中 deg(v) 为节点 v 的邻居数，反映节点的局部连接密度；②加权度中心性（weighted degree）：')
    r3c.font.size = Pt(10.5)
    r3c.font.name = '宋体'
    add_omml_to_para(p3, FORMULA_5)
    r3d = p3.add_run('，即关联边权之和，反映节点在共现频次层面的影响力；③中介中心性（betweenness centrality）：')
    r3d.font.size = Pt(10.5)
    r3d.font.name = '宋体'
    add_omml_to_para(p3, FORMULA_6)
    r3e = p3.add_run('，其中 ')
    r3e.font.size = Pt(10.5)
    r3e.font.name = '宋体'
    add_omml_to_para(p3, make_omml(om_sub(om_run('σ'), om_run('st'))))
    r3f = p3.add_run(' 为节点 s 到 t 的最短路径总数，')
    r3f.font.size = Pt(10.5)
    r3f.font.name = '宋体'
    add_omml_to_para(p3, make_omml(om_sub(om_run('σ'), om_run('st')) + om_run('(v)')))
    r3g = p3.add_run(' 为经过 v 的最短路径数，反映节点作为信息流通「桥梁」的能力；④特征向量中心性（eigenvector centrality）：')
    r3g.font.size = Pt(10.5)
    r3g.font.name = '宋体'
    add_omml_to_para(p3, FORMULA_7)
    r3h = p3.add_run('，求解邻接矩阵 A 的主特征向量，节点得分与其邻居的中心性加权相关。上述计算均调用 NetworkX 内置函数（degree_centrality、betweenness_centrality、eigenvector_centrality），其中 betweenness 计算采用 Brandes 算法（时间复杂度 O(|V||E|)）。')
    r3h.font.size = Pt(10.5)
    r3h.font.name = '宋体'
    p3.paragraph_format.first_line_indent = Cm(0.74)
    p3.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p3._element)
    insert_after = p3

    # （4）Louvain聚类
    p4 = doc.add_paragraph()
    r4 = p4.add_run('（4）聚类分析（Louvain 社区发现）：')
    r4.bold = True
    r4.font.size = Pt(10.5)
    r4.font.name = '宋体'
    r4b = p4.add_run('采用 Louvain 社区发现算法（python-louvain 0.16）对共现网络进行模块划分。该算法以模块度（modularity）最大化为目标：')
    r4b.font.size = Pt(10.5)
    r4b.font.name = '宋体'
    add_omml_to_para(p4, FORMULA_8)
    r4c = p4.add_run('，其中 ')
    r4c.font.size = Pt(10.5)
    r4c.font.name = '宋体'
    add_omml_to_para(p4, FORMULA_9)
    r4d = p4.add_run(' 为网络总边权，')
    r4d.font.size = Pt(10.5)
    r4d.font.name = '宋体'
    add_omml_to_para(p4, make_omml(om_sub(om_run('A'), om_run('ij'))))
    r4e = p4.add_run(' 为节点 i,j 的邻接矩阵元素，')
    r4e.font.size = Pt(10.5)
    r4e.font.name = '宋体'
    add_omml_to_para(p4, make_omml(om_sub(om_run('k'), om_run('i'))))
    r4f = p4.add_run('、')
    r4f.font.size = Pt(10.5)
    r4f.font.name = '宋体'
    add_omml_to_para(p4, make_omml(om_sub(om_run('k'), om_run('j'))))
    r4g = p4.add_run(' 为节点度，δ(·) 为克罗内克函数（当 i,j 属于同一社区时取 1，否则为 0）。算法通过两阶段迭代实现：阶段一将每个节点视为独立社区，依次将节点移动到能使模块度增量 ΔQ 最大的邻居社区；阶段二将同一社区内节点收缩为超级节点，重构网络后重复阶段一，直至模块度不再提升。本研究设置分辨率参数 γ=1.0（默认），最终得到 8 个社区。')
    r4g.font.size = Pt(10.5)
    r4g.font.name = '宋体'
    p4.paragraph_format.first_line_indent = Cm(0.74)
    p4.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p4._element)
    insert_after = p4

    # （5）功能富集分析
    p5 = doc.add_paragraph()
    r5 = p5.add_run('（5）功能富集分析：')
    r5.bold = True
    r5.font.size = Pt(10.5)
    r5.font.name = '宋体'
    r5b = p5.add_run('选取网络中度中心性排名前 50 的基因作为核心基因集，通过 Enrichr 在线数据库（https://maayanlab.cloud/Enrichr/）进行 GO（Gene Ontology）生物学过程（Biological Process）与 KEGG 通路富集分析。富集显著性采用 Fisher 精确检验（右侧检验），计算公式为：')
    r5b.font.size = Pt(10.5)
    r5b.font.name = '宋体'
    add_omml_to_para(p5, FORMULA_10)
    r5c = p5.add_run('，其中 N 为背景基因总数（本数据库有效基因 1997 个），M 为某一 GO/KEGG 条目注释的基因数，n 为核心基因集大小（n=50），k 为核心基因集中属于该条目的基因数。累积 p 值计算为：')
    r5c.font.size = Pt(10.5)
    r5c.font.name = '宋体'
    add_omml_to_para(p5, FORMULA_11)
    r5d = p5.add_run('。当 p<0.05 时，认为该通路富集显著。多重检验校正采用 Benjamini-Hochberg 法计算校正后 p 值（FDR），以 FDR<0.05 作为显著性阈值。')
    r5d.font.size = Pt(10.5)
    r5d.font.name = '宋体'
    p5.paragraph_format.first_line_indent = Cm(0.74)
    p5.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p5._element)
    insert_after = p5

    # （6）超几何分布检验
    p6 = doc.add_paragraph()
    r6 = p6.add_run('（6）显著性检验（超几何分布）：')
    r6.bold = True
    r6.font.size = Pt(10.5)
    r6.font.name = '宋体'
    r6b = p6.add_run('为评估穴位-基因共现是否显著高于随机期望，采用超几何分布检验（scipy.stats.hypergeom，v1.17.1）。设总体文献数 N=344，其中包含特定基因 g 的文献数为 M（基因 g 的边际频次），从总体中随机抽取 n 篇文献（n 为包含特定穴位 a 的文献数，即穴位 a 的边际频次），观察到的共现文献数为 k=C(a,g)。超几何分布概率质量函数为：')
    r6b.font.size = Pt(10.5)
    r6b.font.name = '宋体'
    add_omml_to_para(p6, FORMULA_10)
    r6c = p6.add_run('。单侧 p 值计算为：')
    r6c.font.size = Pt(10.5)
    r6c.font.name = '宋体'
    add_omml_to_para(p6, FORMULA_11)
    r6d = p6.add_run('。当 p<0.05 时，认为该穴位-基因对的共现显著高于随机期望。需要强调的是，显著性检验评估的是「研究关注度」的统计学意义，而非直接的生物学因果关系；共现显著仅表明该穴位-基因对在文献中被共同提及的频率高于随机水平，不能推断穴位对基因表达具有直接调控作用。')
    r6d.font.size = Pt(10.5)
    r6d.font.name = '宋体'
    p6.paragraph_format.first_line_indent = Cm(0.74)
    p6.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p6._element)
    insert_after = p6

    # （7）软件环境
    p7 = doc.add_paragraph()
    r7 = p7.add_run('（7）软件环境：')
    r7.bold = True
    r7.font.size = Pt(10.5)
    r7.font.name = '宋体'
    r7b = p7.add_run('本研究数据分析均在 Python 3.13 环境下完成，主要依赖包及版本包括：pandas 3.0.1（数据处理）、NetworkX 3.6.1（网络构建与拓扑分析）、python-louvain 0.16（社区发现）、scipy 1.17.1（统计检验）、matplotlib 3.10.8（可视化）、numpy 2.4.3（数值计算）。富集分析通过 Enrichr API 在线提交基因列表获取结果。在解读所有分析结果时，须严格区分「共现关注度」与「实验验证的调控关系」，本数据库的定位是文献数据整合与假设生成工具，而非因果推断平台。')
    r7b.font.size = Pt(10.5)
    r7b.font.name = '宋体'
    p7.paragraph_format.first_line_indent = Cm(0.74)
    p7.paragraph_format.space_after = Pt(6)
    insert_after._element.addnext(p7._element)

    # 删除原简短段落
    target_p._element.getparent().remove(target_p._element)

    # ============================================================
    # Step 4: 插入图片
    # ============================================================
    inserted = []
    for section_key, pic_path, caption, width in FIGURE_PLAN:
        section_para = find_para_by_text(doc, section_key)
        if section_para is None:
            print(f'警告：未找到小节 "{section_key}"')
            continue

        found = False
        target = None
        for p in doc.paragraphs:
            if found and p.text.strip():
                target = p
                break
            if p._element is section_para._element:
                found = True

        if target is None:
            print(f'警告：小节 "{section_key}" 后无内容段落')
            continue

        if not os.path.exists(pic_path):
            print(f'警告：图片不存在 {pic_path}')
            continue

        insert_picture_after(doc, target, pic_path, width, caption)
        inserted.append((section_key, caption))

        if section_key == '4　讨论':
            pic2_path, cap2, w2 = PLATFORM_FIG2
            if os.path.exists(pic2_path):
                last_cap = None
                for p in doc.paragraphs:
                    if '图8' in p.text:
                        last_cap = p
                if last_cap:
                    insert_picture_after(doc, last_cap, pic2_path, w2, cap2)
                    inserted.append((section_key, cap2))

    # ============================================================
    # Step 5: 在讨论部分添加平台使用价值
    # ============================================================
    for p in doc.paragraphs:
        if p.text.strip().startswith('本研究构建了针灸穴位-基因关联数据库'):
            p.add_run(
                f'数据库平台（{PLATFORM_URL}）提供了在线检索、共现网络可视化及文献溯源功能，'
                '可供研究者快速获取特定穴位或基因的关联信息，为后续实验设计提供文献依据。'
            )
            break

    doc.save(OUTPUT_DOC)
    print(f'\n最终文档已保存: {OUTPUT_DOC}')
    print(f'共插入 {len(inserted)} 张图片')


if __name__ == '__main__':
    main()
