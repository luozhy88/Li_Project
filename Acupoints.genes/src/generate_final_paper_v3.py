#!/usr/bin/env python3
"""
生成可以直接投稿的最终版 v3 论文（Word）
- 从干净的 v2 基线文档开始
- 按照正确的论文顺序插入图1-图9
- 修正所有图号
- 添加平台网址
"""
import os
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = 'output/paper_figures'
BASE_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v2.docx'
OUTPUT_DOC = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3_final.docx'
PLATFORM_URL = 'https://acupoint-gene-db.luozhy88.workers.dev'

# 图片插入计划（按论文出现顺序）
FIGURE_PLAN = [
    # (section_keyword, pic_path, caption, width_cm)
    ('3.1　文献收录概况',     f'{OUTPUT_DIR}/fig1_dashboard.png',      '图1  数据库统计概览与文献年份分布', 15),
    ('3.2　穴位分布特征',     f'{OUTPUT_DIR}/fig3_top_acupoints.png',  '图2  Top 10 高频研究穴位', 12),
    ('3.3　高频关联基因与基因家族', f'{OUTPUT_DIR}/fig2_top_genes.png',      '图3  Top 10 高频关联基因', 12),
    ('3.5　穴位-基因共现分析', f'{OUTPUT_DIR}/fig4_network.png',        '图4  穴位-基因共现网络图（共现频次≥3）', 15),
    ('3.6.1　网络拓扑分析',    f'{OUTPUT_DIR}/fig7_topology.png',       '图5  网络拓扑分析：基因度中心性与中介中心性 Top 10', 15),
    ('3.6.2　聚类分析（Louvain社区发现）', f'{OUTPUT_DIR}/fig8_clusters.png', '图6  Louvain 聚类分析：各社区节点构成', 12),
    ('3.6.4　超几何分布显著性检验', f'{OUTPUT_DIR}/fig9_hypergeo.png', '图7  超几何检验：Top 10 显著穴位-基因对', 12),
    ('4　讨论',              f'{OUTPUT_DIR}/fig5_acupoint_query.png', '图8  数据库平台"穴位→基因"查询模块示例（以足三里为例）', 14),
]

# 平台界面第二张图也插入到讨论部分
PLATFORM_FIG2 = (
    f'{OUTPUT_DIR}/fig6_gene_query.png',
    '图9  数据库平台"基因→穴位"查询模块示例（以TNF为例）',
    12
)


def insert_picture_after(doc, target_para, picture_path, width_cm, caption_text):
    """在指定段落后插入图片和图注"""
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(picture_path, width=Cm(width_cm))
    
    cap_para = doc.add_paragraph(caption_text)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_after = Pt(12)
    for run in cap_para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
    
    target_para._element.addnext(pic_para._element)
    pic_para._element.addnext(cap_para._element)


def find_para_by_text(doc, text_substr):
    for p in doc.paragraphs:
        if text_substr in p.text:
            return p
    return None


def add_platform_url_to_text(doc, search_text, append_text):
    """在包含特定文本的段落后追加内容"""
    for p in doc.paragraphs:
        if search_text in p.text:
            p.add_run(append_text)
            return True
    return False


def main():
    doc = Document(BASE_DOC)
    
    # Step 1: 在摘要的结论部分添加平台网址
    # 找到"目的"段落（中英文摘要通常在一个段落里）
    for p in doc.paragraphs:
        if p.text.strip().startswith('目的') and '为揭示针灸作用的分子机制' in p.text:
            # 在段落末尾添加平台信息
            p.add_run(f'同时，基于该数据库构建了在线检索与分析平台（{PLATFORM_URL}），为研究者提供便捷的文献数据查询服务。')
            print('已在摘要中添加平台网址')
            break
    
    # Step 2: 在方法学部分添加平台说明
    # 找到数据库设计段落（含 SQLite 的段落）
    for p in doc.paragraphs:
        if 'SQLite' in p.text and 'articles（文献表）' in p.text:
            # 在该段落后插入新段落说明Web平台
            platform_para = doc.add_paragraph(
                f'此外，基于Cloudflare Worker构建了轻量级Web数据检索平台（{PLATFORM_URL}），'
                '提供数据概览、穴位查基因、基因查穴位、共现网络可视化及文献浏览等功能，'
                '便于研究者在线查询与浏览数据库内容。平台采用服务器端渲染（SSR）策略，'
                '所有数据以JSON格式内联部署，无需后端数据库连接，确保高可用性与快速响应。'
            )
            platform_para.paragraph_format.first_line_indent = Cm(0.74)
            for run in platform_para.runs:
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
            p._element.addnext(platform_para._element)
            print('已在方法学中添加平台说明')
            break
    
    # Step 3: 按顺序插入图片
    inserted = []
    for section_key, pic_path, caption, width in FIGURE_PLAN:
        section_para = find_para_by_text(doc, section_key)
        if section_para is None:
            print(f'警告：未找到小节 "{section_key}"')
            continue
        
        # 找到该标题后面的第一个有文字的段落
        found = False
        target = None
        for p in doc.paragraphs:
            if found and p.text.strip():
                target = p
                break
            if p._element is section_para._element:
                found = True
        
        if target is None:
            print(f'警告：小节 "{section_key}" 后无内容段落')
            continue
        
        if not os.path.exists(pic_path):
            print(f'警告：图片不存在 {pic_path}')
            continue
        
        insert_picture_after(doc, target, pic_path, width, caption)
        inserted.append((section_key, caption))
        print(f'已插入 {caption}')
        
        # 如果这是讨论部分（4　讨论），再插入第二张平台图
        if section_key == '4　讨论':
            pic2_path, cap2, w2 = PLATFORM_FIG2
            if os.path.exists(pic2_path):
                # 找到刚插入的图注段落（target后面的第二个段落应该是图注）
                # 直接在target后面查找最后一个添加的图注段落
                # 简单方法：再次遍历找到"图8"的图注段落
                fig8_cap = None
                for p in doc.paragraphs:
                    if cap2.replace('图9', '图8') in p.text:  # 找到图8的图注
                        fig8_cap = p
                        break
                # 实际上 PLATFORM_FIG2 是图9，我们找到最后一个插入的图注
                # 更简单：在讨论的第一个内容段落后再插入图9
                # 由于图8已经插入了，我们在图8的图注后面插入图9
                # 找到图8的图注
                last_cap = None
                for p in doc.paragraphs:
                    if '图8' in p.text:
                        last_cap = p
                if last_cap:
                    insert_picture_after(doc, last_cap, pic2_path, w2, cap2)
                    inserted.append((section_key, cap2))
                    print(f'已插入 {cap2}')
    
    # Step 4: 在讨论部分添加平台使用价值
    for p in doc.paragraphs:
        if p.text.strip().startswith('本研究构建了针灸穴位-基因关联数据库'):
            p.add_run(
                f'数据库平台（{PLATFORM_URL}）提供了在线检索、共现网络可视化及文献溯源功能，'
                '可供研究者快速获取特定穴位或基因的关联信息，为后续实验设计提供文献依据。'
            )
            print('已在讨论中添加平台价值说明')
            break
    
    doc.save(OUTPUT_DOC)
    print(f'\n最终文档已保存: {OUTPUT_DOC}')
    print(f'共插入 {len(inserted)} 张图片')


if __name__ == '__main__':
    main()
