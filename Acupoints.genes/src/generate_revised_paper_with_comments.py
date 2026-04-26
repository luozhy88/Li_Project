#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成带审稿批注的修改版Word投稿论文
以编辑审稿人角度，在有问题的地方用红色括弧标注
"""

import sqlite3
import os
import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', font_size=Pt(12), bold=False, color=None):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_english_font(run, font_name='Arial', font_size=Pt(12), bold=False, color=None):
    """设置英文字体"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_comment_run(paragraph, text, color=RGBColor(0xFF, 0x00, 0x00), bold=True):
    """添加红色批注文字"""
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run


def add_normal_run(paragraph, text, font_name='宋体', font_size=Pt(12), bold=False):
    """添加正常文字"""
    run = paragraph.add_run(text)
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_paragraph_with_comments(doc, parts, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 first_line_indent=None, line_spacing=1.5,
                                 space_after=Pt(6), space_before=Pt(0)):
    """
    添加带有批注的段落
    parts: list of tuples, each tuple is (text, is_comment, comment_color)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

    for text, is_comment, color in parts:
        if is_comment:
            run = p.add_run(text)
            run.font.color.rgb = color if color else RGBColor(0xFF, 0x00, 0x00)
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_section_heading(doc, text, level=1):
    """添加章节标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    if level == 1:
        set_chinese_font(run, '黑体', Pt(14), bold=True)
    else:
        set_chinese_font(run, '黑体', Pt(12), bold=True)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


def collect_corrected_data():
    """收集修正后的数据"""
    conn = sqlite3.connect('output/acupoint_gene.db')
    cursor = conn.cursor()

    data = {}

    # 总体统计（修正后）
    cursor.execute('SELECT COUNT(*) FROM articles')
    data['total_articles'] = cursor.fetchone()[0]

    # 统计有效基因数（排除已知非基因实体）
    bad_genes = [
        'METABOLOMIC', 'POSTOPERATIVE', 'WEB', 'STROKE', 'SYNDROME', 'PROFILING',
        'IVF', 'SPINAL', 'HOMA', 'TRIAL', 'NOD', 'ECM', 'T2DM', 'BBB', 'VESSEL',
        'RESPONSE', 'LPS', 'CD4', 'CD8', 'VAS', 'SA', 'MC', 'KNEE', 'INSIGHTS',
        'TRANSCRIPTOMICS', 'INTEGRATING', 'CYTOSCAPE', 'A2', 'A1', 'VENN', 'HEALTH',
        'C3', 'AG', 'WOMAC', 'WL', 'TOLUIDINE', 'PITTSBURGH', 'PELVIC', 'OVERLAPPING',
        'LA', 'HAMD', 'GENECARDS', 'DISGENET', 'ANALGESIC', 'ALLEN', 'WENYANG',
        'VISUALIZATION', 'VBM', 'VA', 'SUBSTANCE', 'SMR', 'SC', 'RIF', 'RANDOMIZATION',
        'QTL', 'PROSTATITIS', 'POLYMORPHISMS', 'PLXNB2', 'PEARSON', 'PB', 'NIHSS',
        'NEUROIMAGING', 'MENDELIAN', 'MCS', 'MAXIMAL', 'LUTEAL', 'LTBP3', 'INTERVENTION',
        'INFLUENCES', 'IGA', 'GWAS', 'GG', 'FET', 'FD4', 'DATABASE', 'ANNOTATION', 'ADL',
        'ABDOMINAL', 'A10', 'CP', 'CPPS', 'CYTOSCAPE', 'DISGENET', 'GENECARDS',
        'IFNG', 'IL10', 'IL1B', 'IL2', 'IL4', 'IL6', 'INSIGHTS', 'MAXIMAL', 'MCC',
        'OVERLAPPING', 'PELVIC', 'PROSTATITIS', 'SUBSTANCE', 'VENN', 'VISUALIZATION',
        'ANNOTATION', 'DATABASE', 'DISGENET', 'GENECARDS'
    ]
    
    # 实际统计：先获取所有基因
    cursor.execute('SELECT gene_symbol FROM genes')
    all_genes = [row[0] for row in cursor.fetchall()]
    
    # 识别问题基因
    problematic = []
    for g in all_genes:
        g_upper = g.upper() if g else ''
        # 技术术语
        if g_upper in ['METABOLOMIC', 'PROFILING', 'CYTOSCAPE', 'DATABASE', 'DISGENET', 
                       'GENECARDS', 'VISUALIZATION', 'ANNOTATION', 'WEB', 'ECM']:
            problematic.append((g, '技术术语/工具名'))
        # 疾病/临床术语
        elif g_upper in ['STROKE', 'SYNDROME', 'T2DM', 'POSTOPERATIVE', 'SPINAL', 'NOD',
                         'PROSTATITIS', 'PELVIC', 'SUBSTANCE', 'KNEE', 'ABDOMINAL', 'HEALTH']:
            problematic.append((g, '疾病/临床术语'))
        # 研究类型/方法
        elif g_upper in ['TRIAL', 'IVF', 'HOMA', 'GWAS', 'TRANSCRIPTOMICS', 'NEUROIMAGING',
                         'POLYMORPHISMS', 'MENDELIAN', 'RANDOMIZATION', 'PROFILING']:
            problematic.append((g, '研究方法/类型术语'))
        # 量表/工具缩写
        elif g_upper in ['WOMAC', 'HAMD', 'NIHSS', 'PITTSBURGH', 'ADL', 'VAS', 'VBM']:
            problematic.append((g, '量表/评估工具缩写'))
        # 普通英文单词
        elif g_upper in ['RESPONSE', 'INFLUENCES', 'INSIGHTS', 'OVERLAPPING', 'MAXIMAL',
                         'INTERVENTION', 'LUTEAL', 'INTEGRATING', 'VENN', 'ALLEN', 'ANALGESIC']:
            problematic.append((g, '普通英文单词误识别'))
        # 单字母/双字母/无意义缩写
        elif g_upper in ['A1', 'A2', 'A10', 'C3', 'AG', 'GG', 'PB', 'SA', 'SC', 'VA', 
                         'WL', 'LA', 'MC', 'MCS', 'FD4', 'SP', 'CP', 'CPPS', 'MCC']:
            problematic.append((g, '非基因缩写/无意义字符'))
        # 基因家族/通路
        elif g_upper in ['AKT', 'PI3K', 'MAPK', 'WNT', 'JAK', 'ERK', 'MMP', 'BCL', 'P38', 'STAT']:
            problematic.append((g, '基因家族/通路名称，非单一基因'))
        # 非官方基因符号
        elif g_upper in ['LC3', 'BECLIN', 'BECLIN1']:
            problematic.append((g, '非官方基因符号'))
    
    data['problematic_genes'] = problematic
    data['valid_gene_count'] = len(all_genes) - len(set([p[0] for p in problematic]))
    data['total_genes'] = len(all_genes)

    cursor.execute('SELECT COUNT(*) FROM acupoints')
    data['total_acupoints'] = cursor.fetchone()[0]

    cursor.execute('SELECT COUNT(*) FROM diseases')
    data['total_diseases'] = cursor.fetchone()[0]

    # 年份分布
    cursor.execute('''
        SELECT year, COUNT(*) FROM articles 
        WHERE year IS NOT NULL AND year BETWEEN 2010 AND 2026
        GROUP BY year ORDER BY year
    ''')
    data['year_dist'] = cursor.fetchall()

    # 语言分布
    cursor.execute('SELECT language, COUNT(*) FROM articles GROUP BY language')
    data['lang_dist'] = cursor.fetchall()

    # 文献类型
    cursor.execute('''SELECT article_type, COUNT(*) FROM articles 
                      WHERE article_type LIKE '%Review%' OR article_type LIKE '%Letter%'
                      GROUP BY article_type''')
    data['review_articles'] = cursor.fetchall()

    # 近6年占比（2020-2026）
    cursor.execute('SELECT COUNT(*) FROM articles WHERE year >= 2020')
    data['recent_articles'] = cursor.fetchone()[0]

    # Top穴位（数据库已合并）
    cursor.execute('''
        SELECT a.name_cn, a.code, COUNT(DISTINCT aa.article_id) as cnt
        FROM acupoints a LEFT JOIN article_acupoint aa ON a.id = aa.acupoint_id
        GROUP BY a.id ORDER BY cnt DESC LIMIT 10
    ''')
    data['top_acupoints'] = cursor.fetchall()

    # Top基因（含问题基因）
    cursor.execute('''
        SELECT g.gene_symbol, COUNT(DISTINCT ag.article_id) as cnt
        FROM genes g JOIN article_gene ag ON g.id = ag.gene_id
        GROUP BY g.id ORDER BY cnt DESC LIMIT 20
    ''')
    data['top_genes'] = cursor.fetchall()

    # 疾病统计
    cursor.execute('''
        SELECT d.name, COUNT(DISTINCT ad.article_id) as cnt
        FROM diseases d JOIN article_disease ad ON d.id = ad.disease_id
        GROUP BY d.id ORDER BY cnt DESC LIMIT 15
    ''')
    data['top_diseases'] = cursor.fetchall()

    # 足三里相关基因（含问题）
    cursor.execute('''
        SELECT g.gene_symbol, COUNT(*) as cnt
        FROM acupoints a
        JOIN article_acupoint aa ON a.id = aa.acupoint_id
        JOIN article_gene ag ON aa.article_id = ag.article_id
        JOIN genes g ON ag.gene_id = g.id
        WHERE a.name_cn = '足三里'
        GROUP BY g.id ORDER BY cnt DESC LIMIT 10
    ''')
    data['st36_genes'] = cursor.fetchall()

    # 百会相关基因
    cursor.execute('''
        SELECT g.gene_symbol, COUNT(*) as cnt
        FROM acupoints a
        JOIN article_acupoint aa ON a.id = aa.acupoint_id
        JOIN article_gene ag ON aa.article_id = ag.article_id
        JOIN genes g ON ag.gene_id = g.id
        WHERE a.name_cn = '百会'
        GROUP BY g.id ORDER BY cnt DESC LIMIT 8
    ''')
    data['gv20_genes'] = cursor.fetchall()

    conn.close()
    return data


def generate_revised_paper():
    """生成带审稿批注的修改版Word论文"""
    data = collect_corrected_data()
    doc = Document()
    RED = RGBColor(0xFF, 0x00, 0x00)

    # 设置默认段落样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # ===== 标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 1.5
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run('针灸穴位-基因关联数据库的构建及数据挖掘分析')
    set_chinese_font(title_run, '宋体', Pt(22), bold=True)

    # 英文标题
    en_title_p = doc.add_paragraph()
    en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_title_p.paragraph_format.line_spacing = 1.5
    en_title_p.paragraph_format.space_after = Pt(12)
    en_title_run = en_title_p.add_run(
        'Construction and Data Mining Analysis of an Acupoint-Gene Association Database'
    )
    set_english_font(en_title_run, 'Arial', Pt(14), bold=True)

    # 作者
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(6)
    author_run = author_p.add_run('×××　×××　×××')
    set_chinese_font(author_run, '宋体', Pt(12))

    unit_p = doc.add_paragraph()
    unit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit_p.paragraph_format.space_after = Pt(12)
    unit_run = unit_p.add_run('（××大学　针灸推拿学院，北京　100000）')
    set_chinese_font(unit_run, '宋体', Pt(12))

    # ===== 摘要 =====
    add_section_heading(doc, '摘要', level=1)

    # 摘要正文 - 带多处批注
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)

    add_normal_run(p, '目的　')
    add_normal_run(p, '系统整合针灸穴位与基因表达关联的文献数据，构建针灸穴位-基因关联数据库，并基于该数据库开展数据挖掘分析，为揭示针灸作用的分子机制提供数据支撑。')
    add_comment_run(p, '【表述可行，但需注意后文数据可信度问题】')

    add_normal_run(p, '方法　')
    add_normal_run(p, '检索PubMed数据库中2010年1月至2026年4月发表的针灸穴位与基因表达相关文献，制定系统的检索策略与纳入排除标准，采用文本挖掘方法提取文献中的穴位、基因、疾病等实体信息。')
    add_comment_run(p, '【方法学缺陷1：纳入文献包含74篇中文文献（chi），与后文"语种为英文"的纳入标准自相矛盾，需明确是否实际纳入了中文文献，或修正纳入标准。】')
    add_comment_run(p, '【方法学缺陷2：数据库中实际包含Review（32篇）、Systematic Review（3篇）、Letter（1篇），与"排除综述、Meta分析、系统评价"的排除标准矛盾。】')

    add_normal_run(p, '依据《腧穴名称与定位》（GB/T 12346—2021）对穴位名称进行标准化处理，以NCBI Gene数据库为参照规范基因符号。')
    add_comment_run(p, '【数据清洗声明与实际严重不符：数据库中存在大量非基因实体被错误识别为基因符号，如METABOLOMIC、POSTOPERATIVE、STROKE、SYNDROME、WEB、ECM、T2DM、BBB等约100余个；同时PI3K、AKT、MAPK等基因家族/通路名称被当作单一基因统计；LC3、BECLIN等非官方符号未映射至标准HGNC符号。说明NCBI Gene校验步骤实际未执行或执行不彻底。】')

    add_normal_run(p, '采用SQLite构建关系型数据库，运用Python进行数据统计与关联分析。')
    add_comment_run(p, '【方法学深度不足：作为"数据库+数据挖掘"类论文，仅使用频次统计和共现频率远远不够。审稿人期望看到网络拓扑分析（度中心性、介数中心性）、聚类分析、富集分析（GO/KEGG）、甚至假设检验（如卡方检验判断共现是否高于随机期望）。目前的方法学深度不足以支撑数据库类SCI论文的数据挖掘部分。】')

    add_normal_run(p, '结果　')
    add_normal_run(p, f'数据库共收录文献{data["total_articles"]}篇，涉及穴位{data["total_acupoints"]}个、基因{data["total_genes"]}个、疾病或病症类型{data["total_diseases"]}种。')
    add_comment_run(p, f'【数据可信度存疑：声称基因数量为{data["total_genes"]}个，但经核查其中约{len(data["problematic_genes"])}个为非基因实体或问题实体，有效基因数约为{data["valid_gene_count"]}个。若发表前不完成全面基因校验，该数据将严重误导读者。】')

    add_normal_run(p, '文献发表量呈逐年增长趋势，2020年以来文献占比达')
    recent_pct = round(data['recent_articles'] / data['total_articles'] * 100, 1)
    add_normal_run(p, f'{recent_pct}%。')
    add_comment_run(p, '【年份数据异常：2025年收录138篇，2026年仅前4个月即收录111篇，折合全年约333篇，较2024年的19篇增长约17倍。这种跃升更可能的原因是PubMed中大量"Online ahead of print"文章被提前标注为2026年，或爬虫抓取了未正式分配卷期页码的预发表文献。作者未对此异常增长做任何技术解释，直接归因于"研究热度快速上升"，结论缺乏数据支撑。】')

    add_normal_run(p, '高频研究穴位依次为足三里（109篇）、百会（26篇）、四白（26篇）、三阴交（19篇）、关元（18篇）。核心关联基因包括肿瘤坏死因子（TNF，55篇）、白细胞介素-6（IL6，38篇）、蛋白激酶B（AKT，30篇）、磷脂酰肌醇3-激酶（PI3K，27篇）等。')
    add_comment_run(p, '【基因名称问题：AKT、PI3K均为基因家族/通路总称，并非单一基因符号。AKT应拆分为AKT1/AKT2/AKT3，PI3K应拆分为PIK3CA/PIK3CB/PIK3CD等，或单独标注为"基因家族/通路"。直接在结果中将其称为"核心关联基因"会误导读者认为这些是具体的单一基因。】')

    add_normal_run(p, '炎症、疼痛、应激、肿瘤及脑卒中为研究最为集中的疾病领域。')
    add_comment_run(p, '【分类学问题：炎症（inflammation/inflammatory重复计数，合计212篇）、疼痛（pain）、应激（stress）属于病理生理过程或症状，而非疾病（disease）。将其与肿瘤、脑卒中、关节炎等真正疾病并列统计，会导致疾病谱分类混乱，影响数据解读。建议拆分为"疾病"与"病理过程/症状"两个层级。】')

    add_normal_run(p, '数据挖掘结果显示，足三里-TNF为共现频率最高的穴位-基因关联对。')
    add_comment_run(p, '【逻辑跳跃：共现（co-occurrence）仅说明某篇文献中同时提及该穴位和该基因，不等于该穴位"调控"该基因。基因可能在文中作为背景机制、疾病标志物或对照提及。将共现频次直接解读为"关联强度"并推断"穴位刺激可能具有基因调控的相对特异性"，属于方法学上的过度推断。此外，许多文献研究的是穴位配伍（如百会+太冲、足三里+三阴交），共现分析无法将配伍效应归因于单个穴位，这会严重高估单穴的特异性。】')

    add_normal_run(p, '结论　')
    add_normal_run(p, '针灸穴位-基因关联数据库初步实现了该领域文献数据的系统整合与结构化存储。穴位对基因的调控具有多靶点、多通路特征，炎症与免疫相关基因为当前研究热点。该数据库可为后续针灸分子机制研究及穴位特异性分析提供参考。')
    add_comment_run(p, '【结论中"穴位对基因的调控"表述过于肯定：基于共现分析只能得出"穴位与基因在研究中被共同关注"，不能直接推断"调控"关系。建议在结论中弱化因果推断，改用"关联"或"共现关注度"等表述。】')

    # 关键词
    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.first_line_indent = Cm(0.74)
    kw_p.paragraph_format.line_spacing = 1.5
    add_normal_run(kw_p, '关键词　')
    add_normal_run(kw_p, '针灸；穴位；基因表达；数据库；数据挖掘；分子机制')

    # ===== Abstract =====
    add_section_heading(doc, 'Abstract', level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_english_run(p, 'Objective　', bold=True)
    add_english_run(p, 'To systematically integrate literature data on acupoint-gene expression associations, construct an acupoint-gene association database, and perform data mining analysis to provide support for revealing the molecular mechanisms of acupuncture. ')
    add_comment_run(p, ' [Note: The gene entity recognition errors and inclusion criteria contradictions mentioned in the Chinese abstract equally apply here.] ')

    add_english_run(p, 'Methods　', bold=True)
    add_english_run(p, 'PubMed was searched for literature on acupoint-gene expression published from January 2010 to April 2026. Systematic retrieval strategies and inclusion/exclusion criteria were formulated. Text mining was employed to extract acupoint, gene, and disease entities. Acupoint names were standardized according to GB/T 12346—2021, and gene symbols were normalized against the NCBI Gene database. ')
    add_comment_run(p, ' [CRITICAL: Gene normalization claim is not supported by data. Numerous non-gene entities (METABOLOMIC, POSTOPERATIVE, STROKE, SYNDROME, WEB, T2DM, etc.) and gene family names (AKT, PI3K, MAPK) remain in the database as single-gene entries. NCBI Gene API validation was apparently not performed.] ')

    add_english_run(p, 'A relational database was built using SQLite, and Python was used for statistical analysis and association mining. ')
    add_comment_run(p, ' [Methodological limitation: Only frequency counting and co-occurrence analysis were performed. Network topology analysis (degree/betweenness centrality), enrichment analysis (GO/KEGG), and statistical hypothesis testing are expected for a database-mining paper.] ')

    add_english_run(p, 'Results　', bold=True)
    add_english_run(p, f'The database included {data["total_articles"]} articles, covering {data["total_acupoints"]} acupoints, {data["total_genes"]} genes, and {data["total_diseases"]} disease categories. ')
    add_comment_run(p, f' [Data quality issue: Of the claimed {data["total_genes"]} genes, approximately {len(data["problematic_genes"])} are non-gene entities or problematic entries, leaving only ~{data["valid_gene_count"]} valid gene symbols. This must be corrected before publication.] ')

    add_english_run(p, f'Publication output showed an increasing trend year by year, with {recent_pct}% of articles published since 2020. ')
    add_comment_run(p, ' [Data anomaly unexplained: 138 articles in 2025 and 111 in the first 4 months of 2026 (annualized ~333) vs. only 19 in 2024. This 17-fold jump likely reflects PubMed "ahead of print" dating artifacts rather than genuine research growth. Technical explanation required.] ')

    add_english_run(p, 'The most frequently studied acupoints were Zusanli (ST36, 109 articles), Baihui (GV20, 26), Sibai (ST2, 26), Sanyinjiao (SP6, 19), and Guanyuan (CV4, 18). Core associated genes included TNF (55 articles), IL6 (38), AKT (30), and PI3K (27). ')
    add_comment_run(p, ' [Terminology error: AKT and PI3K are gene families/pathways, not single gene symbols. They should be disaggregated into specific isoforms (AKT1, AKT2, AKT3; PIK3CA, PIK3CB, etc.) or labeled as "gene family/pathway." Presenting them as "core associated genes" is misleading.] ')

    add_english_run(p, 'Inflammation, pain, stress, tumor, and stroke were the most intensively studied diseases. ')
    add_comment_run(p, ' [Classification error: inflammation, pain, and stress are pathophysiological processes or symptoms, not diseases. Combining them with true diseases (tumor, stroke) creates taxonomic confusion. A two-tier system (disease vs. pathological process/symptom) is recommended.] ')

    add_english_run(p, 'Data mining revealed Zusanli-TNF as the most frequent acupoint-gene association pair. ')
    add_comment_run(p, ' [Overinterpretation: Co-occurrence in abstracts does not establish regulatory causation. The gene may be mentioned as background, biomarker, or control. Moreover, many studies use acupoint combinations, making single-acupoint attribution impossible without controlled experiments.] ')

    add_english_run(p, 'Conclusion　', bold=True)
    add_english_run(p, 'The acupoint-gene association database achieves preliminary systematic integration and structured storage of literature data in this field. Acupoint regulation of gene expression exhibits multi-target and multi-pathway characteristics, with inflammation- and immune-related genes being the current research hotspot. This database may serve as a reference for subsequent studies on acupuncture molecular mechanisms and acupoint specificity.')
    add_comment_run(p, ' [The phrase "acupoint regulation of gene expression" overstates the evidence. Co-occurrence analysis only demonstrates research attention, not regulatory causation. The conclusion should be softened to "acupoint-gene associations" or "co-occurrence patterns."] ')

    # Keywords
    kw_p2 = doc.add_paragraph()
    kw_p2.paragraph_format.first_line_indent = Cm(0.74)
    kw_p2.paragraph_format.line_spacing = 1.5
    add_english_run(kw_p2, 'Keywords　', bold=True)
    add_english_run(kw_p2, 'acupuncture; acupoint; gene expression; database; data mining; molecular mechanism')

    # ===== 1 引言 =====
    add_section_heading(doc, '1　引言', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '针灸作为中医学的重要组成部分，其临床疗效已得到广泛验证。近年来，随着基因组学、转录组学及蛋白质组学技术的快速发展，越来越多的研究从分子水平探讨针灸的作用机制，证实穴位刺激可特异性调控基因表达，进而影响下游信号通路与生物学功能[1-3]。然而，现有的针灸-基因研究数据分散于数千篇文献之中，缺乏统一的数据整合平台，研究者难以快速、全面地获取特定穴位所调控的基因信息。')
    add_comment_run(p, '【引言逻辑可接受，但建议补充说明：本文采用文本挖掘方法，提取的是文献中穴位与基因的共现信息，而非经实验验证的调控关系，以避免读者对数据库性质产生误解。】')

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p2, '在数据库资源方面，TCMGeneDIT、TCMSP、HERB等数据库主要聚焦中药成分与靶点的关联，未涉及针灸穴位层面的分子数据[4,5]；美国国立卫生研究院（NIH）资助的TARA项目虽关注穴位解剖与生理特征，但尚未纳入基因表达层面的信息[6]。因此，构建一个专门面向针灸穴位-基因关联的数据库，对于整合现有研究成果、挖掘穴位作用的共性规律与特异性特征，具有重要的科学价值。')
    add_comment_run(p2, '【建议核实：是否已存在其他针灸-基因相关数据库或知识图谱。若审稿人发现已有类似资源而本文声称"填补空白"，可能直接拒稿。建议在引言中客观描述现有资源的特点与不足，而非强调"空白"。】')

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p3, '本研究基于PubMed文献资源，采用文本挖掘与数据标准化技术，构建了针灸穴位-基因关联数据库，并对收录数据进行统计分析与关联挖掘，以期为针灸分子机制研究提供数据基础。')

    # ===== 2 资料与方法 =====
    add_section_heading(doc, '2　资料与方法', level=1)

    # 2.1
    add_section_heading(doc, '2.1　数据来源与检索策略', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '文献数据来源于PubMed数据库（https://pubmed.ncbi.nlm.nih.gov），检索时限为2010年1月至2026年4月。检索策略采用主题词与自由词结合的方式，核心检索式为：(acupuncture[Title/Abstract] OR electroacupuncture[Title/Abstract] OR moxibustion[Title/Abstract]) AND (gene[Title/Abstract] OR expression[Title/Abstract] OR microarray[Title/Abstract] OR transcriptome[Title/Abstract] OR proteomics[Title/Abstract])，并根据具体穴位名称（如Zusanli、ST36、Taichong、LR3等）进行扩展检索。同时补充检索了抑郁症、高血压、帕金森病、炎症、疼痛等疾病相关的针灸-基因文献，以确保数据覆盖的全面性。')
    add_comment_run(p, '【检索策略描述较为完整，但需补充说明：实际检索是否覆盖了中文数据库（CNKI、万方）？数据库中74篇中文文献（chi）的来源是PubMed中标注为中文的文献，还是额外检索了中文数据库？若仅检索PubMed，则中文文献应为英文发表但被PubMed标注语言为中文者，需在方法中说明。】')

    # 2.2
    add_section_heading(doc, '2.2　纳入与排除标准', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '纳入标准：（1）研究类型为实验研究、临床研究或网络药理学研究；')
    add_comment_run(p, '【方法学重大缺陷：将"网络药理学研究"（in silico预测）与实验/临床研究（体外/体内验证）纳入同一数据库并等同处理，会导致预测靶点与实验验证基因混为一谈，严重稀释数据库的证据等级。建议修改为：纳入标准：（1）研究类型为实验研究、临床研究或网络药理学研究，但须在数据库中设置"证据等级"字段（实验验证=Ⅰ级，临床研究=Ⅱ级，网络药理学预测=Ⅲ级），分层统计与展示，避免将预测性靶点直接等同于实验验证的差异表达基因。】')
    add_normal_run(p, '（2）明确报告了穴位名称及差异表达基因或核心靶点基因；')
    add_comment_run(p, '【建议补充：网络药理学文献中的"核心靶点"通常基于网络拓扑计算（如度值、中介中心性）而非实验验证，与差异表达基因的生物学意义不同，需在提取时明确标注数据来源类型。】')
    add_normal_run(p, '（3）语种为英文。')
    add_comment_run(p, '【与实际情况矛盾：数据库中74篇文献语言标注为中文（chi），占总量21.5%。若确实仅纳入英文文献，则需核查这些中文标注文献的实际语言；若实际纳入了中文文献，则纳入标准应修改为"中英文"。此外，3.1节称文献来源期刊以《针刺研究》《中国针灸》为主，这些期刊虽有英文版，但需明确说明纳入的是其英文发表文章，否则存在逻辑不一致。】')

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p2, '排除标准：（1）综述、Meta分析、系统评价及述评类文献；（2）仅涉及中药成分而未涉及穴位刺激的文献；（3）摘要信息不完整，无法提取穴位或基因数据的文献。')
    add_comment_run(p2, '【排除标准执行不彻底：数据库中实际包含Review类文献32篇、Systematic Review 3篇、Letter 1篇，与"排除综述"的标准直接矛盾。此外，排除标准不完整：未明确排除"动物实验未注明物种""细胞实验未注明穴位模拟方式"等低可外推性研究；亦未说明是否纳入"Ahead of print"文献。】')
    add_comment_run(p2, '【建议修改为：排除标准增加：（4）动物实验未明确报告物种及造模方法者；（5）细胞实验未明确穴位刺激模拟方式（如电针条件培养基）者；（6）Ahead of print且无完整摘要者；（7）Review、Systematic Review、Letter、Commentary类文献（已纳入者须剔除）。】')

    # 2.3
    add_section_heading(doc, '2.3　数据提取与标准化', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '采用Python编程语言结合正则表达式与词典匹配方法，从文献标题及摘要中自动提取穴位名称、基因符号、疾病名称等实体信息。穴位名称参照《腧穴名称与定位》（GB/T 12346—2021）进行规范化处理，将英文代码（如ST36、LR3）、拼音（如Zusanli、Taichong）及别名统一转换为标准中文名称，并标注经络归属与国际标准代码。')
    add_comment_run(p, '【穴位合并声明与实际基本相符：数据库层面（acupoints表）已将ST36/Zusanli/Sanli合并为"足三里"（109篇），Baihui/GV20合并为"百会"（26篇），但原始提取的summary.json中仍显示拆分计数（ST36:96、Zusanli:86、Sanli:86）。若论文中引用的统计数据来源于未合并前的summary.json，则会导致共现分析的分母和分子均失真。需确认论文所有统计数据均基于合并后的数据库表，而非旧版summary.json。】')

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p2, '基因符号以NCBI Gene数据库为参照进行标准化，去除重复项及非基因实体（如技术术语、疾病缩写等）。')
    add_comment_run(p2, '【原始数据清洗存在严重问题，此声明与实际情况严重不符。实际提取的基因列表中仍包含大量非基因实体：METABOLOMIC（代谢组学，技术术语）、POSTOPERATIVE（术后，临床术语）、WEB（网页）、STROKE（脑卒中，疾病名）、SYNDROME（综合征）、PROFILING（分析技术）、IVF（体外受精）、SPINAL（脊柱）、HOMA（稳态模型评估）、TRIAL（临床试验）、NOD（结节）、ECM（细胞外基质）等。此外，PI3K、AKT、MAPK均为基因家族/通路总称，并非单一基因符号；LC3官方符号为MAP1LC3A/B/C，BECLIN官方符号为BECN1。】')
    add_comment_run(p2, '【建议修改为：基因符号以NCBI Gene/HGNC数据库为参照进行标准化，采用双向校验机制：首先通过字典匹配去除已知非基因实体（如技术术语、疾病缩写、实验操作词等），其次将家族性名称（如PI3K、AKT、MAPK）拆分为具体基因亚型（PIK3CA、PIK3CB、AKT1、AKT2等）或单独标注为"基因家族/通路"，再次对剩余候选基因进行NCBI Gene API批量校验，仅保留官方Symbol，确保所有实体均为有效基因。目前数据库中约15%的基因-文献关联涉及非基因实体，必须在发表前彻底清理。】')

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p3, '疾病名称采用MeSH主题词进行归类整理。')
    add_comment_run(p3, '【疾病归类存在问题：inflammatory与inflammation被当作两个独立实体分别计数（127篇和85篇），实为同一概念的不同词形；pain、stress被归类为"疾病"，实为症状/病理过程。建议建立疾病-症状分层体系，并对同义词进行合并。】')

    # 2.4
    add_section_heading(doc, '2.4　数据库构建与数据分析', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '采用SQLite构建轻量级关系型数据库，设计articles（文献表）、acupoints（穴位表）、genes（基因表）、diseases（疾病表）及article_acupoint、article_gene、article_disease（关联表），实现文献、穴位、基因、疾病四维数据的结构化存储。运用Python pandas与collections模块进行频次统计与关联分析，计算穴位-基因共现频率，识别高频关联对。')
    add_comment_run(p, '【作为"数据库+数据挖掘"类论文，仅使用频次统计和共现频率远远不够。审稿人通常会期望看到：（1）网络拓扑分析（度中心性、介数中心性）识别网络枢纽节点；（2）KEGG/GO富集分析揭示功能通路特征；（3）超几何分布或卡方检验评估穴位-基因共现是否显著高于随机期望；（4）聚类分析或模块识别。目前的方法学深度不足以支撑一篇数据库类SCI论文的数据挖掘部分，建议补充上述分析。】')

    # ===== 3 结果 =====
    add_section_heading(doc, '3　结果', level=1)

    # 3.1
    add_section_heading(doc, '3.1　文献收录概况', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, f'经检索与筛选，数据库共收录符合纳入标准的文献{data["total_articles"]}篇，涉及穴位{data["total_acupoints"]}个、基因{data["total_genes"]}个、疾病或病症类型{data["total_diseases"]}种。')
    add_comment_run(p, f'【数据可信度警告：{data["total_genes"]}个基因中包含约{len(data["problematic_genes"])}个非基因实体或问题实体（如METABOLOMIC、POSTOPERATIVE、STROKE、SYNDROME、WEB、T2DM、BBB等），有效基因数约为{data["valid_gene_count"]}个。若发表前不修正，该基础数据将严重误导后续分析结论。】')
    add_comment_run(p, '【语言矛盾：74篇文献语言为中文（占21.5%），与"纳入标准：语种为英文"矛盾。若实际纳入了中文文献，应修改纳入标准并说明；若这些为中文期刊的英文版文章，应在方法中明确说明。】')

    add_normal_run(p, '文献发表时间跨度为2010年至2026年，总体呈增长态势。2020年以前年均收录量不足10篇，2024年增至19篇，2025年与2026年（截至4月）分别收录138篇和111篇，近6年文献量占总量的')
    recent_pct_str = f'{recent_pct}%'
    add_normal_run(p, recent_pct_str)
    add_normal_run(p, '，表明该领域研究热度正在快速上升。')
    add_comment_run(p, '【数据异常未解释：2026年仅前4个月即收录111篇，折合全年约333篇，较2024年的19篇增长约17倍。这种跃升更可能的原因是：（1）PubMed中大量"Online ahead of print"或"Epub ahead of print"文章被提前标注为2026年；（2）爬虫抓取了未正式分配卷期页码的预发表文献；（3）部分文章的PubMed日期标注为正式在线发表日期而非纸质出版日期。作者未对此异常增长做任何技术解释，直接归因于"研究热度快速上升"，结论缺乏数据支撑。】')
    add_comment_run(p, '【建议修改为：2024年增至19篇，2025年与2026年（截至4月）分别收录138篇和111篇。需注意的是，2026年数据仅覆盖4个月，且PubMed中部分文献为"Online ahead of print"状态，其实际正式出版年份可能滞后。因此，2025—2026年的激增可能部分反映了在线优先出版模式的普及，而非纯粹的文献增量。建议在数据库中增加"出版状态"字段（正式出版/Epub ahead of print），并在趋势分析中以"正式出版年份"为主进行绘图，以提高时间序列的可靠性。】')

    add_normal_run(p, '文献来源期刊以《针刺研究》（Zhen ci yan jiu）、《中国针灸》（Zhongguo zhen jiu）、Chinese Medicine、Evidence-based Complementary and Alternative Medicine等为主。')
    add_comment_run(p, '【未提供具体的期刊分布数据（如各期刊占比、影响因子分布），且未讨论是否存在期刊偏倚（如某些期刊更倾向于发表阳性结果）。建议补充期刊分布表，并讨论潜在发表偏倚对数据库代表性的影响。】')

    # 3.2
    add_section_heading(doc, '3.2　穴位分布特征', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '按文献关联频次排序，研究最为集中的前10个穴位依次为：足三里（ST36，109篇）、百会（GV20，26篇）、四白（ST2，26篇）、三阴交（SP6，19篇）、关元（CV4，18篇）、内关（PC6，18篇）、攒竹（BL2，13篇）、上关（GB3，12篇）、长强（GV1，11篇）、天枢（ST25，11篇）。足三里以绝对优势位居首位，其文献量占总量的31.7%，这与足三里作为强身保健要穴的经典定位相符，也反映了其在免疫调节、抗炎、神经保护等分子机制研究中的核心地位。百会、四白、三阴交等头面与下肢穴位亦受到较多关注，提示研究者对脑病、妇科病及面部疾病相关穴位的基因调控机制具有浓厚兴趣。从经络分布看，足阳明胃经穴位占比最高（足三里、四白、天枢、地仓），其次为督脉（百会、长强、大椎、水沟）与足太阳膀胱经（攒竹、肾俞、大杼），手厥阴心包经（内关）、任脉（关元、气海）等亦有涉及。')
    add_comment_run(p, '【穴位统计基本可信（数据库层面已合并同义词），但需注意：原始summary.json中同一穴位的不同名称（ST36/Zusanli/Sanli、Baihui/GV20等）未被合并，若论文中任何统计引用了summary.json而非数据库表，会导致计数失真。建议全文统一使用数据库合并后的统计数据，并删除或更新旧版summary.json。】')

    # 3.3
    add_section_heading(doc, '3.3　高频关联基因', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '数据库共提取基因实体2083个，去重后保留差异表达基因或核心靶点基因。高频关联基因以炎症与免疫调控分子为主：肿瘤坏死因子（TNF，55篇）、白细胞介素-6（IL6，38篇）位居前两位，提示炎症调控是针灸作用机制研究的核心切入点。信号转导分子如蛋白激酶B（AKT，30篇）、磷脂酰肌醇3-激酶（PI3K，27篇）、丝裂原活化蛋白激酶（MAPK，17篇）、哺乳动物雷帕霉素靶蛋白（MTOR，14篇）及核苷酸结合寡聚化结构域样受体蛋白3（NLRP3，14篇）亦为研究热点，表明细胞信号通路的激活或抑制在针灸效应中发挥关键介导作用。此外，自噬相关基因（LC3、BECLIN）、转录因子（FOS、STAT3）、凋亡调控因子（BCL2家族）等亦出现较高频次，反映了针灸作用的多靶点、多层次特征。')
    add_comment_run(p, '【严重数据质量问题：（1）声称"提取基因实体2083个"，但其中约100余个为非基因实体（METABOLOMIC、POSTOPERATIVE、STROKE、SYNDROME、WEB、T2DM、BBB等），有效基因数远低于此。（2）AKT、PI3K、MAPK均为基因家族/通路总称，并非单一基因符号，不能称为"差异表达基因"或"核心靶点基因"。（3）LC3官方符号为MAP1LC3A/B/C，BECLIN官方符号为BECN1，使用非官方符号会降低数据库的专业性和可检索性。（4）"BCL2家族"表述模糊，数据库中同时存在BCL2和BCL两个独立条目，BCL并非有效基因符号。】')
    add_comment_run(p, '【修改建议：在结果中增加"数据质量说明"小节，明确指出当前版本存在基因实体识别误差，并给出已识别的问题实体清单及修正计划。对于AKT、PI3K、MAPK等家族名称，应修改为"AKT家族（含AKT1/AKT2/AKT3）"等表述，或将其从基因统计中剔除，单独列为"高频信号通路"。】')

    # 3.4
    add_section_heading(doc, '3.4　疾病领域分布', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '疾病实体统计结果显示，炎症（127篇）与疼痛（69篇）为研究最集中的两大领域，其次为应激（59篇）、肿瘤（38篇）、脑卒中（31篇）、关节炎（30篇）及戒断综合征（25篇）。抑郁症（20篇）、焦虑症（13篇）、骨关节炎（13篇）、肠易激综合征（12篇）及阿尔茨海默病（11篇）等神经精神及消化系统疾病亦占一定比例。上述疾病分布与针灸临床优势病种高度吻合，说明分子机制研究正在从实验向临床转化，具有明确的应用导向。')
    add_comment_run(p, '【分类学问题：炎症（inflammation/inflammatory在数据库中被拆分为两个实体，合计212篇，此处仅引用127篇，存在数据不一致）、疼痛（pain）、应激（stress）属于病理生理过程或症状，而非疾病（disease）。将其与肿瘤、脑卒中、关节炎等真正疾病并列统计，会导致疾病谱分类混乱，影响数据解读。建议将结果拆分为"疾病类"和"病理过程/症状类"两个层级分别统计。】')
    add_comment_run(p, '【此外，"炎症127篇"与数据库中inflammatory 127篇+inflammation 85篇=212篇存在矛盾，需核查原文引用的是哪个字段，并解释为何与数据库总计不一致。】')

    # 3.5
    add_section_heading(doc, '3.5　穴位-基因关联挖掘', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '通过共现分析识别穴位与基因的关联强度。足三里-TNF为共现频次最高的穴位-基因对（22次），其次是足三里-IL6（16次）、足三里-AKT（10次）、足三里-PI3K（10次）。足三里相关基因呈现明显的炎症与免疫调控特征，涉及核因子-κB（NF-κB）、Toll样受体4（TLR4）等关键通路。百会穴相关的高频基因包括TNF、BDNF、GFAP等，提示其在神经保护及神经炎症调控中的重要作用。三阴交的高频关联基因以TNF、TG（甲状腺球蛋白）为主，与其在妇科疾病及内分泌调节中的应用背景一致。不同穴位的高频关联基因存在明显差异，初步提示穴位刺激可能具有基因调控的相对特异性。')
    add_comment_run(p, '【方法学过度推断（最严重问题之一）：共现（co-occurrence）仅说明某篇文献中同时提及该穴位和该基因，不等于该穴位"调控"该基因。基因可能在文中作为背景机制、疾病标志物或对照提及。将共现频次直接解读为"关联强度"并推断"穴位刺激可能具有基因调控的相对特异性"，属于方法学上的过度推断。】')
    add_comment_run(p, '【配伍效应无法拆分：许多文献研究的是穴位配伍（如百会+太冲、足三里+三阴交），共现分析无法将配伍效应归因于单个穴位，这会严重高估单穴的特异性。例如，若某文献同时研究"百会+太冲"并提及BDNF，共现分析会将BDNF同时归因于百会和太冲，但原文可能并未区分两穴各自对BDNF的独立贡献。】')
    add_comment_run(p, '【建议修改：（1）明确声明"共现频次反映的是研究关注度，而非直接的调控因果关系"；（2）增加验证步骤：仅统计基因出现在"结果/结论"段落（而非背景/方法）的共现；对共现频次最高的前50对进行人工抽样复核；（3）将"穴位特异性"结论改为"研究关注度的穴位差异"，并说明由于配伍效应和共现分析的局限性，"穴位特异性"的结论需在严格控制配伍因素及采用直接对比实验后方可确立。】')

    # ===== 4 讨论 =====
    add_section_heading(doc, '4　讨论', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '本研究构建了')
    add_normal_run(p, '首个')
    add_comment_run(p, '【"首个"表述风险极高。若审稿人检索到已有类似数据库（即使功能不完全相同），可能直接拒稿。建议弱化此表述，突出本数据库的特色（如文本挖掘方法、特定时间跨度、穴位标准化策略等），而非绝对意义上的"首个"。】')
    add_normal_run(p, '面向针灸穴位-基因关联的文献数据库，实现了344篇文献、41个穴位、2083个基因及52种疾病的系统整合。数据库的建成填补了该领域缺乏专用数据平台的空白，为后续针灸分子机制的系统研究提供了基础数据支撑。')
    add_comment_run(p, '【"填补了该领域缺乏专用数据平台的空白"与"首个"问题类似，表述过于绝对。建议改为"为整合现有分散的针灸-基因研究数据提供了一种可扩展的文献整合方案"等更谦逊的表述。】')

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p2, '从数据分析结果来看，足三里在基因调控研究中占据核心地位，其高频关联基因TNF、IL6、AKT、PI3K等均属于炎症与信号转导领域的关键分子。这一发现与足三里"培补后天、扶正祛邪"的传统功效认识相呼应。现代研究表明，针刺足三里可通过调控TNF-α、IL-6等促炎因子的表达，发挥抗炎与免疫调节作用[7,8]；同时，AKT/PI3K信号通路的激活与细胞存活、代谢调节密切相关[9]，可能是足三里发挥多系统调节作用的分子基础。')
    add_comment_run(p2, '【讨论中的"基因调控""关联基因"等表述延续了结果部分的过度推断问题。建议将"基因调控研究"改为"共现关注度研究"，将"高频关联基因"改为"高频共现基因"，以准确反映分析方法的本质。此外，AKT/PI3K作为通路名称，在讨论中作为信号通路来阐述是合适的，但在结果部分将其列为"基因"则是不当的，需注意前后表述的一致性。】')

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p3, '百会穴作为督脉要穴，其高频关联基因以神经相关分子为主。BDNF是维持神经元存活与突触可塑性的关键神经营养因子，针刺百会可通过上调BDNF表达改善脑缺血再灌注损伤及抑郁样行为[10,11]。GFAP作为星形胶质细胞活化的标志物，其表达下调提示针刺百会可能通过抑制星形胶质细胞过度活化而减轻神经炎症[12]。这些发现为百会穴主治神志病、脑病的传统应用提供了分子层面的佐证。')
    add_comment_run(p3, '【讨论逻辑基本合理，但需注意：本文数据库中百会与BDNF、GFAP的共现，仅说明这些文献同时提及百会和这些分子，不能证明数据库本身"发现"了这些关联。讨论中引用具体文献[10-12]来支持关联是恰当的，但这属于对已有文献的综述，而非对数据库挖掘结果的独特发现。建议明确区分"数据库共现统计"与"文献实验验证"，避免读者误以为共现分析本身提供了新的机制证据。】')

    p4 = doc.add_paragraph()
    p4.paragraph_format.line_spacing = 1.5
    p4.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p4, '值得注意的是，不同穴位的高频关联基因存在差异，初步提示穴位对基因的调控可能具有一定的相对特异性。足三里偏向炎症与免疫，百会偏向神经保护，三阴交则涉及内分泌与代谢调节。这种差异是否与穴位的经络归属、解剖位置及主治功能相关，尚需在更大样本量和更严格的实验条件下进一步验证。未来可引入网络药理学方法，构建穴位-基因-通路多维网络，深入挖掘穴位特异性的分子标识。')
    add_comment_run(p4, '【这是全文最关键的过度推断之一。基于共现分析得出"穴位对基因的调控可能具有相对特异性"是不成立的，原因如下：（1）共现≠调控；（2）配伍效应无法拆分；（3）文献选择偏倚（某些穴位被研究得更多，自然会与更多基因共现）。建议修改为："不同穴位的高频关联基因存在明显差异，初步提示研究者在不同穴位方向上关注的分子靶点存在差异；但由于共现分析无法区分单穴效应与配伍效应，且部分文献为配伍取穴，故穴位特异性的结论需在严格控制配伍因素及采用直接对比实验（穴位vs非穴位、经脉vs非经脉）后方可确立。未来可引入网络药理学方法构建穴位-基因-通路多维网络，但网络分析结果仍需实验验证。"]')

    p5 = doc.add_paragraph()
    p5.paragraph_format.line_spacing = 1.5
    p5.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p5, '本研究存在以下局限。第一，数据提取主要依赖词典匹配与正则表达式，对于复杂句式或隐含信息的识别能力有限，部分文献可能存在穴位或基因提取遗漏。第二，纳入文献以英文为主，未能充分覆盖中文数据库（如CNKI、万方）中的相关研究，可能影响数据的全面性。第三，文献中基因表达的上调/下调方向、倍数变化及统计学显著性等定量信息未能完全提取，数据库目前以定性关联为主。第四，部分高频实体（如COA、TG）是否为准确的基因符号尚需人工复核。')
    add_comment_run(p5, '【局限部分对数据质量问题的认识严重不足。作者仅提及"COA、TG尚需人工复核"，但实际上问题远比这严重：约100余个非基因实体被错误识别为基因（METABOLOMIC、POSTOPERATIVE、STROKE、SYNDROME等），大量基因家族/通路被当作单一基因（AKT、PI3K、MAPK），非官方符号未映射（LC3→MAP1LC3A/B/C，BECLIN→BECN1）。这些系统性错误对共现分析结果的可信度构成实质性影响，必须在局限部分如实披露。】')
    add_comment_run(p5, '【此外，局限部分未提及以下关键问题：（1）纳入标准与实际情况的语言矛盾（74篇中文文献）；（2）排除标准执行不彻底（数据库中仍含Review、Letter等）；（3）2025—2026年数据异常激增未做技术解释；（4）网络药理学预测数据与实验数据未分层，导致证据等级混杂。这些同样是重要的方法学局限，应在修订稿中补充。】')

    add_normal_run(p5, '后续研究将引入自然语言处理模型提升实体识别精度，并补充中文文献与定量数据，逐步完善数据库内容。')

    # ===== 5 结论 =====
    add_section_heading(doc, '5　结论', level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_normal_run(p, '本研究构建了针灸穴位-基因关联数据库，初步实现了该领域文献数据的系统整合。数据挖掘结果表明，炎症与免疫相关基因为当前针灸分子机制研究的热点，足三里、百会等穴位在基因调控网络中处于核心节点，不同穴位的高频关联基因存在差异，提示穴位作用可能具有分子层面的相对特异性。该数据库可为针灸机制研究、穴位选优及临床方案设计提供数据参考。')
    add_comment_run(p, '【结论中"在基因调控网络中处于核心节点"和"穴位作用可能具有分子层面的相对特异性"均基于共现分析的过度推断。建议修改为："数据挖掘结果表明，炎症与免疫相关分子为当前针灸分子机制研究的热点关注领域，足三里、百会等穴位在与基因共现的网络中处于高关注度节点。不同穴位的高频共现基因存在差异，提示研究关注方向存在穴位差异性，但受限于共现分析方法学，尚不能推断穴位对基因的特异性调控作用。该数据库可为针灸机制研究的文献调研、穴位选优的方向性参考及临床方案设计提供数据基础。"]')

    # ===== 参考文献 =====
    add_section_heading(doc, '参考文献', level=1)
    refs = [
        '[1] 张颖, 李晶, 王瑜, 等. 针刺对抑郁症模型大鼠海马区基因表达谱的影响[J]. 针刺研究, 2024, 49(3): 245-252.',
        '[2] Ma SM, Zhang L, Chen Y, et al. Transcriptome analysis of gene expression in spontaneously hypertensive rats following acupuncture at Taichong (LR3)[J]. J Tradit Chin Med, 2019, 39(4): 562-570.',
        '[3] Choi EM, Jiang F, Longo LD. Acupuncture at Taichong (LR3) modulates gene expression in the substantia nigra of Parkinson disease mice[J]. Neurochem Res, 2011, 36(11): 2129-2137.',
        '[4] Chen CY. TCM Database@Taiwan: the world\'s largest traditional Chinese medicine database for drug screening in silico[J]. PLoS One, 2011, 6(1): e15939.',
        '[5] Ru J, Li P, Wang J, et al. TCMSP: a database of systems pharmacology for drug discovery from herbal medicines[J]. J Cheminform, 2014, 6: 13.',
        '[6] Schnyer R, Abrams M. The TARA project: a resource for acupuncture research[J]. J Altern Complement Med, 2020, 26(5): 385-390.',
        '[7] 刘志诚, 孙志洁, 李玫, 等. 针刺对肥胖大鼠下丘脑IL-6、TNF-α及瘦素受体表达的影响[J]. 中国针灸, 2018, 38(6): 611-616.',
        '[8] Torres-Rosas R, Yehia G, Peña G, et al. Dopamine mediates vagal modulation of the immune system by electroacupuncture[J]. Nat Med, 2014, 20(3): 291-297.',
        '[9] 刘静, 赵吉平, 李彬, 等. 电针足三里对功能性消化不良大鼠胃窦PI3K/AKT/mTOR信号通路的影响[J]. 针刺研究, 2023, 48(2): 133-140.',
        '[10] 王顺, 张轶丹, 姜丽芳, 等. 针刺百会、大椎对脑缺血再灌注损伤大鼠BDNF表达的影响[J]. 中国中西医结合杂志, 2015, 35(8): 973-977.',
        '[11] 尹磊淼, 杨永清, 王宇, 等. 针刺抗抑郁的分子机制研究进展[J]. 中国中西医结合杂志, 2022, 42(5): 631-636.',
        '[12] 李明, 张伟, 陈刚, 等. 电针对脑缺血大鼠星形胶质细胞活化及GFAP表达的影响[J]. 针刺研究, 2021, 46(4): 289-294.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_normal_run(p, ref)

    # 保存
    output_path = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_审稿批注版.docx'
    doc.save(output_path)
    print(f'已生成带审稿批注的修改版论文: {output_path}')
    return output_path


def add_english_run(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run


if __name__ == '__main__':
    generate_revised_paper()
