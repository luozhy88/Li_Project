#!/usr/bin/env python3
"""
替换 v3 Word 文档中的图片为新生成的带中文的图表
策略：找到包含 "图X" 文本的段落，然后查找其前面的图片段落并替换
"""
import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = 'output/paper_figures'
V3_PATH = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3.docx'
V3_OUTPUT = 'output/针灸穴位-基因关联数据库构建及数据挖掘分析_投稿稿_正式修改版_v3.docx'

FIGURE_MAP = {
    '图1': f'{OUTPUT_DIR}/fig1_dashboard.png',
    '图2': f'{OUTPUT_DIR}/fig2_top_genes.png',
    '图3': f'{OUTPUT_DIR}/fig3_top_acupoints.png',
    '图4': f'{OUTPUT_DIR}/fig4_network.png',
    '图5': f'{OUTPUT_DIR}/fig5_acupoint_query.png',
    '图6': f'{OUTPUT_DIR}/fig6_gene_query.png',
}

def has_image(para):
    return len(para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0

def replace_images_in_docx(docx_path, output_path):
    doc = Document(docx_path)
    replaced = []
    
    # 遍历所有段落，找到包含 "图X" 的段落
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        matched_key = None
        for key in FIGURE_MAP:
            if text.startswith(key):
                matched_key = key
                break
        
        if not matched_key:
            continue
        
        pic_path = FIGURE_MAP[matched_key]
        if not os.path.exists(pic_path):
            print(f'段落 {i} ({matched_key}): 图片文件不存在 {pic_path}')
            continue
        
        # 向前查找最近的包含图片的段落（最多往前5个）
        target_para = None
        for j in range(i-1, max(-1, i-6), -1):
            if has_image(doc.paragraphs[j]):
                target_para = doc.paragraphs[j]
                break
        
        if target_para is None:
            print(f'段落 {i} ({matched_key}): 未找到前面的图片段落')
            continue
        
        # 删除段落中的所有 run
        for run in target_para.runs:
            run._element.getparent().remove(run._element)
        
        # 添加新图片
        target_para.add_run().add_picture(pic_path, width=Cm(15))
        target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        replaced.append((i, matched_key, pic_path))
        print(f'段落 {i} ({matched_key}): 已替换为 {pic_path}')
    
    doc.save(output_path)
    print(f'\n已保存: {output_path}')
    print(f'共替换 {len(replaced)} 张图片')

if __name__ == '__main__':
    replace_images_in_docx(V3_PATH, V3_OUTPUT)
